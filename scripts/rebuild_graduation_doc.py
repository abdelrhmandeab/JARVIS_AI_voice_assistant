"""Build and update the Jarvis documentation book.

Phase 0 creates the thesis-style Word skeleton at
``docs/jarvis_documentation_book.docx``. Later phases can import the helper
functions in this file to replace a single ``[[FILL:<id>]]`` placeholder.
"""

from __future__ import annotations

import argparse
import ast
import json
import re
import sqlite3
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "jarvis_documentation_book.docx"
FIGURE_DIR = ROOT / "docs" / "generated_figures"

PALETTE = {
    "ink": "1a2230",
    "muted": "5b6675",
    "line": "dde3ec",
    "page": "ffffff",
    "soft": "f6f8fb",
    "accent": "1f6feb",
    "accent2": "0b3d91",
    "ok": "188a42",
    "warn": "b26a00",
    "risk": "b3261e",
}

FRONT_MATTER = [
    ("title", "Title Page"),
    ("ack", "Acknowledgments"),
    ("abstract", "Abstract"),
    ("contents", "Table of Contents"),
    ("figures", "List of Figures"),
    ("tables", "List of Tables"),
    ("abbr", "Abbreviations"),
]

CHAPTERS = [
    ("c1", "Chapter 1 Introduction"),
    ("c2", "Chapter 2 Literature Review"),
    ("c3", "Chapter 3 Methodology"),
    ("c4", "Chapter 4 Implementation - Root"),
]

IMPLEMENTATION_SECTIONS = [
    ("c4-core", "4.2 core/"),
    ("c4-memdata", "4.3 data/memory/"),
    ("c4-audio", "4.4 audio/"),
    ("c4-wakeword", "4.5 wake-word dataset and training"),
    ("c4-nlp", "4.6 nlp/"),
    ("c4-llm", "4.7 llm/"),
    ("c4-oscontrol", "4.8 os_control/"),
    ("c4-tools", "4.9 tools/ and 4.10 utils/"),
    ("c4-utils", "4.10 utils/ cross-reference placeholder"),
    ("c4-ui", "4.11 ui/ and desktop/"),
    ("c4-tests", "4.12 tests/, 4.13 scripts/, and 4.14 models/"),
]

BACK_MATTER = [
    ("c5", "Chapter 5 Conclusion & Future Work"),
    ("refs", "References"),
    ("appendix", "Appendices"),
]

FRONT_MATTER_TITLES = {title for _, title in FRONT_MATTER}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.upper())


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_paragraph_border(paragraph, color: str = PALETTE["line"], size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, field_code: str, placeholder: str) -> None:
    """Insert a Word field with fallback placeholder text."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    add_field(paragraph, "PAGE", "1")


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_placeholder(document: Document, placeholder_id: str, bookmark_id: int):
    paragraph = document.add_paragraph(style="Placeholder")
    paragraph.add_run(f"[[FILL:{placeholder_id}]]")
    add_bookmark(paragraph, placeholder_id, bookmark_id)
    return paragraph


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(12)
    normal.font.color.rgb = rgb(PALETTE["ink"])

    title = styles["Title"]
    title.font.name = "Palatino Linotype"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Palatino Linotype")
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = rgb(PALETTE["accent2"])

    for style_name, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)]:
        style = styles[style_name]
        style.font.name = "Palatino Linotype"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Palatino Linotype")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(PALETTE["accent2"] if style_name == "Heading 1" else PALETTE["ink"])
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    if "Caption" in styles:
        caption = styles["Caption"]
    else:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.bold = False
    caption.font.color.rgb = rgb(PALETTE["muted"])
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    placeholder = styles.add_style("Placeholder", WD_STYLE_TYPE.PARAGRAPH)
    placeholder.font.name = "Consolas"
    placeholder._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    placeholder.font.size = Pt(10)
    placeholder.font.color.rgb = rgb(PALETTE["warn"])
    placeholder.paragraph_format.space_before = Pt(4)
    placeholder.paragraph_format.space_after = Pt(10)

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(10)
    code.font.color.rgb = rgb(PALETTE["ink"])
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)

    for name, color in [
        ("Callout", PALETTE["accent"]),
        ("Callout OK", PALETTE["ok"]),
        ("Callout Warning", PALETTE["warn"]),
        ("Callout Risk", PALETTE["risk"]),
    ]:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.left_indent = Inches(0.15)
        style.paragraph_format.right_indent = Inches(0.15)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section.footer.paragraphs[0])


def add_styled_table(document: Document, headers: Sequence[str], rows: Iterable[Sequence[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        shade_cell(header_cells[index], PALETTE["accent2"])
        set_cell_text(header_cells[index], header, bold=True, color=PALETTE["page"])
    for row_index, row_values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_index, value in enumerate(row_values):
            if row_index % 2 == 0:
                shade_cell(cells[col_index], PALETTE["soft"])
            set_cell_text(cells[col_index], str(value))
    return table


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run(text)


def add_callout(document: Document, text: str, kind: str = "Callout") -> None:
    paragraph = document.add_paragraph(style=kind)
    paragraph.add_run(text)
    add_paragraph_border(paragraph)


def add_code_excerpt(document: Document, code: str) -> None:
    for line in code.rstrip().splitlines():
        paragraph = document.add_paragraph(style="Code Block")
        paragraph.add_run(line)


def add_diagram_image(document: Document, image_path: str | Path, caption: str, width_inches: float = 6.2) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    add_caption(document, caption)


def find_placeholder_paragraph(document: Document, placeholder_id: str):
    marker = f"[[FILL:{placeholder_id}]]"
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            return paragraph
    raise ValueError(f"Placeholder not found: {marker}")


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def iter_doc_block_elements(document: Document):
    body = document._body._element
    for child in body:
        if child.tag == qn("w:sectPr"):
            continue
        yield child


def replace_placeholder_with_document(doc_path: str | Path, placeholder_id: str, fragment: Document) -> None:
    """Replace one placeholder paragraph with all body blocks from a fragment."""
    document = Document(str(doc_path))
    placeholder = find_placeholder_paragraph(document, placeholder_id)
    placeholder_element = placeholder._element
    parent = placeholder_element.getparent()
    insert_at = parent.index(placeholder_element)
    parent.remove(placeholder_element)

    for offset, element in enumerate(list(iter_doc_block_elements(fragment))):
        parent.insert(insert_at + offset, element)

    document.save(str(doc_path))


def replace_placeholder_with_paragraphs(doc_path: str | Path, placeholder_id: str, texts: Sequence[str]) -> None:
    """Replace one visible placeholder with simple paragraphs.

    Later phases can extend this helper for richer blocks. It deliberately
    targets only the matching marker and leaves all other sections intact.
    """
    document = Document(str(doc_path))
    paragraph = find_placeholder_paragraph(document, placeholder_id)
    paragraph.text = ""
    for index, text in enumerate(texts):
        if index == 0:
            paragraph.add_run(text)
        else:
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            inserted = paragraph._parent.add_paragraph()
            inserted._p = new_p
            inserted.add_run(text)
    document.save(str(doc_path))


def _text_width(draw: "ImageDraw.ImageDraw", text: str, font) -> int:
    """Measure rendered text width so diagram boxes can be sized to fit."""
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0]


def _node_box_width(draw: "ImageDraw.ImageDraw", title: str, subtitle: str, font, small_font, min_width: int = 330, padding: int = 56) -> int:
    """Compute a box width wide enough for both the title and subtitle lines."""
    widest = max(_text_width(draw, title, font), _text_width(draw, subtitle, small_font))
    return max(min_width, widest + padding)


def add_root_layout_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_1_1_repository_layout.png"
    width, height = 1800, 1160
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 54)
        font = ImageFont.truetype("arial.ttf", 34)
        small_font = ImageFont.truetype("arial.ttf", 26)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230",
        "muted": "#5b6675",
        "line": "#dde3ec",
        "soft": "#f6f8fb",
        "accent": "#1f6feb",
        "accent2": "#0b3d91",
        "green": "#188a42",
        "amber": "#b26a00",
        "grey": "#68707c",
    }

    draw.text((70, 45), "Figure 4.1.1 - Repository layout at the root boundary", fill=colors["accent2"], font=title_font)
    draw.text((70, 112), "Phase 1 reads the root plus core/config.py and the top-level data/ tree only.", fill=colors["muted"], font=small_font)

    center = (760, 210, 1040, 310)
    draw.rounded_rectangle(center, radius=18, fill=colors["accent2"], outline=colors["accent2"])
    draw.text((825, 240), "main.py", fill="white", font=font)

    nodes = [
        ("core/", "orchestration + config", "#1f6feb", 90, 420),
        ("audio/", "wake, STT, TTS, VAD", "#1f6feb", 500, 420),
        ("nlp/", "understanding cascade", "#7c3aed", 910, 420),
        ("llm/", "answering + prompts", "#7c3aed", 1320, 420),
        ("os_control/", "verified Windows effects", "#b26a00", 90, 620),
        ("tools/", "weather, web, calculator", "#188a42", 500, 620),
        ("ui/ + desktop/", "optional tray/UI bridge", "#68707c", 910, 620),
        ("data/", "runtime artifacts", "#188a42", 1320, 620),
        ("tests/", "safety and regression checks", "#68707c", 90, 820),
        ("scripts/", "developer/training utilities", "#b26a00", 500, 820),
        ("models/", "wake-word ONNX assets", "#188a42", 910, 820),
        ("requirements*.txt", "runtime + training tiers", "#68707c", 1320, 820),
    ]

    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 330, y + 120), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 120), fill=color)
        draw.text((x + 30, y + 22), title, fill=colors["ink"], font=font)
        draw.text((x + 30, y + 70), subtitle, fill=colors["muted"], font=small_font)
        draw.line((900, 310, x + 165, y), fill=colors["line"], width=3)

    image.save(path, quality=95)
    return path


def add_core_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_2_1_core_flow.png"
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 52)
        font = ImageFont.truetype("arial.ttf", 30)
        small_font = ImageFont.truetype("arial.ttf", 23)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230",
        "muted": "#5b6675",
        "line": "#dde3ec",
        "soft": "#f6f8fb",
        "blue": "#1f6feb",
        "purple": "#7c3aed",
        "amber": "#b26a00",
        "green": "#188a42",
        "grey": "#68707c",
    }

    columns = [
        [("orchestrator.run()", "startup, wake/listen/process loop", colors["blue"]),
         ("RuntimeCoordinator", "interrupt phase gates", colors["grey"])],
        [("command_router.route_command()", "parser -> code-switch -> semantic -> keyword -> LLM/tool", colors["purple"]),
         ("memory manager/store", "fast slots + LLM context API", colors["green"])],
        [("route_verifier.verify()", "schema, slots, question guard, risk, policy", colors["amber"]),
         ("metrics/logger/doctor", "timing, structured logs, diagnostics", colors["grey"])],
        [("dispatch + handlers/", "domain handlers and OS/tool adapters", colors["amber"]),
         ("response shaping", "persona, templates, voice normalization", colors["purple"])],
    ]

    measure_image = Image.new("RGB", (10, 10), "white")
    measure_draw = ImageDraw.Draw(measure_image)
    col_widths = []
    for column in columns:
        col_width = 330
        for title, subtitle, _ in column:
            col_width = max(col_width, _node_box_width(measure_draw, title, subtitle, font, small_font))
        col_widths.append(col_width)

    gap_x, gap_y = 40, 135
    box_h = 130
    col_x = [70]
    for w in col_widths[:-1]:
        col_x.append(col_x[-1] + w + gap_x)
    width = col_x[-1] + col_widths[-1] + 70
    height = 980

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.text((70, 45), "Figure 4.2.1 - core/ internal data flow", fill=colors["ink"], font=title_font)
    draw.text((70, 110), "Runtime loop, routing gate, dispatch, memory, and observability.", fill=colors["muted"], font=small_font)

    row_y = [255, 255 + box_h + gap_y]
    centers_top = []
    centers_bottom = []
    for col_index, column in enumerate(columns):
        x = col_x[col_index]
        w = col_widths[col_index]
        for row_index, (title, subtitle, color) in enumerate(column):
            y = row_y[row_index]
            draw.rounded_rectangle((x, y, x + w, y + box_h), radius=18, fill=colors["soft"], outline=colors["line"], width=3)
            draw.rectangle((x, y, x + 12, y + box_h), fill=color)
            draw.text((x + 28, y + 25), title, fill=colors["ink"], font=font)
            draw.text((x + 28, y + 72), subtitle, fill=colors["muted"], font=small_font)
            if row_index == 0:
                centers_top.append((x, x + w))
            else:
                centers_bottom.append((x, x + w))

    for i in range(len(columns) - 1):
        y = row_y[0] + box_h // 2
        draw.line((centers_top[i][1], y, centers_top[i + 1][0], y), fill=colors["line"], width=5)
        ex = centers_top[i + 1][0]
        draw.polygon([(ex, y), (ex - 14, y - 8), (ex - 14, y + 8)], fill=colors["line"])
        y2 = row_y[1] + box_h // 2
        draw.line((centers_bottom[i][1], y2, centers_bottom[i + 1][0], y2), fill=colors["line"], width=5)
        ex2 = centers_bottom[i + 1][0]
        draw.polygon([(ex2, y2), (ex2 + 14, y2 - 8), (ex2 + 14, y2 + 8)], fill=colors["line"])

    for col_index in range(len(columns)):
        cx = col_x[col_index] + col_widths[col_index] // 2
        y_top = row_y[0] + box_h
        y_bottom = row_y[1]
        draw.line((cx, y_top, cx, y_bottom), fill=colors["line"], width=5)
        draw.polygon([(cx, y_bottom), (cx - 8, y_bottom - 14), (cx + 8, y_bottom - 14)], fill=colors["line"])

    draw.text((70, height - 160), "Phase boundary: core/ documents in-process logic. Persisted memory files are documented in Section 4.3.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def py_files_under_core() -> list[Path]:
    return sorted(
        path
        for path in (ROOT / "core").rglob("*.py")
        if "__pycache__" not in path.parts
    )


CORE_FILE_RESPONSIBILITIES = {
    "core/action_planner.py": "Plans and executes multi-step internal command sequences.",
    "core/adaptive_wake.py": "Records confirmed/false wake samples and runs background adaptive retraining hooks.",
    "core/clarification_builder.py": "Builds slot-specific and ambiguity clarification prompts.",
    "core/command_classifier.py": "NLU helper layer for extracting actions, slots, numbers, URLs, schedule payloads, and command targets.",
    "core/command_parser.py": "Regex/keyword parser that emits ParsedCommand objects and structured args.",
    "core/command_router.py": "Main routing cascade, verifier handoff, dispatch, LLM fallback, follow-up handling, and response shaping.",
    "core/config.py": "Central environment/configuration surface and data-directory setup.",
    "core/data_migration.py": "One-time migration of legacy root runtime artifacts into data/.",
    "core/demo_mode.py": "Thread-safe demo-mode flag used for presentation overlays.",
    "core/dialogue_manager.py": "Dialogue state machine for idle, follow-up, confirmation, and command phases.",
    "core/doctor.py": "Dependency, feature-tier, latency, model, and health diagnostics.",
    "core/elevation.py": "Windows Administrator/elevation detection.",
    "core/eval_routing.py": "Decision-only route harness for NLU evaluation.",
    "core/hardware_detect.py": "RAM/GPU detection and Qwen3/Faster-Whisper runtime recommendations.",
    "core/identity.py": "Varied bilingual self-introduction replies.",
    "core/intent_confidence.py": "Intent confidence, entity scoring, clarification resolution, and ambiguity safeguards.",
    "core/intent_schema.py": "Single registry of routable intents, required slots, risk tiers, and fast-execute flags.",
    "core/knowledge_base.py": "Offline knowledge-base indexing/retrieval service API.",
    "core/language_gate.py": "English/Arabic script detection and supported-language normalization.",
    "core/logger.py": "Shared logging, structured event logging, and Windows-safe rotation.",
    "core/memory_manager.py": "Fast RAM context vs richer LLM context assembly.",
    "core/memory_store.py": "SQLite and vector-memory persistence adapters used by SessionMemory.",
    "core/memory_types.py": "MemoryContext data shape shared by fast and LLM memory paths.",
    "core/metrics.py": "Latency, quality, clarification, diagnostic, and stage timing trackers.",
    "core/orchestrator.py": "Runtime startup ordering, wake/listen/process loop, concurrency, and cleanup.",
    "core/persona.py": "Active persona configuration and prompt-block formatting.",
    "core/response_shaper.py": "Voice-oriented response shortening, cleanup, and persona-safe shaping.",
    "core/response_templates.py": "Bilingual response templates for common control, clarification, and safety states.",
    "core/route_verifier.py": "Single gate that converts a routed candidate into execute/clarify/confirm/llm.",
    "core/runtime_coordinator.py": "Wake-word interrupt gate for TTS/LLM cancellation by runtime phase.",
    "core/session_memory.py": "In-process memory API over turns, slots, preferences, references, pending tasks, and persistence.",
    "core/shutdown.py": "Signal handling and graceful cleanup for TTS, queues, and background services.",
    "core/tts_prosody.py": "Deterministic TTS punctuation/prosody polishing.",
    "core/tts_voices.py": "Voice profile registry and TTS voice environment aliases.",
    "core/voice_normalizer.py": "Deterministic spoken-form normalization for numbers, dates, weather, places, and search results.",
    "core/handlers/__init__.py": "Handler package marker.",
    "core/handlers/advanced_operations.py": "Command chains, batch file operations, and semantic search handler helpers.",
    "core/handlers/audit.py": "Formats audit-log, verification, and reseal responses.",
    "core/handlers/batch.py": "Delegates batch operations to os_control batch services.",
    "core/handlers/file_navigation.py": "Handles file navigation/search/open responses and last-file memory.",
    "core/handlers/job_queue.py": "Formats and handles background job queue commands.",
    "core/handlers/knowledge_base.py": "Handles knowledge-base sync/search/status commands.",
    "core/handlers/memory.py": "Handles memory summary, preference, and stored-context commands.",
    "core/handlers/persona.py": "Handles persona status and persona-switch commands.",
    "core/handlers/policy.py": "Handles policy status/profile commands.",
    "core/handlers/search_index.py": "Handles search-index refresh/status commands.",
    "core/handlers/voice.py": "Handles runtime voice/STT/audio UX profile commands.",
}


def core_inventory_rows() -> list[tuple[str, str]]:
    rows = []
    for path in py_files_under_core():
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, CORE_FILE_RESPONSIBILITIES.get(rel, "Core Python module present in this folder.")))
    return rows


def code_excerpt(path: str, start: int, end: int) -> str:
    lines = (ROOT / path).read_text(encoding="utf-8").splitlines()
    selected = lines[start - 1 : end]
    if len(selected) > 20:
        selected = selected[:20]
    return "\n".join(selected)


def function_excerpt(path: str, function_name: str, max_lines: int = 20) -> str:
    text = (ROOT / path).read_text(encoding="utf-8")
    lines = text.splitlines()
    tree = ast.parse(text)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == function_name:
            end = min(getattr(node, "end_lineno", node.lineno + max_lines - 1), node.lineno + max_lines - 1)
            return "\n".join(lines[node.lineno - 1 : end])
    raise ValueError(f"Function {function_name} not found in {path}")


def core_config_rows() -> list[tuple[str, str, str]]:
    defaults = get_config_defaults()
    keys = [
        ("JARVIS_NLU_INTENT_ROUTING_ENABLED", "Enable the routing cascade before LLM fallback."),
        ("JARVIS_NLU_PARSER_FASTPATH_ENABLED", "Allow direct parser fast-path routing."),
        ("JARVIS_NLU_PARSER_FASTPATH_CONFIDENCE_FLOOR", "Minimum confidence floor for fast parser acceptance."),
        ("JARVIS_CODESWITCH_ROUTER_ENABLED", "Enable mixed English/Egyptian-Arabic shortcut routing."),
        ("JARVIS_SEMANTIC_ROUTER_ENABLED", "Enable embedding-based semantic routing."),
        ("JARVIS_SEMANTIC_MIN_CONFIDENCE", "Minimum semantic score for candidate acceptance."),
        ("JARVIS_SEMANTIC_MIN_MARGIN", "Minimum gap between top semantic candidate and runner-up."),
        ("JARVIS_ROUTE_VERIFIER_ENABLED", "Enable verifier logging and decision calculation."),
        ("JARVIS_FAST_COMMAND_MIN_CONFIDENCE", "Verifier/eval fast-command tuning input."),
        ("JARVIS_CLARIFY_FROM_TEMPLATES", "Use template-based clarification prompts."),
        ("JARVIS_CLARIFY_MAX_ROUNDS", "Maximum clarification turns before fallback."),
        ("JARVIS_STRUCTURED_LLM_NLU_ENABLED", "Enable gated structured-LLM NLU fallback."),
        ("JARVIS_MEMORY_FAST_CONTEXT_ENABLED", "Expose fast RAM-only context to router paths."),
        ("JARVIS_MEMORY_LLM_CONTEXT_ENABLED", "Allow recent turns and semantic recall for LLM context."),
        ("JARVIS_MEMORY_SHORT_TERM_TURNS", "Number of recent turns in LLM-bound memory context."),
        ("JARVIS_MEMORY_VECTOR_RECALL_ENABLED", "Enable bounded vector recall for LLM/uncertain routes."),
        ("JARVIS_DOCTOR_STARTUP_ASYNC", "Run startup diagnostics off the critical path."),
        ("JARVIS_STARTUP_BACKGROUND_PREWARM_ENABLED", "Load optional services in background."),
        ("JARVIS_PREWARM_SEMANTIC_ROUTER_BLOCKING", "Force semantic-router warmup onto startup path if true."),
        ("JARVIS_PREWARM_LLM_BLOCKING", "Force LLM warmup onto startup path if true."),
        ("JARVIS_FOLLOWUP_ENABLED", "Enable follow-up listening window after responses."),
        ("JARVIS_ROUTE_TIMING_LOG", "Emit per-stage route timings."),
        ("JARVIS_MEMORY_TIMING_LOG", "Emit memory timing lines."),
    ]
    return [(key, defaults.get(key, "see core/config.py"), purpose) for key, purpose in keys]


def add_pseudocode(document: Document, title: str, lines: Sequence[str]) -> None:
    document.add_paragraph(title)
    add_code_excerpt(document, "\n".join(lines))


def add_memory_data_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_3_1_memory_data_flow.png"
    width, height = 1900, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 52)
        font = ImageFont.truetype("arial.ttf", 30)
        small_font = ImageFont.truetype("arial.ttf", 23)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230",
        "muted": "#5b6675",
        "line": "#dde3ec",
        "soft": "#f6f8fb",
        "blue": "#1f6feb",
        "purple": "#7c3aed",
        "amber": "#b26a00",
        "green": "#188a42",
        "grey": "#68707c",
    }

    draw.text((70, 45), "Figure 4.3.1 - Persisted memory artifacts", fill=colors["ink"], font=title_font)
    draw.text((70, 110), "Observed files inside data/memory/: SQLite rows, JSON slots, and WAL companions.", fill=colors["muted"], font=small_font)

    nodes = [
        ("Conversation turn", "user/assistant/lang/intent", colors["blue"], 80, 280),
        ("turns table", "id, timestamp, user, assistant, language, intent", colors["green"], 510, 210),
        ("slots table", "name, JSON value, updated_at", colors["green"], 510, 450),
        ("jarvis_memory.db", "primary SQLite artifact", colors["green"], 960, 320),
        ("db-wal / db-shm", "WAL journaling sidecars", colors["grey"], 1380, 210),
        ("jarvis_memory.json", "legacy/debug mirror", colors["amber"], 1380, 450),
        ("Later prompt context", "fast slots + recent turns + language history", colors["purple"], 960, 670),
    ]
    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 360, y + 135), radius=18, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 135), fill=color)
        draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, y + 74), subtitle, fill=colors["muted"], font=small_font)

    arrows = [
        ((440, 345), (510, 278)),
        ((440, 345), (510, 515)),
        ((870, 278), (960, 385)),
        ((870, 515), (960, 385)),
        ((1320, 385), (1380, 278)),
        ((1320, 385), (1380, 515)),
        ((1140, 455), (1140, 670)),
    ]
    for start, end in arrows:
        draw.line((*start, *end), fill=colors["line"], width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])

    draw.text((80, 850), "Data-only phase note: code mechanisms were covered in Section 4.2; this chapter documents files and schemas observed on disk.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def memory_file_inventory_rows() -> list[tuple[str, str, str, str]]:
    memory_dir = ROOT / "data" / "memory"
    rows = []
    descriptions = {
        "jarvis_memory.db": ("SQLite database", "Primary observed store with turns and slots tables.", "core memory logic reads/writes it, per Section 4.2."),
        "jarvis_memory.db-shm": ("SQLite WAL shared-memory file", "SQLite sidecar used while WAL mode is active.", "SQLite runtime companion for jarvis_memory.db."),
        "jarvis_memory.db-wal": ("SQLite write-ahead log", "SQLite sidecar for pending WAL pages; observed length may be zero when checkpointed.", "SQLite runtime companion for jarvis_memory.db."),
        "jarvis_memory.json": ("JSON", "Legacy/debug mirror containing preferred_language, turns, pending_clarification, and context_slots.", "Legacy import/debug-export target, per prior core chapter."),
    }
    for path in sorted(memory_dir.iterdir(), key=lambda p: p.name.lower()):
        if path.is_dir():
            rows.append((path.name + "/", "directory", "Subfolder present in memory store.", "TODO(verify)"))
            continue
        fmt, stores, reader = descriptions.get(path.name, ("file", "Memory-related file present on disk.", "TODO(verify)"))
        rows.append((path.name, fmt, f"{stores} Size: {path.stat().st_size} bytes.", reader))
    return rows


def sqlite_memory_summary() -> dict:
    db_path = ROOT / "data" / "memory" / "jarvis_memory.db"
    summary = {
        "exists": db_path.exists(),
        "tables": [],
        "indexes": [],
        "counts": {},
        "slots": [],
        "turn_sample": [],
        "page_count": "",
        "page_size": "",
        "journal_mode": "",
    }
    if not db_path.exists():
        return summary
    db_uri = db_path.as_posix().replace("'", "''")
    con = sqlite3.connect(f"file:{db_uri}?mode=ro&immutable=1", uri=True)
    con.row_factory = sqlite3.Row
    try:
        summary["page_count"] = str(con.execute("PRAGMA page_count").fetchone()[0])
        summary["page_size"] = str(con.execute("PRAGMA page_size").fetchone()[0])
        summary["journal_mode"] = str(con.execute("PRAGMA journal_mode").fetchone()[0])
        for row in con.execute("SELECT name, type, sql FROM sqlite_master WHERE type IN ('table','index') ORDER BY type,name"):
            item = dict(row)
            if item["type"] == "table":
                summary["tables"].append(item)
            else:
                summary["indexes"].append(item)
        for table in ("turns", "slots"):
            try:
                summary["counts"][table] = int(con.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0])
            except sqlite3.Error:
                summary["counts"][table] = None
        try:
            for row in con.execute("SELECT name, updated_at, value FROM slots ORDER BY name"):
                value = row["value"]
                category = "scalar"
                try:
                    decoded = json.loads(value) if value is not None else None
                    if isinstance(decoded, dict):
                        category = f"object ({len(decoded)} keys)"
                    elif isinstance(decoded, list):
                        category = f"array ({len(decoded)} items)"
                    elif isinstance(decoded, str):
                        category = "string"
                    elif isinstance(decoded, (int, float)):
                        category = "number"
                    elif decoded is None:
                        category = "null"
                    elif isinstance(decoded, bool):
                        category = "boolean"
                except Exception:
                    category = "raw text"
                summary["slots"].append((row["name"], category, str(row["updated_at"])))
        except sqlite3.Error:
            pass
        try:
            for row in con.execute(
                "SELECT id, timestamp, language, intent, length(user) AS user_len, length(assistant) AS assistant_len "
                "FROM turns ORDER BY id DESC LIMIT 5"
            ):
                summary["turn_sample"].append(dict(row))
        except sqlite3.Error:
            pass
    finally:
        con.close()
    return summary


def json_memory_summary() -> dict:
    path = ROOT / "data" / "memory" / "jarvis_memory.json"
    summary = {"exists": path.exists(), "keys": [], "turn_count": 0, "context_slots": [], "preferred_language": ""}
    if not path.exists():
        return summary
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        summary["error"] = str(exc)
        return summary
    if isinstance(data, dict):
        summary["keys"] = sorted(data.keys())
        summary["preferred_language"] = str(data.get("preferred_language") or "")
        turns = data.get("turns")
        if isinstance(turns, list):
            summary["turn_count"] = len(turns)
        slots = data.get("context_slots")
        if isinstance(slots, dict):
            for key, value in sorted(slots.items()):
                if isinstance(value, dict):
                    category = f"object ({len(value)} keys)"
                elif isinstance(value, list):
                    category = f"array ({len(value)} items)"
                else:
                    category = type(value).__name__
                summary["context_slots"].append((key, category))
    return summary


def add_audio_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_4_1_audio_flow.png"
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 52)
        font = ImageFont.truetype("arial.ttf", 30)
        small_font = ImageFont.truetype("arial.ttf", 23)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }

    columns = [
        [("wake_word.py", "openWakeWord ONNX + EMA/peak gate", colors["blue"]),
         ("cues.py", "short wake feedback tones", colors["grey"])],
        [("mic.py / streaming_stt.py", "VAD capture + partials", colors["blue"]),
         ("vad.py", "Silero ONNX or energy fallback", colors["amber"])],
        [("stt.py", "ElevenLabs + Faster-Whisper fallback", colors["purple"]),
         ("tts.py", "ElevenLabs/Edge sentence playback", colors["green"])],
        [("core router", "text leaves audio/", colors["grey"]),
         ("wake_enrollment.py", "sample scoring + threshold recommendation", colors["amber"])],
    ]

    measure_image = Image.new("RGB", (10, 10), "white")
    measure_draw = ImageDraw.Draw(measure_image)
    col_widths = []
    for column in columns:
        col_width = 360
        for title, subtitle, _ in column:
            col_width = max(col_width, _node_box_width(measure_draw, title, subtitle, font, small_font))
        col_widths.append(col_width)

    gap_x, gap_y = 40, 135
    box_h = 135
    col_x = [80]
    for w in col_widths[:-1]:
        col_x.append(col_x[-1] + w + gap_x)
    width = col_x[-1] + col_widths[-1] + 80
    height = 980

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.text((70, 45), "Figure 4.4.1 - audio/ runtime pipeline", fill=colors["ink"], font=title_font)
    draw.text((70, 110), "Wake decision, VAD capture, STT, and TTS playback.", fill=colors["muted"], font=small_font)

    row_y = [260, 260 + box_h + gap_y]
    tops = []
    bottoms = []
    for col_index, column in enumerate(columns):
        x = col_x[col_index]
        w = col_widths[col_index]
        for row_index, (title, subtitle, color) in enumerate(column):
            y = row_y[row_index]
            draw.rounded_rectangle((x, y, x + w, y + box_h), radius=18, fill=colors["soft"], outline=colors["line"], width=3)
            draw.rectangle((x, y, x + 12, y + box_h), fill=color)
            draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
            draw.text((x + 28, y + 74), subtitle, fill=colors["muted"], font=small_font)
            (tops if row_index == 0 else bottoms).append((x, x + w))

    for i in range(len(columns) - 1):
        y = row_y[0] + box_h // 2
        draw.line((tops[i][1], y, tops[i + 1][0], y), fill=colors["line"], width=5)
        ex = tops[i + 1][0]
        draw.polygon([(ex, y), (ex - 14, y - 8), (ex - 14, y + 8)], fill=colors["line"])

    for col_index in (1, 2):
        cx = col_x[col_index] + col_widths[col_index] // 2
        y_top = row_y[0] + box_h
        y_bottom = row_y[1]
        draw.line((cx, y_bottom, cx, y_top), fill=colors["line"], width=5)
        draw.polygon([(cx, y_top), (cx - 8, y_top + 14), (cx + 8, y_top + 14)], fill=colors["line"])

    draw.text((80, height - 130), "Training data/model provenance is deferred to Section 4.5; this section covers runtime audio behavior.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_wake_training_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_5_1_wake_training_flow.png"
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 52)
        font = ImageFont.truetype("arial.ttf", 30)
        small_font = ImageFont.truetype("arial.ttf", 23)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }

    measure_image = Image.new("RGB", (10, 10), "white")
    measure_draw = ImageDraw.Draw(measure_image)
    box_h = 135

    left_nodes = [
        ("positive_train/val", "English + Arabic wake clips", colors["blue"]),
        ("negative_train/val", "background, near-phrase, false-positive clips", colors["amber"]),
    ]
    left_w = max(_node_box_width(measure_draw, t, s, font, small_font) for t, s, _ in left_nodes)

    chain_nodes = [
        ("features/*.npy", "float32 tensors (N, 41, 96)", colors["green"]),
        ("model training", "code not present in this folder", colors["grey"]),
        ("deployed ONNX", "loaded at runtime by audio/wake_word.py", colors["purple"]),
    ]
    chain_widths = [_node_box_width(measure_draw, t, s, font, small_font) for t, s, _ in chain_nodes]

    gap_x = 60
    left_x = 80
    chain_x = [left_x + left_w + 140]
    for w in chain_widths[:-1]:
        chain_x.append(chain_x[-1] + w + gap_x)
    width = chain_x[-1] + chain_widths[-1] + 80
    height = 980

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.text((70, 45), "Figure 4.5.1 - wake-word dataset/training artifacts", fill=colors["ink"], font=title_font)
    draw.text((70, 110), "Observed unified wake-word folder: WAV splits plus precomputed feature tensors.", fill=colors["muted"], font=small_font)

    left_y = [280, 560]
    for (title, subtitle, color), y in zip(left_nodes, left_y):
        draw.rounded_rectangle((left_x, y, left_x + left_w, y + box_h), radius=18, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((left_x, y, left_x + 12, y + box_h), fill=color)
        draw.text((left_x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((left_x + 28, y + 74), subtitle, fill=colors["muted"], font=small_font)

    chain_y = 420
    for i, ((title, subtitle, color), x, w) in enumerate(zip(chain_nodes, chain_x, chain_widths)):
        draw.rounded_rectangle((x, chain_y, x + w, chain_y + box_h), radius=18, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, chain_y, x + 12, chain_y + box_h), fill=color)
        draw.text((x + 28, chain_y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, chain_y + 74), subtitle, fill=colors["muted"], font=small_font)

    converge_x = chain_x[0]
    converge_y = chain_y + box_h // 2
    for y in left_y:
        draw.line((left_x + left_w, y + box_h // 2, converge_x, converge_y), fill=colors["line"], width=5)
    ex, ey = converge_x, converge_y
    draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])

    for i in range(len(chain_nodes) - 1):
        start_x = chain_x[i] + chain_widths[i]
        end_x = chain_x[i + 1]
        y = chain_y + box_h // 2
        draw.line((start_x, y, end_x, y), fill=colors["line"], width=5)
        draw.polygon([(end_x, y), (end_x - 14, y - 8), (end_x - 14, y + 8)], fill=colors["line"])

    draw.text((80, height - 130), "This phase documents data/provenance only. Runtime wake-decision logic is in Section 4.4.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def wake_training_root() -> Path:
    return ROOT / "wake word data" / "jarvis_unified_training"


def add_nlp_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_6_1_nlp_flow.png"
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 29)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }

    measure_image = Image.new("RGB", (10, 10), "white")
    measure_draw = ImageDraw.Draw(measure_image)
    box_h = 125

    chain_specs = [
        ("Recognized text", "input from audio/core", colors["grey"]),
        ("codeswitching.py", "Arabic normalization, script counts, numbers", colors["blue"]),
        ("code_switch_router.py", "verb + entity shortcut to ParsedCommand", colors["green"]),
    ]
    split_specs = [
        ("semantic_router.py", "top-k embedding similarity", colors["purple"]),
        ("fuzzy + keyword", "noisy STT keyword evidence", colors["amber"]),
    ]
    tail_specs = [
        ("nlu.py", "slot/entity enrichment + missing slots", colors["blue"]),
    ]

    chain_widths = [_node_box_width(measure_draw, t, s, font, small_font, min_width=290) for t, s, _ in chain_specs]
    split_width = max(_node_box_width(measure_draw, t, s, font, small_font, min_width=290) for t, s, _ in split_specs)
    tail_width = _node_box_width(measure_draw, tail_specs[0][0], tail_specs[0][1], font, small_font, min_width=290)

    gap_x = 50
    x = 70
    chain_x = []
    for w in chain_widths:
        chain_x.append(x)
        x += w + gap_x
    split_x = x
    x += split_width + gap_x
    tail_x = x
    width = tail_x + tail_width + 70
    height = 900

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.text((70, 45), "Figure 4.6.1 - nlp/ understanding flow", fill=colors["ink"], font=title_font)
    draw.text((70, 108), "Folder-local text understanding stages before command execution.", fill=colors["muted"], font=small_font)

    chain_y = 300
    for (title, subtitle, color), bx, bw in zip(chain_specs, chain_x, chain_widths):
        draw.rounded_rectangle((bx, chain_y, bx + bw, chain_y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((bx, chain_y, bx + 12, chain_y + box_h), fill=color)
        draw.text((bx + 28, chain_y + 24), title, fill=colors["ink"], font=font)
        draw.text((bx + 28, chain_y + 70), subtitle, fill=colors["muted"], font=small_font)
    for i in range(len(chain_specs) - 1):
        y = chain_y + box_h // 2
        start_x = chain_x[i] + chain_widths[i]
        end_x = chain_x[i + 1]
        draw.line((start_x, y, end_x, y), fill=colors["line"], width=5)
        draw.polygon([(end_x, y), (end_x - 14, y - 8), (end_x - 14, y + 8)], fill=colors["line"])

    split_y = [220, 425]
    for (title, subtitle, color), y in zip(split_specs, split_y):
        draw.rounded_rectangle((split_x, y, split_x + split_width, y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((split_x, y, split_x + 12, y + box_h), fill=color)
        draw.text((split_x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((split_x + 28, y + 70), subtitle, fill=colors["muted"], font=small_font)

    last_chain_x = chain_x[-1] + chain_widths[-1]
    last_chain_mid = chain_y + box_h // 2
    for y in split_y:
        draw.line((last_chain_x, last_chain_mid, split_x, y + box_h // 2), fill=colors["line"], width=5)
        ex, ey = split_x, y + box_h // 2
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])

    tail_y = 300
    draw.rounded_rectangle((tail_x, tail_y, tail_x + tail_width, tail_y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
    draw.rectangle((tail_x, tail_y, tail_x + 12, tail_y + box_h), fill=tail_specs[0][2])
    draw.text((tail_x + 28, tail_y + 24), tail_specs[0][0], fill=colors["ink"], font=font)
    draw.text((tail_x + 28, tail_y + 70), tail_specs[0][1], fill=colors["muted"], font=small_font)

    for y in split_y:
        draw.line((split_x + split_width, y + box_h // 2, tail_x, tail_y + box_h // 2), fill=colors["line"], width=5)
        ex, ey = tail_x, tail_y + box_h // 2
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])

    callout_left = chain_x[1]
    callout_right = tail_x + tail_width
    draw.rounded_rectangle((callout_left, 650, callout_right, 755), radius=14, fill="#ffffff", outline=colors["line"], width=3)
    draw.text((callout_left + 35, 682), "Acceptance rule used by caller: best >= tau and (best - second) >= delta; near ties defer.", fill=colors["ink"], font=font)
    draw.text((70, height - 85), "No standalone nlp/text_normalizer.py is present; normalization evidence is in codeswitching.py and fuzzy_matcher.py.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_llm_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_7_1_llm_flow.png"
    width, height = 1900, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 29)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((70, 45), "Figure 4.7.1 - llm/ answer and streaming flow", fill=colors["ink"], font=title_font)
    draw.text((70, 108), "Prompt package, local model stream, sentence chunking, and optional structured/tool fallback.", fill=colors["muted"], font=small_font)
    nodes = [
        ("prompt_builder.py", "persona + language + memory + KB", colors["blue"], 70, 270),
        ("prompts/*.txt", "micro, slim, full templates", colors["grey"], 70, 520),
        ("ollama_client.py", "Ollama generate payload + stream", colors["purple"], 480, 270),
        ("sentence_buffer.py", "EN/AR speakable chunks", colors["green"], 900, 270),
        ("TTS callback", "on_sentence(text)", colors["green"], 1320, 270),
        ("structured_nlu.py", "strict JSON command fallback", colors["amber"], 480, 560),
        ("tool_caller.py", "tool-call to ParsedCommand bridge", colors["amber"], 900, 560),
    ]
    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 330, y + 125), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 125), fill=color)
        draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, y + 70), subtitle, fill=colors["muted"], font=small_font)
    for start, end in [((400, 332), (480, 332)), ((810, 332), (900, 332)), ((1230, 332), (1320, 332)), ((235, 520), (235, 395)), ((810, 622), (900, 622)), ((645, 395), (645, 560))]:
        draw.line((*start, *end), fill=colors["line"], width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])
    draw.rounded_rectangle((1290, 545, 1710, 690), radius=16, fill="#ffffff", outline=colors["line"], width=3)
    draw.text((1320, 575), "Interrupt/cancel", fill=colors["ink"], font=font)
    draw.text((1320, 620), "cancel_event is checked per streamed chunk", fill=colors["muted"], font=small_font)
    draw.line((1320, 545, 760, 395), fill=colors["line"], width=4)
    draw.text((70, 815), "Prompt memory contents are cross-referenced to Section 4.3; this phase documents only how llm/ inserts the block.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_os_control_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_8_1_os_control_flow.png"
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 28)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "red": "#b3261e",
        "grey": "#68707c",
    }

    measure_image = Image.new("RGB", (10, 10), "white")
    measure_draw = ImageDraw.Draw(measure_image)
    box_h = 130

    # Columns 0-3 have only a top-row node; column 4 has both a top and bottom node.
    top_row = [
        ("Routed action", "system/app/file/control request", colors["grey"]),
        ("policy.py + risk_policy.py", "allowed? risk tier?", colors["amber"]),
        ("confirmation.py", "pending action + PIN if needed", colors["red"]),
        ("adapter modules", "native, PowerShell, WinRT, Explorer", colors["blue"]),
        ("verify OS state", "readback/poll/window check", colors["green"]),
    ]
    bottom_pair = [
        ("action_log.py + persistence.py", "hash-chained audit + rollback", colors["grey"]),
        ("adapter_result.py", "success/failure/confirm shape", colors["purple"]),
    ]

    col_widths = [_node_box_width(measure_draw, t, s, font, small_font, min_width=300) for t, s, _ in top_row]
    # The bottom row spans under columns 3-4; make sure those two columns are wide enough
    # for the bottom labels too, since the bottom pair sits directly beneath them.
    bottom_widths = [_node_box_width(measure_draw, t, s, font, small_font, min_width=300) for t, s, _ in bottom_pair]
    col_widths[3] = max(col_widths[3], bottom_widths[0])
    col_widths[4] = max(col_widths[4], bottom_widths[1])

    gap_x = 50
    col_x = [70]
    for w in col_widths[:-1]:
        col_x.append(col_x[-1] + w + gap_x)
    width = col_x[-1] + col_widths[-1] + 70
    height = 980

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.text((70, 45), "Figure 4.8.1 - os_control/ verified side-effect flow", fill=colors["ink"], font=title_font)
    draw.text((70, 108), "Routed action, policy/risk, confirmation/PIN, adapter execution, state verification, and audit logging.", fill=colors["muted"], font=small_font)

    top_y = 315
    for (title, subtitle, color), x, w in zip(top_row, col_x, col_widths):
        draw.rounded_rectangle((x, top_y, x + w, top_y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, top_y, x + 12, top_y + box_h), fill=color)
        draw.text((x + 28, top_y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, top_y + 72), subtitle, fill=colors["muted"], font=small_font)

    for i in range(len(top_row) - 1):
        y = top_y + box_h // 2
        start_x = col_x[i] + col_widths[i]
        end_x = col_x[i + 1]
        draw.line((start_x, y, end_x, y), fill=colors["line"], width=5)
        draw.polygon([(end_x, y), (end_x - 14, y - 8), (end_x - 14, y + 8)], fill=colors["line"])

    bottom_y = 585
    bottom_x = [col_x[3], col_x[4]]
    bottom_w = [col_widths[3], col_widths[4]]
    for (title, subtitle, color), x, w in zip(bottom_pair, bottom_x, bottom_w):
        draw.rounded_rectangle((x, bottom_y, x + w, bottom_y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, bottom_y, x + 12, bottom_y + box_h), fill=color)
        draw.text((x + 28, bottom_y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, bottom_y + 72), subtitle, fill=colors["muted"], font=small_font)

    # verify OS state -> adapter_result.py (vertical, column 4)
    cx4 = col_x[4] + col_widths[4] // 2
    draw.line((cx4, top_y + box_h, cx4, bottom_y), fill=colors["line"], width=5)
    draw.polygon([(cx4, bottom_y), (cx4 - 8, bottom_y - 14), (cx4 + 8, bottom_y - 14)], fill=colors["line"])

    # adapter_result.py -> action_log.py + persistence.py (horizontal, bottom row, right to left)
    by = bottom_y + box_h // 2
    draw.line((bottom_x[1], by, bottom_x[0] + bottom_w[0], by), fill=colors["line"], width=5)
    ex = bottom_x[0] + bottom_w[0]
    draw.polygon([(ex, by), (ex + 14, by - 8), (ex + 14, by + 8)], fill=colors["line"])

    # action_log.py + persistence.py -> adapter modules (vertical, column 3)
    cx3 = col_x[3] + col_widths[3] // 2
    draw.line((cx3, bottom_y, cx3, top_y + box_h), fill=colors["line"], width=5)
    draw.polygon([(cx3, top_y + box_h), (cx3 - 8, top_y + box_h - 14), (cx3 + 8, top_y + box_h - 14)], fill=colors["line"])

    draw.rounded_rectangle((70, 760, width - 70, 870), radius=14, fill="#ffffff", outline=colors["line"], width=3)
    draw.text((105, 790), "Contract: an OS action should report success only when the adapter returns success and, where implemented, read-back verification confirms the state.", fill=colors["ink"], font=font)
    draw.text((105, 835), "Degraded paths return explicit failure/admin/manual-settings messages instead of claiming success.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_tools_flow_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_9_1_tools_live_data_flow.png"
    width, height = 1900, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 29)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((70, 45), "Figure 4.9.1 - tools/ live data flow", fill=colors["ink"], font=title_font)
    draw.text((70, 108), "Query detection, tool execution, raw result handling, voice-ready normalization, and prompt context.", fill=colors["muted"], font=small_font)
    nodes = [
        ("User query", "weather/search/math request", colors["grey"], 80, 325),
        ("live_data.py", "detect intent + choose tool", colors["blue"], 430, 325),
        ("weather.py", "Open-Meteo current weather", colors["green"], 800, 230),
        ("web_search.py", "DuckDuckGo ranked snippets", colors["purple"], 800, 430),
        ("raw result", "empty string on failure", colors["amber"], 1160, 325),
        ("voice-ready block", "[WEATHER] / [WEB_SEARCH]", colors["blue"], 1500, 325),
    ]
    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 300, y + 125), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 125), fill=color)
        draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, y + 70), subtitle, fill=colors["muted"], font=small_font)
    for start, end in [((380, 388), (430, 388)), ((730, 360), (800, 292)), ((730, 415), (800, 492)), ((1100, 292), (1160, 360)), ((1100, 492), (1160, 415)), ((1460, 388), (1500, 388))]:
        draw.line((*start, *end), fill=colors["line"], width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])
    draw.rounded_rectangle((70, 675, width - 70, 800), radius=14, fill="#ffffff", outline=colors["line"], width=3)
    draw.text((105, 700), "calculator.py is a separate fast path: math-looking text -> safe expression", fill=colors["ink"], font=font)
    draw.text((105, 745), "cleanup -> minimal eval namespace -> formatted number.", fill=colors["ink"], font=font)
    draw.text((70, 850), "Voice normalizer internals are outside this phase; tools/live_data.py only calls the normalizer before returning blocks.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_ui_bridge_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_11_1_ui_bridge_protocol.png"
    width, height = 1900, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 29)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((70, 45), "Figure 4.11.1 - UI bridge protocol", fill=colors["ink"], font=title_font)
    draw.text((70, 108), "Optional Python bridge between the headless engine and Tauri/React desktop windows.", fill=colors["muted"], font=small_font)
    nodes = [
        ("Engine", "dialogue state + route_command", colors["blue"], 80, 330),
        ("ui/events.py", "event/command names + JSON", colors["grey"], 430, 210),
        ("ui/bridge.py", "FastAPI WebSocket /ws", colors["purple"], 430, 430),
        ("useJarvisSocket", "connect, dispatch, reconnect", colors["green"], 820, 430),
        ("Zustand store", "state, config, transcripts", colors["green"], 1180, 430),
        ("Tauri windows", "overlay + dashboard", colors["amber"], 1500, 330),
    ]
    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 310, y + 125), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 125), fill=color)
        draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, y + 70), subtitle, fill=colors["muted"], font=small_font)
    arrows = [
        ((390, 390), (430, 492)),  # engine to bridge
        ((740, 492), (820, 492)),
        ((1130, 492), (1180, 492)),
        ((1490, 492), (1500, 392)),
        ((1500, 360), (740, 245)),  # commands back through protocol naming
        ((585, 335), (585, 430)),
    ]
    for start, end in arrows:
        draw.line((*start, *end), fill=colors["line"], width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])
    draw.rounded_rectangle((245, 680, 1660, 780), radius=14, fill="#ffffff", outline=colors["line"], width=3)
    draw.text((280, 710), "Events out: state_changed, transcripts, response, amplitude, metrics, health, error, config.", fill=colors["ink"], font=font)
    draw.text((280, 750), "Commands in: text_command, mute_toggle, config_request, health_request, setting_update, feature_flag.", fill=colors["muted"], font=small_font)
    draw.text((70, 830), "Headless guarantee: bridge/tray startup is optional; missing UI dependencies or absent desktop app do not stop the engine.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_qa_assets_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_4_12_1_qa_assets_flow.png"
    width, height = 1900, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 29)
        small_font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()

    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((70, 45), "Figure 4.12.1 - QA, scripts, and model assets", fill=colors["ink"], font=title_font)
    draw.text((70, 108), "Repository support assets that prove routing, memory, voice text shaping, setup, and wake-model packaging.", fill=colors["muted"], font=small_font)
    nodes = [
        ("tests/", "unittest safety + behavior checks", colors["blue"], 80, 260),
        ("fixtures/", "154 labeled NLU eval cases", colors["blue"], 80, 470),
        ("pytest / unittest", "local suite execution", colors["purple"], 480, 365),
        ("core + nlp + llm", "routing, memory, voice utilities", colors["green"], 860, 260),
        ("scripts/", "doc build, setup, wake tooling", colors["amber"], 860, 470),
        ("models/", "jarvis_unified ONNX assets", colors["grey"], 1260, 365),
        ("runtime docs", "Phase 11 section + figure", colors["purple"], 1600, 365),
    ]
    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 310, y + 125), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 125), fill=color)
        draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, y + 70), subtitle, fill=colors["muted"], font=small_font)
    arrows = [
        ((390, 322), (480, 410)),
        ((390, 532), (480, 455)),
        ((790, 425), (860, 322)),
        ((790, 425), (860, 532)),
        ((1170, 532), (1260, 455)),
        ((1570, 425), (1600, 425)),
        ((1170, 322), (1260, 410)),
    ]
    for start, end in arrows:
        draw.line((*start, *end), fill=colors["line"], width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=colors["line"])
    draw.rounded_rectangle((210, 690, 1690, 790), radius=14, fill="#ffffff", outline=colors["line"], width=3)
    draw.text((245, 720), "Hard QA targets in this repository: no unsafe eval executions, no question-to-command false fires, and no near-tie auto-execution.", fill=colors["ink"], font=font)
    draw.text((245, 760), "Wake data/training mechanics are cross-referenced to Section 4.5; this phase inventories the script and model assets only.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def wake_training_inventory_rows() -> list[tuple[str, str, str, str]]:
    root = wake_training_root()
    rows = []
    descriptions = {
        "features": ("directory", "Precomputed NumPy feature tensors.", "Consumes WAV splits; feeds training/evaluation code not present in this folder."),
        "positive_train": ("directory", "Positive wake-word training WAV clips.", "Source class for positive_features_train.npy."),
        "positive_val": ("directory", "Positive validation/test WAV clips.", "Source class for positive_features_test.npy."),
        "negative_train": ("directory", "Negative/background training WAV clips.", "Source class for negative_features_train.npy."),
        "negative_val": ("directory", "Negative validation/test WAV clips.", "Source class for negative_features_test.npy."),
    }
    for child in sorted(root.iterdir(), key=lambda p: p.name.lower()):
        if child.is_dir():
            file_count = len([p for p in child.rglob("*") if p.is_file()])
            fmt, stores, role = descriptions.get(child.name, ("directory", "Dataset subfolder.", "Observed dataset artifact."))
            rows.append((child.name + "/", fmt, f"{stores} Files: {file_count}.", role))
        else:
            rows.append((child.name, child.suffix.lstrip(".") or "file", f"Size: {child.stat().st_size} bytes.", "Direct file in training folder."))
    missing = []
    for name in ("eval", "training"):
        if not (root / name).exists():
            missing.append(name + "/")
    if missing:
        rows.append((", ".join(missing), "absent", "No separate eval/ or training/ code folder exists inside jarvis_unified_training.", "Documented as absent rather than inferred."))
    return rows


def wake_feature_rows() -> list[tuple[str, str, str, str]]:
    rows = []
    try:
        import numpy as np
    except Exception:
        np = None
    features = wake_training_root() / "features"
    for path in sorted(features.glob("*.npy")):
        shape = "unknown"
        dtype = "unknown"
        if np is not None:
            arr = np.load(path, mmap_mode="r")
            shape = str(tuple(arr.shape))
            dtype = str(arr.dtype)
        label = "positive" if path.name.startswith("positive") else "negative"
        split = "train" if "train" in path.name else "test/validation"
        rows.append((path.name, label, split, f"{shape}, {dtype}, {path.stat().st_size} bytes"))
    return rows


def wake_wav_split_rows() -> list[tuple[str, str, str, str, str]]:
    rows = []
    root = wake_training_root()
    for sub in ("positive_train", "positive_val", "negative_train", "negative_val"):
        files = list((root / sub).rglob("*.wav"))
        ascii_count = sum(1 for f in files if all(ord(ch) < 128 for ch in f.name))
        non_ascii_count = len(files) - ascii_count
        total_bytes = sum(f.stat().st_size for f in files)
        rows.append((sub + "/", str(len(files)), str(ascii_count), str(non_ascii_count), str(total_bytes)))
    return rows


def wake_wav_sample_rows() -> list[tuple[str, str, str, str, str]]:
    import contextlib
    import wave
    rows = []
    root = wake_training_root()
    for sub in ("positive_train", "positive_val", "negative_train", "negative_val"):
        sample = next(iter(sorted((root / sub).rglob("*.wav"))), None)
        if sample is None:
            rows.append((sub, "none", "", "", ""))
            continue
        try:
            with contextlib.closing(wave.open(str(sample), "rb")) as wav:
                rows.append((sub, sample.name[:64], str(wav.getnchannels()), str(wav.getframerate()), f"{wav.getnframes() / float(wav.getframerate()):.3f}s"))
        except Exception as exc:
            rows.append((sub, sample.name[:64], "error", "error", str(exc)[:80]))
    return rows


AUDIO_FILE_RESPONSIBILITIES = {
    "audio/cues.py": "Generates short cached sine-wave cues and plays them through sounddevice when available.",
    "audio/mic.py": "Records utterances with VAD, pre-roll, adaptive silence, start timeout, and WAV output.",
    "audio/streaming_stt.py": "Records with VAD while producing partial STT windows and Arabic stability gating.",
    "audio/stt.py": "Owns STT backend selection, ElevenLabs cloud client, Faster-Whisper fallback, language lock, validation, and metadata.",
    "audio/tts.py": "SpeechEngine for hybrid ElevenLabs/Edge-TTS/console playback, sentence queue streaming, voice prep, and interruption.",
    "audio/vad.py": "Silero VAD ONNX wrapper with energy fallback and batch/streaming helpers.",
    "audio/wake_enrollment.py": "Records/scans user wake samples and prints threshold recommendations.",
    "audio/wake_word.py": "Unified bilingual wake-word listener using openWakeWord ONNX, EMA/peak/confirm gates, cooldown, and interrupt integration.",
}


def audio_inventory_rows() -> list[tuple[str, str]]:
    rows = []
    for path in sorted((ROOT / "audio").glob("*.py")):
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, AUDIO_FILE_RESPONSIBILITIES.get(rel, "Audio source file.")))
    return rows


def audio_config_rows() -> list[tuple[str, str]]:
    return [
        ("Wake word", "WAKE_WORD_UNIFIED_ONNX_PATH, WAKE_WORD_THRESHOLD, WAKE_WORD_PEAK_THRESHOLD, WAKE_WORD_CONFIRM_FRAMES, WAKE_WORD_EMA_WINDOW, WAKE_WORD_CHUNK_SIZE, WAKE_WORD_MIN_RMS, WAKE_WORD_DETECTION_COOLDOWN_SECONDS, WAKE_WORD_RECORD_START_DELAY_MS, WAKE_WORD_INPUT_DEVICE, WAKE_WORD_USER_SAMPLES_DIR."),
        ("Recording/VAD", "MAX_RECORD_DURATION, VAD_ENERGY_THRESHOLD, VAD_BACKEND, VAD_SILERO_THRESHOLD, VAD_COMMAND_SILENCE_SECONDS, VAD_CHAT_SILENCE_SECONDS, VAD_MIN_SPEECH_SECONDS, VAD_PREROLL_SECONDS, VAD_START_TIMEOUT_SECONDS."),
        ("STT", "STT_BACKEND, ELEVENLABS_API_KEY, ELEVENLABS_BASE_URL, STT_LANGUAGE_HINT, STT_LANGUAGE_LOCK, STT_FORBID_OTHER_LANGUAGES, STT_MIN_CONFIDENCE, STT_ELEVENLABS_* timeout/cooldown/model flags, WHISPER_MODEL, WHISPER_DEVICE, WHISPER_COMPUTE_TYPE."),
        ("Streaming partials", "STT_PARTIAL_MIN_SECONDS, STT_PARTIAL_WINDOW_SECONDS, STT_PARTIAL_INTERVAL_SECONDS, STT_PARTIAL_WHISPER_MODEL."),
        ("TTS", "TTS_DEFAULT_BACKEND, TTS_QUALITY_MODE, TTS_SENTENCE_STREAMING_ENABLED, TTS_SENTENCE_SYNTH_WORKERS, TTS_SENTENCE_FIRST_FLUSH_MIN_CHARS, TTS_SENTENCE_GAP_MS, TTS_PARAGRAPH_GAP_MS, TTS_ELEVENLABS_TIMEOUT_SECONDS, TTS_EDGE_MIXED_SCRIPT_CHUNKING."),
        ("Cues/interrupts", "WAKE_INTERRUPT_ACK_SOUND, WAKE_INTERRUPT_ACK_FREQ_HZ, WAKE_INTERRUPT_ACK_DURATION_MS, WAKE_INTERRUPT_BLOCKED_TONE_ENABLED, WAKE_INTERRUPT_BLOCKED_TONE_FREQ_HZ, WAKE_INTERRUPT_BLOCKED_TONE_DURATION_MS."),
    ]


def add_phase4_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "audio/ is the voice I/O boundary. It detects the bilingual wake word, captures speech with VAD, transcribes audio "
        "with cloud/local STT paths, streams partial transcripts for early routing, and speaks responses through hybrid TTS. "
        "The wake-word model training data is intentionally deferred to Section 4.5."
    )
    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], audio_inventory_rows())
    add_caption(document, "Table 4.4.1: Inventory of every source file in audio/.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_audio_flow_diagram(), "Figure 4.4.1: Audio runtime pipeline from wake detection through STT and TTS.")

    document.add_heading("Key Modules In Depth", level=3)
    document.add_paragraph("wake_word.py loads one unified openWakeWord-compatible ONNX model and applies a two-level decision gate: EMA-smoothed score plus raw peak score within the confirm window.")
    add_code_excerpt(document, code_excerpt("audio/wake_word.py", 220, 241))
    add_caption(document, "Listing 4.4.1: Unified wake-word model loading.")
    add_code_excerpt(document, code_excerpt("audio/wake_word.py", 301, 318))
    add_caption(document, "Listing 4.4.2: Runtime wake thresholds and EMA setup.")
    add_code_excerpt(document, code_excerpt("audio/wake_word.py", 401, 425))
    add_caption(document, "Listing 4.4.3: EMA, peak, confirm-frame, cooldown, and interrupt gate.")

    document.add_paragraph("mic.py and streaming_stt.py capture WAV audio with pre-roll, start timeout, adaptive silence, and optional partial transcription windows.")
    add_code_excerpt(document, code_excerpt("audio/mic.py", 132, 170))
    add_caption(document, "Listing 4.4.4: VAD recording setup.")
    add_code_excerpt(document, code_excerpt("audio/streaming_stt.py", 287, 305))
    add_caption(document, "Listing 4.4.5: Arabic partial-transcript stability gate.")

    document.add_paragraph("stt.py keeps runtime STT settings and metadata, builds the ElevenLabs HTTP client, validates language, and falls back to local Faster-Whisper when needed.")
    add_code_excerpt(document, code_excerpt("audio/stt.py", 90, 110))
    add_caption(document, "Listing 4.4.6: Reused ElevenLabs HTTP client with timeout settings.")
    add_code_excerpt(document, code_excerpt("audio/stt.py", 145, 150))
    add_caption(document, "Listing 4.4.7: Runtime STT language hint normalization.")

    document.add_paragraph("vad.py wraps Silero ONNX when available and rejects low-energy chunks before inference; if the model/runtime is missing, energy gating remains available.")
    add_code_excerpt(document, code_excerpt("audio/vad.py", 101, 126))
    add_caption(document, "Listing 4.4.8: Silero VAD initialization and fallback state.")
    add_code_excerpt(document, code_excerpt("audio/vad.py", 165, 176))
    add_caption(document, "Listing 4.4.9: Streaming VAD chunk decision.")

    document.add_paragraph("tts.py speaks through SpeechEngine. Hybrid mode tries ElevenLabs first, then Edge-TTS, then console fallback; sentence-queue playback overlaps synthesis and playback for lower perceived latency.")
    add_code_excerpt(document, code_excerpt("audio/tts.py", 853, 879))
    add_caption(document, "Listing 4.4.10: Sentence queue streaming worker.")
    add_code_excerpt(document, code_excerpt("audio/tts.py", 1288, 1308))
    add_caption(document, "Listing 4.4.11: Hybrid TTS fallback order.")

    document.add_paragraph("wake_enrollment.py is a runtime support tool: it scores recorded samples against the deployed wake model and recommends a threshold based on the quietest usable sample.")
    add_code_excerpt(document, code_excerpt("audio/wake_enrollment.py", 130, 167))
    add_caption(document, "Listing 4.4.12: Wake enrollment threshold recommendation.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(document, "Wake-word decision:", [
        "load unified ONNX model",
        "for each audio chunk:",
        "    score = model.predict(chunk)",
        "    ema_score = alpha * score + (1-alpha) * previous_ema",
        "    frame_ok = ema_score >= threshold and rms >= min_rms",
        "    require confirm_frames consecutive frame_ok chunks",
        "    require max(raw_scores in confirm window) >= peak_threshold",
        "    require cooldown has elapsed",
        "    save recent audio and, if speaking/thinking, request interrupt",
    ])
    add_pseudocode(document, "STT language lock and fallback:", [
        "normalize runtime language hint to ar/en/auto",
        "transcribe with configured backend",
        "validate transcript against allowed English/Egyptian-Arabic scripts and confidence floors",
        "on cloud failure or weak result, use local Faster-Whisper fallback when configured",
        "record backend/language/confidence metadata for diagnostics",
    ])
    add_pseudocode(document, "Sentence-level TTS streaming:", [
        "iterate sentences from LLM sentence buffer",
        "submit synthesis futures to a bounded worker pool",
        "as soon as the next future is ready, play waveform",
        "insert sentence or paragraph gap samples",
        "allow interrupt() to stop queued playback",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Area", "Config constants used by audio/"], audio_config_rows())
    add_caption(document, "Table 4.4.2: Audio configuration surface imported by audio/ modules; defaults are centralized in core/config.py.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "Audio behavior", "Evidence"], [
        ("Normal", "Wake fires after EMA/peak/confirm/cooldown gates, recording captures speech until adaptive silence, STT returns text, and TTS streams sentences.", "wake_word.py, mic.py, streaming_stt.py, stt.py, tts.py."),
        ("Degraded", "Missing openwakeword/sounddevice raises honest runtime errors; missing Silero/onnxruntime falls back to energy VAD; ElevenLabs/Edge failures fall back to local or console paths.", "wake_word.py, mic.py, vad.py, stt.py, tts.py."),
        ("Adversarial/Edge", "Low-RMS noise cannot pass wake gate; sustained mild scores must also clear peak threshold; Arabic partials wait for stable repeated text; TTS can be interrupted mid-stream.", "wake_word.py, streaming_stt.py, tts.py."),
    ])
    add_caption(document, "Table 4.4.3: Audio behavior under normal, degraded, and edge conditions.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "audio/ underpins hands-free operation by turning microphone input into verified bilingual text and assistant text "
        "back into speech. Its degradation paths keep failures explicit, while the wake gate and language-aware partials reduce false activation and unstable early execution."
    )


def fill_phase4(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-audio", add_phase4_content)


def add_phase5_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "The wake-word dataset/training artifacts document the offline evidence behind the unified on-device wake model. "
        "In this workspace the deployed-relevant folder is wake word data/jarvis_unified_training. It contains class-split "
        "WAV clips and precomputed feature tensors, but no separate eval/ or training/ source-code folder inside that target."
    )

    document.add_heading("File and Subfolder Inventory", level=3)
    add_styled_table(document, ["Artifact", "Format", "What it stores", "Pipeline role"], wake_training_inventory_rows())
    add_caption(document, "Table 4.5.1: Observed artifacts in wake word data/jarvis_unified_training.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_wake_training_flow_diagram(), "Figure 4.5.1: Wake-word data artifacts from WAV splits to feature tensors and deployed ONNX boundary.")

    document.add_heading("Key Artifacts In Depth", level=3)
    document.add_paragraph(
        "features/ contains four NumPy tensors. Every observed tensor uses float32 values and a per-sample shape of 41 by 96, "
        "which indicates fixed-size extracted wake-word features rather than raw waveform storage."
    )
    add_styled_table(document, ["Feature file", "Class", "Split", "Observed shape/dtype/size"], wake_feature_rows())
    add_caption(document, "Table 4.5.2: Precomputed feature tensors under features/.")

    document.add_paragraph(
        "The WAV split folders separate positive wake examples from negative/background examples and train from validation/test data. "
        "Filenames show both ASCII and non-ASCII names, reflecting English and Arabic phrase coverage without needing to open audio content."
    )
    add_styled_table(document, ["Split folder", "WAV files", "ASCII filenames", "Non-ASCII filenames", "Total bytes"], wake_wav_split_rows())
    add_caption(document, "Table 4.5.3: WAV class/split counts in the unified wake-word dataset.")

    document.add_paragraph(
        "A small sample of WAV headers confirms mono 16 kHz clips, with positive examples around 2.48 seconds and negative background clips "
        "commonly longer. The table uses headers only, not audio playback or transcript assumptions."
    )
    add_styled_table(document, ["Split", "Sample file", "Channels", "Rate", "Duration"], wake_wav_sample_rows())
    add_caption(document, "Table 4.5.4: Header-level sample metadata from each WAV split.")

    document.add_paragraph(
        "No Python training loop, loss definition, threshold-selection report, or held-out metric file is present inside "
        "jarvis_unified_training. The folder therefore proves the existence and shape of the dataset/features, while the exact "
        "training/evaluation implementation remains outside this phase's read boundary."
    )

    document.add_heading("Algorithms Reflected By The Data", level=3)
    add_pseudocode(document, "Observed feature-preparation procedure:", [
        "collect positive wake clips under positive_train/ and positive_val/",
        "collect negative/background clips under negative_train/ and negative_val/",
        "extract fixed-size features from each WAV clip",
        "write class/split arrays:",
        "    positive_features_train.npy",
        "    positive_features_test.npy",
        "    negative_features_train.npy",
        "    negative_features_test.npy",
        "each row has feature shape 41 x 96",
    ])
    add_pseudocode(document, "Observed train/validation balance:", [
        "positive train features: 3090",
        "positive test features: 858",
        "negative train features: 7173",
        "negative test features: 1726",
        "negative examples outnumber positive examples in both train and test splits",
        "this supports false-wake resistance, but exact loss/threshold tuning is not present in this folder",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Surface", "Observed value", "Meaning"], [
        ("Dataset root", "wake word data/jarvis_unified_training", "Unified wake-word dataset/training artifact folder documented in this phase."),
        ("Feature format", "*.npy float32 tensors", "Precomputed features consumed by training/evaluation tooling."),
        ("Raw audio format", "*.wav, mono 16 kHz in inspected samples", "Source clips for positive/negative class splits."),
        ("Runtime deployment link", "deployed ONNX model path is documented in Section 4.4", "This phase documents data provenance; runtime loading is separate."),
        ("Absent folders", "eval/ and training/ are not present inside this target", "No local metric reports or training scripts are claimed from this folder."),
    ])
    add_caption(document, "Table 4.5.5: Dataset/training configuration surface visible from the artifact folder.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "Artifact behavior", "Defense note"], [
        ("Normal", "Positive/negative train and validation WAV splits exist, and features/ contains matching class/split tensors.", "Enough provenance is present to discuss data balance and feature shapes."),
        ("Degraded", "No eval/ or training/ code folder is present inside jarvis_unified_training.", "Training loop and metric claims are intentionally marked as not present rather than inferred."),
        ("Adversarial/Edge", "Negative feature rows and WAV clips outnumber positive examples; filenames include background, ambient, near-phrase, and false-positive-style names.", "The data design visibly targets false-wake resistance."),
        ("Reproducibility", "Raw WAV and feature tensors are present, but no manifest/checksum/metric report was found directly in the target folder.", "A later improvement would add a manifest and evaluation report next to the data."),
    ])
    add_caption(document, "Table 4.5.6: Wake-word dataset behavior under normal, degraded, and edge conditions.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "The unified wake-word data artifacts support the hands-free objective by pairing English and Arabic positive examples with a larger negative/background set. "
        "The precomputed feature tensors make the data pipeline inspectable, while the absence of local training/eval code in this folder is an explicit reproducibility limitation to address before final archival."
    )


def fill_phase5(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-wakeword", add_phase5_content)


NLP_FILE_RESPONSIBILITIES = {
    "nlp/__init__.py": "Package marker documenting NLP utilities for robust intent understanding on noisy STT output.",
    "nlp/codeswitching.py": "Mixed Arabic/English normalization, verb/entity maps, script counts, Arabic numerals, and code-switched intent hints.",
    "nlp/code_switch_router.py": "Fast code-switch shortcut that maps verb+entity phrases to ParsedCommand objects before semantic routing.",
    "nlp/semantic_router.py": "Sentence-transformers semantic route catalog and top-k confidence scoring.",
    "nlp/fuzzy_matcher.py": "Arabic/English text normalization plus exact/fuzzy keyword scoring for noisy STT text.",
    "nlp/keyword_engine.py": "Bilingual keyword catalog grouped by intent with action and target terms.",
    "nlp/intent_classifier.py": "Rule-based classifier over keyword/fuzzy matches with confidence, tie handling, and suggestions.",
    "nlp/nlu.py": "Entity extraction, slot enrichment, required-slot detection, and NLUResult packaging.",
    "nlp/entity_types.py": "Shared EntityType enum and value-to-enum lookup map.",
}


def nlp_inventory_rows() -> list[tuple[str, str]]:
    rows = []
    for path in sorted((ROOT / "nlp").glob("*.py"), key=lambda p: p.name.lower()):
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, NLP_FILE_RESPONSIBILITIES.get(rel, "NLP Python module present in this folder.")))
    missing = ROOT / "nlp" / "text_normalizer.py"
    if not missing.exists():
        rows.append(("nlp/text_normalizer.py", "Not present in this workspace; normalization is documented from codeswitching.py and fuzzy_matcher.py instead."))
    return rows


def nlp_config_rows() -> list[tuple[str, str, str]]:
    return [
        ("SEMANTIC_CONFIDENCE_THRESHOLD", "0.75 in nlp/semantic_router.py", "Semantic wrapper accepts the best route only above this local threshold."),
        ("Semantic top-k k", "classify_semantic_topk(text, k=3)", "Caller can request multiple candidate routes to apply a separate margin decision."),
        ("JARVIS_SEMANTIC_MIN_MARGIN", "Imported/used by the core caller, not defined in nlp/", "Margin guard for near ties; this phase observes top-k support in nlp/ and documents the caller rule without reading core."),
        ("MATCH_THRESHOLD", "70 in nlp/intent_classifier.py", "Minimum fuzzy keyword score for action/target matches."),
        ("MIN_INTENT_SCORE", "1 in nlp/intent_classifier.py", "Lowest accepted weighted keyword score."),
        ("MIN_CONFIDENCE", "0.35 in nlp/intent_classifier.py", "Lowest accepted keyword-classifier confidence."),
        ("SUGGEST_THRESHOLD", "0.45 in nlp/intent_classifier.py", "Minimum suggestion confidence from the looser suggest_intent path."),
        ("Entity types", "app, system_feature, path, number, duration, email, person, date", "Slot taxonomy declared by nlp/entity_types.py."),
    ]


def add_phase6_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "The nlp/ package is Jarvis's text-understanding layer between recognized speech text and the command router. "
        "It normalizes noisy English/Arabic input, gives mixed-language commands a fast shortcut, offers semantic top-k "
        "route candidates, falls back to fuzzy keyword evidence, and enriches parsed commands with slots and entity types."
    )

    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], nlp_inventory_rows())
    add_caption(document, "Table 4.6.1: Inventory of Python files observed directly under nlp/.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_nlp_flow_diagram(), "Figure 4.6.1: NLP flow from normalized text through code-switch, semantic top-k, fuzzy/keyword, and slot extraction stages.")

    document.add_heading("Key Modules In Depth", level=3)
    document.add_paragraph(
        "No standalone nlp/text_normalizer.py is present. The observed normalization responsibility is split: "
        "codeswitching.py handles Arabic letter folding, tashkeel stripping, Arabic-Indic numerals, script counts, and "
        "mixed Arabic/Latin entity extraction; fuzzy_matcher.py performs lower-level punctuation and whitespace cleanup for keyword matching."
    )
    add_code_excerpt(document, code_excerpt("nlp/codeswitching.py", 207, 230))
    add_caption(document, "Listing 4.6.1: Arabic diacritic, letter, and numeral normalization helpers.")
    add_code_excerpt(document, code_excerpt("nlp/codeswitching.py", 369, 421))
    add_caption(document, "Listing 4.6.2: normalize_codeswitched() return shape and initial script/verb analysis.")

    document.add_paragraph(
        "code_switch_router.py is the fast shortcut for mixed-language command forms. It reuses verb sets from intent confidence, "
        "normalizes system actions, and emits ParsedCommand objects for open, close, delete, rename, move, and system-action paths when a verb/entity pattern is confident."
    )
    add_code_excerpt(document, code_excerpt("nlp/code_switch_router.py", 97, 148))
    add_caption(document, "Listing 4.6.3: try_codeswitch() system action and open/close mapping.")
    add_code_excerpt(document, code_excerpt("nlp/code_switch_router.py", 155, 192))
    add_caption(document, "Listing 4.6.4: try_codeswitch() delete, rename, and move mapping.")

    document.add_paragraph(
        "semantic_router.py loads sentence-transformers lazily and exposes classify_semantic_topk(). That API returns ranked "
        "intent names and scores; the older classify_semantic() wrapper accepts only the top result when it reaches the local 0.75 threshold."
    )
    add_code_excerpt(document, code_excerpt("nlp/semantic_router.py", 537, 557))
    add_caption(document, "Listing 4.6.5: Semantic router availability and local confidence threshold.")
    add_code_excerpt(document, code_excerpt("nlp/semantic_router.py", 608, 655))
    add_caption(document, "Listing 4.6.6: Semantic top-k scoring and back-compatible single-result wrapper.")

    document.add_paragraph(
        "fuzzy_matcher.py and keyword_engine.py provide the non-embedding fallback surface. keyword_engine.py stores bilingual "
        "intent keywords, while fuzzy_matcher.py scores exact and approximate containment. intent_classifier.py combines action and target matches, rejects weak evidence, and demotes near ties to unknown."
    )
    add_code_excerpt(document, code_excerpt("nlp/fuzzy_matcher.py", 30, 45))
    add_caption(document, "Listing 4.6.7: Fuzzy matcher normalization and score entry point.")
    add_code_excerpt(document, code_excerpt("nlp/fuzzy_matcher.py", 78, 92))
    add_caption(document, "Listing 4.6.8: Fuzzy keyword containment and sorted match collection.")
    add_code_excerpt(document, code_excerpt("nlp/intent_classifier.py", 89, 119))
    add_caption(document, "Listing 4.6.9: Keyword classifier acceptance and near-tie rejection.")

    document.add_paragraph(
        "nlu.py is the slot layer. It preserves existing parser args, then enriches missing app, duration, timer label, "
        "file, location, time, and reminder-message slots. entity_types.py names the small shared taxonomy consumed by route definitions."
    )
    add_code_excerpt(document, code_excerpt("nlp/nlu.py", 211, 238))
    add_caption(document, "Listing 4.6.10: NLUResult assembly with required-slot detection.")
    add_code_excerpt(document, code_excerpt("nlp/nlu.py", 247, 305))
    add_caption(document, "Listing 4.6.11: Domain-specific entity enrichment.")
    add_code_excerpt(document, code_excerpt("nlp/entity_types.py", 1, 16))
    add_caption(document, "Listing 4.6.12: EntityType enum and lookup map.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(document, "Semantic margin acceptance:", [
        "topk = classify_semantic_topk(text, k=3)",
        "if topk is empty: continue to fuzzy/keyword path",
        "best = topk[0]",
        "second = topk[1] if present else score 0.0",
        "if best.score >= tau and (best.score - second.score) >= delta:",
        "    accept best.intent",
        "else:",
        "    defer because the semantic evidence is weak or ambiguous",
    ])
    add_pseudocode(document, "Code-switch resolution:", [
        "normalize Arabic letters, remove tashkeel, convert Arabic-Indic digits",
        "tokenize and count Arabic vs Latin script segments",
        "detect an Arabic or English verb in the early tokens",
        "extract Latin app/entity text or Arabic mapped entity",
        "try system-action normalization first",
        "map verb families to ParsedCommand intents/actions",
        "return None if the verb/entity pattern is not confident",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Setting", "Observed value", "NLP role"], nlp_config_rows())
    add_caption(document, "Table 4.6.2: NLP constants and imported configuration surfaces visible from nlp/.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "NLP behavior", "Evidence"], [
        ("Normal", "Mixed English/Arabic input is normalized, code-switch shortcuts can return ParsedCommand, semantic top-k ranks paraphrases, and NLU fills required slots.", "codeswitching.py, code_switch_router.py, semantic_router.py, nlu.py."),
        ("Degraded", "If sentence-transformers is missing, semantic_router.py logs the disabled state and returns no semantic candidates; fuzzy/keyword and code-switch modules remain available.", "semantic_router.py _ensure_loaded(); fuzzy_matcher.py; intent_classifier.py."),
        ("Adversarial/Edge", "Near keyword ties are returned as unknown, semantic callers can reject small margins, and question-shaped phrases such as informational Chrome questions are not claimed as app-open commands by nlp/ alone.", "intent_classifier.py tie guard; semantic_router.py top-k support."),
        ("Code-switch example", "A phrase such as mixed Arabic plus a Latin app name can produce an app entity hint or direct app command when the verb/entity pattern matches.", "codeswitching.py entity extraction; code_switch_router.py try_codeswitch(); nlu.py _cs_entity_hint."),
    ])
    add_caption(document, "Table 4.6.3: NLP behavior under normal, degraded, and adversarial conditions.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "nlp/ advances the bilingual-assistant objective by making noisy and mixed-script speech actionable without forcing every "
        "utterance into the LLM path. Its layered design keeps high-confidence shortcuts cheap, gives semantic routing ranked evidence, "
        "uses keyword/fuzzy logic when embeddings are unavailable, and centralizes slot enrichment before execution decisions."
    )


def fill_phase6(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-nlp", add_phase6_content)


LLM_FILE_RESPONSIBILITIES = {
    "llm/ollama_client.py": "Local Ollama generate client, runtime model settings, streaming, keep-alive, fallback model attempts, think-tag filtering, and cancel checks.",
    "llm/prompt_builder.py": "Prompt package assembly with persona, language pinning, prompt tiers, memory block insertion, KB context, few-shot caps, and num_ctx autosizing.",
    "llm/sentence_buffer.py": "Language-aware streaming chunk accumulator for English and Arabic TTS boundaries.",
    "llm/structured_nlu.py": "Strict JSON command-understanding fallback over Ollama chat, intended to be gated by the caller and verified before execution.",
    "llm/tool_caller.py": "Ollama tool-call schema, optional Claude tool format bridge, and conversion from tool calls to ParsedCommand objects.",
    "llm/prompts/full_prompt.txt": "High-tier prompt template with examples and language/style constraints.",
    "llm/prompts/slim_prompt.txt": "Medium-tier prompt template used by the default prompt tier map.",
    "llm/prompts/micro_prompt.txt": "Minimal/low-tier prompt template with compact examples.",
}


def llm_inventory_rows() -> list[tuple[str, str]]:
    rows = []
    for path in sorted((ROOT / "llm").rglob("*"), key=lambda p: p.as_posix().lower()):
        if not path.is_file() or "__pycache__" in path.parts:
            continue
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, LLM_FILE_RESPONSIBILITIES.get(rel, "LLM artifact present in this folder.")))
    if not (ROOT / "llm" / "claude_client.py").exists():
        rows.append(("llm/claude_client.py", "Not present as source in this workspace; tool_caller.py contains an optional import path but this phase does not document a source file that is absent."))
    return rows


def llm_prompt_template_rows() -> list[tuple[str, str, str]]:
    rows = []
    for path in sorted((ROOT / "llm" / "prompts").glob("*.txt")):
        text = path.read_text(encoding="utf-8")
        first_nonblank = next((line.strip() for line in text.splitlines() if line.strip()), "")
        rows.append((path.relative_to(ROOT).as_posix(), f"{len(text.splitlines())} lines / {len(text)} chars", first_nonblank[:120]))
    return rows


def llm_config_rows() -> list[tuple[str, str]]:
    return [
        ("LLM_MODEL", "Default model name used by ollama_client.py, structured_nlu.py, and tool_caller.py when no runtime model is set."),
        ("LLM_FALLBACK_MODELS", "Additional non-streaming model candidates tried by ollama_client.py."),
        ("LLM_OLLAMA_BASE_URL", "Base URL used to build /api/generate and /api/chat endpoints."),
        ("LLM_OLLAMA_NUM_CTX", "Fallback context ceiling used by runtime getters and prompt_builder.py."),
        ("LLM_TEMPERATURE / LLM_TOP_P / LLM_REPEAT_PENALTY", "Generation options inserted into Ollama request payloads."),
        ("LLM_MAX_RESPONSE_TOKENS", "Mapped to Ollama num_predict."),
        ("LLM_STOP_TOKENS", "Decoded into the Ollama stop-token list."),
        ("LLM_TIMEOUT_SECONDS", "HTTP timeout and streaming hard-timeout input."),
        ("LLM_CTX_AUTOSIZE", "Enables prompt-size-based num_ctx selection in prompt_builder.py."),
        ("LLM_FEWSHOT_MIN / LLM_FEWSHOT_MAX", "Bounds few-shot example counts per prompt tier."),
        ("LLM_LANG_PIN_ENABLED", "Controls language pin insertion and older template hint filtering."),
        ("MEMORY_PROMPT_BLOCK_ENABLED / MEMORY_MAX_CONTEXT_CHARS", "Controls whether prompt_builder.py inserts one bounded memory block."),
        ("KB_TOP_K / KB_MAX_CONTEXT_CHARS", "Bounds retrieved knowledge-base context in prompt_builder.py."),
        ("SENTENCE_BUFFER_*", "Controls English/Arabic soft and hard word limits plus connector holding in sentence_buffer.py."),
        ("CLAUDE_DEFAULT_MODEL", "Imported only by tool_caller.py's optional Claude tool-call path."),
    ]


def add_phase7_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "The llm/ package is Jarvis's answer-generation and LLM-assisted fallback layer. It builds a single prompt package "
        "for an answer turn, streams local Ollama output through a sentence buffer for TTS, and also contains constrained "
        "structured/tool-call paths for ambiguous commands. This section documents only llm/ source and templates; the memory "
        "content inserted into prompts is cross-referenced to Section 4.3."
    )

    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], llm_inventory_rows())
    add_caption(document, "Table 4.7.1: Inventory of llm/ Python modules and prompt templates.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_llm_flow_diagram(), "Figure 4.7.1: LLM flow from prompt package through Ollama streaming and sentence-buffered TTS callbacks.")

    document.add_heading("Key Modules In Depth", level=3)
    document.add_paragraph(
        "ollama_client.py constructs the /api/generate payload with model, prompt, stream mode, keep_alive, num_ctx, "
        "temperature, top_p, repeat_penalty, num_predict, and stop tokens. For qwen3 thinking-mode models, it also disables "
        "thinking where supported and appends /no_think if needed."
    )
    add_code_excerpt(document, code_excerpt("llm/ollama_client.py", 61, 99))
    add_caption(document, "Listing 4.7.1: Ollama request payload options, keep-alive, stop tokens, and think suppression.")
    add_code_excerpt(document, code_excerpt("llm/ollama_client.py", 217, 253))
    add_caption(document, "Listing 4.7.2: Streaming entry point, cancel-event default, SentenceBuffer, and payload creation.")
    add_code_excerpt(document, code_excerpt("llm/ollama_client.py", 276, 336))
    add_caption(document, "Listing 4.7.3: Stream loop with cancel checks, hard timeout, think-tag filtering, and sentence callbacks.")
    add_code_excerpt(document, code_excerpt("llm/ollama_client.py", 360, 384))
    add_caption(document, "Listing 4.7.4: Streaming timeout/connect-error handling and non-streaming fallback attempt.")

    document.add_paragraph(
        "prompt_builder.py owns prompt size and structure. It maps model tiers to prompt templates, estimates token counts, "
        "autosizes num_ctx under the runtime ceiling, injects persona/language rules, inserts exactly one memory prompt block "
        "when enabled, and returns a package containing prompt text, context metadata, and selected num_ctx."
    )
    add_code_excerpt(document, code_excerpt("llm/prompt_builder.py", 20, 27))
    add_caption(document, "Listing 4.7.5: Prompt tier to template mapping.")
    add_code_excerpt(document, code_excerpt("llm/prompt_builder.py", 131, 180))
    add_caption(document, "Listing 4.7.6: Token estimate and num_ctx autosizing.")
    add_code_excerpt(document, code_excerpt("llm/prompt_builder.py", 344, 385))
    add_caption(document, "Listing 4.7.7: Prompt package assembly with memory and knowledge context.")

    document.add_paragraph(
        "sentence_buffer.py turns streamed token fragments into speakable chunks. English and Arabic use different boundary "
        "character sets and soft/hard word limits, and connector holding prevents chunks from ending on conjunction-like words."
    )
    add_code_excerpt(document, code_excerpt("llm/sentence_buffer.py", 28, 56))
    add_caption(document, "Listing 4.7.8: SentenceBuffer constants and runtime thresholds.")
    add_code_excerpt(document, code_excerpt("llm/sentence_buffer.py", 92, 119))
    add_caption(document, "Listing 4.7.9: Boundary flush, hard flush, and connector holding.")

    document.add_paragraph(
        "structured_nlu.py is a strict JSON fallback, not a direct executor. Its module docstring states that the caller gates it "
        "and passes successful output through verification. The file grounds the prompt in known schema names, asks for one JSON object, "
        "coerces confidence into 0.0-1.0, and returns None on network, status, JSON, or missing-intent failures."
    )
    add_code_excerpt(document, code_excerpt("llm/structured_nlu.py", 30, 61))
    add_caption(document, "Listing 4.7.10: Structured NLU prompt grounded in known intents.")
    add_code_excerpt(document, code_excerpt("llm/structured_nlu.py", 90, 148))
    add_caption(document, "Listing 4.7.11: Strict JSON request, failure handling, and coercion boundary.")

    document.add_paragraph(
        "tool_caller.py defines a default tool schema for app, file, timer, search, volume, brightness, lock, sleep, and screenshot commands. "
        "It can call Ollama chat with tools, normalize returned tool calls, optionally convert the schema to Claude's input_schema format, "
        "and map recognized tool calls into ParsedCommand objects."
    )
    add_code_excerpt(document, code_excerpt("llm/tool_caller.py", 12, 68))
    add_caption(document, "Listing 4.7.12: Default tool schema excerpt.")
    add_code_excerpt(document, code_excerpt("llm/tool_caller.py", 180, 233))
    add_caption(document, "Listing 4.7.13: Ollama chat tool-call request and normalization.")
    add_code_excerpt(document, code_excerpt("llm/tool_caller.py", 235, 282))
    add_caption(document, "Listing 4.7.14: Optional Claude tool-call bridge; source file is absent in this workspace.")
    add_code_excerpt(document, code_excerpt("llm/tool_caller.py", 300, 388))
    add_caption(document, "Listing 4.7.15: Tool calls converted to ParsedCommand objects.")

    document.add_heading("Prompt Templates", level=3)
    add_styled_table(document, ["Template", "Size", "First nonblank line"], llm_prompt_template_rows())
    add_caption(document, "Table 4.7.2: Prompt template files under llm/prompts/.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(document, "num_ctx autosize:", [
        "estimate tokens from English and Arabic character counts",
        "ceiling = runtime context ceiling for the selected model tier",
        "if LLM_CTX_AUTOSIZE is disabled: use ceiling",
        "if tokens <= 256: select 512",
        "else if tokens <= 512: select 1024",
        "else if tokens <= 1024: select 2048",
        "else if tokens <= 2048: select 4096",
        "else: select ceiling",
        "return selected value clamped between 512 and ceiling",
    ])
    add_pseudocode(document, "sentence-boundary flush:", [
        "append streamed token to buffer",
        "choose English or Arabic boundary set",
        "if latest boundary exists and word count reaches soft limit:",
        "    flush through boundary unless the candidate ends with a connector",
        "if word count reaches hard limit:",
        "    hold when configured and buffer ends with a connector",
        "    otherwise flush at the hard-word position",
        "on stream end or timeout with partial text: flush remainder",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Setting", "LLM role observed in llm/"], llm_config_rows())
    add_caption(document, "Table 4.7.3: Configuration constants imported by llm/ modules; defaults are centralized outside this phase.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "LLM behavior", "Evidence"], [
        ("Normal", "Prompt builder creates a tiered prompt package with persona/language/memory/KB context; Ollama streams tokens; SentenceBuffer emits speakable chunks to on_sentence.", "prompt_builder.py, ollama_client.py, sentence_buffer.py."),
        ("Degraded", "Streaming returns honest local-model errors on non-200, timeout, empty response, or Ollama connection failure; non-streaming ask_llm fallback is attempted where coded.", "ollama_client.py ask_llm_streaming()."),
        ("Optional fallback", "Tool-caller contains a Claude tool-call path, but llm/claude_client.py is absent as source in this workspace, so this phase documents only the optional import boundary.", "tool_caller.py call_tool_tier_claude(); llm/ file inventory."),
        ("Adversarial/Edge", "Runaway output is bounded by num_predict, stop tokens, request timeout, and hard streaming timeout; qwen3 think tags are filtered; wake-word interrupt can set the cancel event and stop the stream loop.", "ollama_client.py _build_request_payload(); ask_llm_streaming()."),
        ("Structured command ambiguity", "structured_nlu.py returns strict JSON or None and is documented as caller-gated/verifier-bound rather than directly executable.", "structured_nlu.py module docstring and understand_structured()."),
    ])
    add_caption(document, "Table 4.7.4: LLM behavior under normal, degraded, optional, and adversarial conditions.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "llm/ advances the assistant by separating answer generation from deterministic command routing while still providing controlled "
        "fallbacks for ambiguous language. Prompt tiers let smaller models use shorter instructions, memory and KB context are inserted once, "
        "streaming chunks reach TTS quickly, and cancellation/timeout/stop-token controls keep generation bounded."
    )


def fill_phase7(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-llm", add_phase7_content)


OS_CONTROL_FILE_RESPONSIBILITIES = {
    "os_control/__init__.py": "Package marker for OS control helpers.",
    "os_control/action_log.py": "Append/read/verify/reseal facade for action audit logs.",
    "os_control/adapter_result.py": "Shared success/failure/confirmation response shape and router tuple conversion.",
    "os_control/app_ops.py": "Application catalog, app resolution, open/close execution, and confirmed app close payloads.",
    "os_control/app_scanner.py": "Installed application discovery from registry, Start Menu shortcuts, App Paths, Store apps, cache, and watcher hooks.",
    "os_control/batch_ops.py": "Batch command manager with disallowed sensitive intents.",
    "os_control/calendar_ops.py": "Outlook COM calendar draft creation; opens an event window for user confirmation.",
    "os_control/capture_ops.py": "Screenshot and screen recording wrapper over native capture, ffmpeg, and Game Bar paths.",
    "os_control/clipboard_ops.py": "Clipboard read/write/clear with pyperclip availability checks.",
    "os_control/confirmation.py": "Pending action confirmation manager, PIN sentinel flow, expiry, cancellation, and attempt handling.",
    "os_control/email_ops.py": "Email draft creation through Outlook COM or Gmail fallback; does not send mail.",
    "os_control/explorer_ops.py": "Open/reveal/open-file behavior in Windows Explorer with location and fuzzy filename resolution.",
    "os_control/file_ops.py": "Validated file navigation, search, create, move, copy, rename, delete, rollback, and confirmation payloads.",
    "os_control/job_queue.py": "Persistent delayed command queue over the shared persistence database.",
    "os_control/native_ops.py": "Native Windows helpers for volume, brightness, lock/sleep, screenshots, and read-back verification.",
    "os_control/note_ops.py": "Dictated note saving to a configured note directory with automatic names.",
    "os_control/path_resolver.py": "Human folder aliases, drive aliases, path resolution, and path humanization.",
    "os_control/persistence.py": "SQLite-backed action log, confirmations, rollback stack, and job queue persistence.",
    "os_control/policy.py": "Policy engine for command category enable/disable decisions.",
    "os_control/powershell_bridge.py": "Vetted PowerShell template registry and non-arbitrary template runner.",
    "os_control/radio_ops.py": "Wi-Fi, Bluetooth, and airplane mode control through WinRT first and PowerShell fallback.",
    "os_control/reminder_ops.py": "Windows reminder scheduling through Task Scheduler with in-process fallback.",
    "os_control/risk_policy.py": "System, file, and app operation risk-tier mapping plus coverage validation.",
    "os_control/screen_context.py": "Foreground/visible-window and optional screenshot-vision screen description.",
    "os_control/search_index.py": "SQLite file search index worker and query service.",
    "os_control/second_factor.py": "Spoken PIN/passphrase normalization, hashing, verification, and lockout tracking.",
    "os_control/settings_ops.py": "Windows Settings page resolver and ms-settings launcher.",
    "os_control/sysinfo_ops.py": "Battery and system status helpers.",
    "os_control/system_ops.py": "System command registry, alias normalization, confirmation request, native/template dispatch, and result shaping.",
    "os_control/temporal_parser.py": "Natural datetime and recurrence parser for timer/reminder-style commands.",
    "os_control/timer_ops.py": "Countdown timers, alarms, persistence, re-arming, and notification playback.",
    "os_control/windows_toggles.py": "Night Light, Do Not Disturb, Energy Saver, and Live Captions toggles with verification/fallback behavior.",
}


def os_control_inventory_rows() -> list[tuple[str, str]]:
    rows = []
    for path in sorted((ROOT / "os_control").glob("*.py"), key=lambda p: p.name.lower()):
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, OS_CONTROL_FILE_RESPONSIBILITIES.get(rel, "OS control Python module present in this folder.")))
    return rows


def os_control_config_rows() -> list[tuple[str, str]]:
    return [
        ("ALLOW_DESTRUCTIVE_SYSTEM_COMMANDS", "Blocks destructive system commands in system_ops.py unless enabled."),
        ("SECOND_FACTOR_REQUIRED_FOR_DESTRUCTIVE / SECOND_FACTOR_PIN", "Controls whether destructive confirmations require spoken PIN verification."),
        ("CONFIRMATION_TIMEOUT_SECONDS / SENSITIVE_PIN_PENDING_TIMEOUT_SECONDS", "Bounds pending confirmation/PIN action lifetime."),
        ("SECOND_FACTOR_MAX_ATTEMPTS_PER_TOKEN / SECOND_FACTOR_LOCKOUT_SECONDS", "Controls second-factor failed-attempt lockout behavior."),
        ("CONFIRMATION_MAX_ATTEMPTS_PER_TOKEN / CONFIRMATION_LOCKOUT_SECONDS", "Controls confirmation attempt rate limiting."),
        ("CONTROLS_VERIFY_STATE", "Controls read-back verification in native_ops.py and radio_ops.py."),
        ("CONTROLS_ADMIN_HINT", "Allows system_ops.py to add an administrator hint when network radio actions fail without elevation."),
        ("FEATURE_FLAGS", "Gates native volume and media direct-dispatch paths in system_ops.py and auto app discovery in app_ops.py."),
        ("VOLUME_BACKEND", "Allows explicit app-volume mode; system volume path otherwise excludes waveOut readback."),
        ("FILE_DEFAULT_SEARCH_ROOTS / DEFAULT_SEARCH_PATH / DEFAULT_WORKING_DIRECTORY", "Drive file search and current-directory defaults in file_ops.py."),
        ("FILE_HUMANIZE_PATHS / FILE_SPEAK_PATHS / MAX_FILE_RESULTS", "Control file response wording and search result bounds."),
        ("ALLOW_PERMANENT_DELETE / ROLLBACK_DIR_NAME", "Control permanent delete policy and soft-delete rollback storage."),
        ("APP_* resolution constants", "Tune app catalog TTL, refresh-on-miss, and scoring bonuses in app_ops.py."),
        ("TOGGLE_* methods / LIVE_CAPTION_HOTKEY", "Select registry/hotkey/settings fallback behavior in windows_toggles.py."),
        ("AIRPLANE_RESTORE_RADIOS", "Controls whether airplane-off restores the pre-airplane radio snapshot."),
        ("SCREENSHOT_DIR / SCREENRECORD_*", "Configure capture_ops.py screenshot and recording destinations/backends."),
        ("NOTE_DIR / NOTE_BASENAME", "Configure note_ops.py save location and default naming."),
        ("TIMER_*", "Configure timer persistence, Clock app launch, and TTS announcement behavior."),
        ("ACTION_LOG_FILE", "Path used by action_log.py's text append facade."),
        ("JOB_MAX_RETRIES_DEFAULT", "Default retry count used by job_queue.py."),
    ]


def add_phase8_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "os_control/ is the Windows side-effect boundary behind the routing and safety layers. It contains the adapters that "
        "open apps, manipulate files, change system state, manage timers/reminders, capture the screen, and record audit evidence. "
        "The folder's recurring contract is: decide policy/risk, ask confirmation when required, execute through a constrained adapter, "
        "verify state where the adapter supports readback, then return a structured success/failure/confirmation result."
    )

    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], os_control_inventory_rows())
    add_caption(document, "Table 4.8.1: Inventory of every Python file directly under os_control/.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_os_control_flow_diagram(), "Figure 4.8.1: Verified OS side-effect flow from routed action through policy, confirmation, adapter execution, verification, and audit logging.")

    document.add_heading("Key Modules In Depth", level=3)
    document.add_paragraph(
        "system_ops.py is the system-command hub. SYSTEM_COMMANDS declares command metadata such as template name, description, "
        "destructive status, and confirmation requirement. request_system_command_result() applies the policy gate, normalizes args, "
        "routes non-confirmed commands directly, and creates a pending confirmation for confirmed commands."
    )
    add_code_excerpt(document, code_excerpt("os_control/system_ops.py", 47, 80))
    add_caption(document, "Listing 4.8.1: SYSTEM_COMMANDS registry excerpt.")
    add_code_excerpt(document, code_excerpt("os_control/system_ops.py", 1594, 1654))
    add_caption(document, "Listing 4.8.2: System command policy, risk, confirmation, and second-factor request path.")
    add_code_excerpt(document, code_excerpt("os_control/system_ops.py", 1972, 2028))
    add_caption(document, "Listing 4.8.3: System command execution with destructive block and native-first dispatch.")

    document.add_paragraph(
        "native_ops.py is the low-level control adapter for volume, brightness, lock/sleep, and capture. The volume setter is explicit "
        "about the system-volume contract: pycaw and PowerShell media-key paths are eligible, while waveOut is excluded from system-volume "
        "verification because it changes this process's volume rather than the master endpoint."
    )
    add_code_excerpt(document, code_excerpt("os_control/native_ops.py", 552, 611))
    add_caption(document, "Listing 4.8.4: System volume set with read-back verification and waveOut exclusion.")
    add_code_excerpt(document, code_excerpt("os_control/native_ops.py", 699, 739))
    add_caption(document, "Listing 4.8.5: Brightness getter/setter chain with read-back check.")

    document.add_paragraph(
        "radio_ops.py uses WinRT first for Wi-Fi and Bluetooth, then PowerShell fallback. Both paths report success only after the requested "
        "state is read back when verification is enabled. Airplane mode snapshots existing radio state and treats partial restore/change as failure."
    )
    add_code_excerpt(document, code_excerpt("os_control/radio_ops.py", 132, 160))
    add_caption(document, "Listing 4.8.6: WinRT radio set with polling verification.")
    add_code_excerpt(document, code_excerpt("os_control/radio_ops.py", 312, 405))
    add_caption(document, "Listing 4.8.7: WinRT-to-PowerShell radio fallback and airplane-mode all-or-nothing result.")

    document.add_paragraph(
        "windows_toggles.py covers Windows 11 toggles that do not share one public API. Night Light and Do Not Disturb use registry write/read-back "
        "paths with Settings URI fallback; Energy Saver uses registry or powercfg paths; Live Captions uses hotkey/window detection and settings fallback."
    )
    add_code_excerpt(document, code_excerpt("os_control/windows_toggles.py", 126, 172))
    add_caption(document, "Listing 4.8.8: Night Light registry write, shell notification, verification, and Settings fallback.")
    add_code_excerpt(document, code_excerpt("os_control/windows_toggles.py", 180, 224))
    add_caption(document, "Listing 4.8.9: Do Not Disturb registry write/read-back and fallback.")

    document.add_paragraph(
        "file_ops.py, path_resolver.py, and explorer_ops.py are the human-file layer. Path resolver maps spoken folder and drive aliases to concrete "
        "paths and short speakable names. Explorer operations can open or reveal paths. File write operations validate policy and paths, resolve bare "
        "names from configured roots, return ambiguity instead of guessing, and create confirmation payloads for move, rename, delete, and copy."
    )
    add_code_excerpt(document, code_excerpt("os_control/path_resolver.py", 204, 286))
    add_caption(document, "Listing 4.8.10: Spoken location resolution and humanized path output.")
    add_code_excerpt(document, code_excerpt("os_control/explorer_ops.py", 86, 155))
    add_caption(document, "Listing 4.8.11: Explorer fuzzy filename matching with extension-optional support and scan cap.")
    add_code_excerpt(document, code_excerpt("os_control/file_ops.py", 946, 1116))
    add_caption(document, "Listing 4.8.12: File move/rename/delete/copy request paths and confirmation payloads.")

    document.add_paragraph(
        "app_scanner.py and app_ops.py keep application control separate from file navigation. app_ops.py blocks generic folder words from resolving "
        "as apps, scores app candidates by aliases, running state, availability, and usage, refreshes a stale catalog on miss, and requires confirmation for app-close operations."
    )
    add_code_excerpt(document, code_excerpt("os_control/app_ops.py", 554, 681))
    add_caption(document, "Listing 4.8.13: App candidate scoring, folder-word blocklist, ambiguity margin, and refresh-on-miss.")
    add_code_excerpt(document, code_excerpt("os_control/app_ops.py", 947, 1006))
    add_caption(document, "Listing 4.8.14: App close confirmation payload.")

    document.add_paragraph(
        "confirmation.py and second_factor.py implement the spoken PIN path. The current flow stores a single pending PIN action behind a fixed sentinel, "
        "normalizes spoken digits, checks lockout state, compares hashes with hmac.compare_digest(), discards expired or locked actions, and logs accepted or rejected attempts."
    )
    add_code_excerpt(document, code_excerpt("os_control/confirmation.py", 42, 154))
    add_caption(document, "Listing 4.8.15: Pending PIN action creation and verify_pin_and_execute().")
    add_code_excerpt(document, code_excerpt("os_control/second_factor.py", 45, 160))
    add_caption(document, "Listing 4.8.16: Spoken PIN normalization, attempt limits, and second-factor verification.")

    document.add_paragraph(
        "risk_policy.py centralizes risk labels for system, file, and app actions. adapter_result.py standardizes success, failure, and confirmation responses. "
        "persistence.py stores a hash-chained action log, pending confirmations, rollback actions, and background jobs in SQLite."
    )
    add_code_excerpt(document, code_excerpt("os_control/risk_policy.py", 1, 67))
    add_caption(document, "Listing 4.8.17: Risk-tier mappings and coverage validation.")
    add_code_excerpt(document, code_excerpt("os_control/adapter_result.py", 1, 92))
    add_caption(document, "Listing 4.8.18: Adapter result and router tuple response shape.")
    add_code_excerpt(document, code_excerpt("os_control/persistence.py", 37, 123))
    add_caption(document, "Listing 4.8.19: SQLite tables for action logs, confirmations, rollback, and jobs.")

    document.add_paragraph(
        "The remaining adapters broaden the OS surface without changing the safety pattern: timer_ops.py persists/re-arms timers, reminder_ops.py uses Task Scheduler "
        "with in-process fallback, capture_ops.py writes screenshots/recordings under configured dirs, clipboard_ops.py fails honestly when pyperclip is unavailable, "
        "email_ops.py drafts but does not send mail, note_ops.py writes .txt notes, screen_context.py describes visible windows or uses screenshot vision, and search_index.py keeps a file index."
    )
    add_styled_table(document, ["Adapter group", "Modules", "Observed behavior"], [
        ("System state", "system_ops.py, native_ops.py, radio_ops.py, windows_toggles.py, powershell_bridge.py", "Constrained command registry, native-first execution, vetted PowerShell templates, WinRT/registry/hotkey paths, and read-back checks where implemented."),
        ("Files and Explorer", "file_ops.py, path_resolver.py, explorer_ops.py, search_index.py", "Human folder aliases, extension-optional search, ambiguity handling, confirmed writes, rollback storage, and indexed search."),
        ("Applications", "app_scanner.py, app_ops.py", "Static plus scanned catalog, usage/running/availability scoring, folder-word blocklist, and confirmed app close."),
        ("Safety and audit", "confirmation.py, second_factor.py, policy.py, risk_policy.py, persistence.py, action_log.py, adapter_result.py", "Policy blocks, risk tiers, PIN/attempt lockout, structured adapter results, hash-chained action log, confirmations, rollback, and jobs."),
        ("Personal productivity", "timer_ops.py, reminder_ops.py, clipboard_ops.py, email_ops.py, note_ops.py, calendar_ops.py", "Timers/reminders, clipboard operations, mail/calendar drafts, and note saving with explicit failure paths."),
        ("Screen and status", "capture_ops.py, screen_context.py, settings_ops.py, sysinfo_ops.py", "Screenshot/recording, visible-window description, Settings URI opening, and battery/system status."),
    ])
    add_caption(document, "Table 4.8.2: os_control/ adapter families.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(document, "Verify-after-execute contract:", [
        "normalize action and arguments",
        "check policy and destructive/confirmation gates",
        "execute through the safest available adapter",
        "if CONTROLS_VERIFY_STATE is enabled and readback exists:",
        "    poll or read OS state after execution",
        "    if actual state is unavailable or mismatched: report failure or try next backend",
        "log success, blocked, pending, or failed action",
        "return adapter_result success/failure/confirmation payload",
    ])
    add_pseudocode(document, "PIN confirmation flow:", [
        "create pending action with PIN_REQUIRED sentinel and expiry",
        "when user speaks PIN: normalize digit words and Arabic-Indic digits",
        "if no pending action or expired: reject and clear state",
        "if confirmation attempts are locked: discard pending action",
        "compare hashed spoken secret to configured PIN/passphrase hash",
        "on success: discard pending action, clear attempts, return stored payload",
        "on repeated failure: lock out and discard pending action",
    ])
    add_pseudocode(document, "Risk-tier decision:", [
        "system action: explicit override wins",
        "else destructive system command is high risk",
        "else confirmation-required system command is medium risk",
        "file delete/permanent delete is high; move/rename are medium",
        "app close is medium",
        "unmapped operation defaults to low, but validate_risk_policy_coverage can flag required mappings",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Setting", "Role visible from os_control/"], os_control_config_rows())
    add_caption(document, "Table 4.8.3: Configuration symbols imported or referenced by os_control/ modules; defaults are outside this phase boundary.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "os_control behavior", "Evidence"], [
        ("Normal", "A routed action passes policy/risk checks, executes through the matching adapter, verifies state where supported, logs the result, and returns a structured adapter payload.", "system_ops.py, native_ops.py, radio_ops.py, adapter_result.py, action_log.py."),
        ("Degraded", "If WinRT radio control is unavailable or fails, radio_ops.py falls back to PowerShell; if registry toggle does not verify, windows_toggles.py opens Settings URI and returns False; if pyperclip or Outlook COM are unavailable, clipboard/email paths return explicit fallback/error messages.", "radio_ops.py, windows_toggles.py, clipboard_ops.py, email_ops.py."),
        ("Elevation/Admin", "Network radio failures can receive an admin hint from system_ops.py when configured; file and destructive system operations are blocked by policy/config rather than forced.", "system_ops.py, file_ops.py."),
        ("Adversarial/Edge", "Wrong PIN attempts are rate-limited and can discard the pending action; permanent delete is blocked unless enabled; ambiguous file/app resolution returns ambiguity instead of choosing silently; read-back mismatch reports failure.", "second_factor.py, confirmation.py, file_ops.py, app_ops.py, native_ops.py."),
        ("Auditability", "Action logs are stored with previous hash plus payload digest, confirmations and rollback actions share the persistence database, and adapter_result.py preserves metadata for the router.", "persistence.py, action_log.py, adapter_result.py."),
    ])
    add_caption(document, "Table 4.8.4: os_control behavior under normal, degraded, elevation, adversarial, and audit situations.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "os_control/ advances verified control and safety objectives by keeping real Windows side effects behind explicit adapters, policy/risk gates, "
        "confirmation and PIN checks, read-back verification, and audit persistence. The folder also makes failure states useful: a control that cannot "
        "verify, lacks permissions, hits ambiguity, or fails a PIN does not claim success."
    )


def fill_phase8(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-oscontrol", add_phase8_content)


TOOLS_FILE_RESPONSIBILITIES = {
    "tools/__init__.py": "Tools package marker.",
    "tools/weather.py": "Open-Meteo current-weather client with WMO weather-code text mapping and empty-string failure behavior.",
    "tools/web_search.py": "DuckDuckGo text search wrapper with timeout, blocklist filtering, trusted-domain boost, recency boost, and formatted snippets.",
    "tools/live_data.py": "Weather/search intent detector, parallel live-data aggregator, and voice-ready labeled block formatter.",
    "tools/calculator.py": "Safe quick calculator for math-looking expressions, including Arabic-Indic digit normalization.",
    "tools/evaluate_nlu.py": "Decision-only routing evaluation harness for intent accuracy, slots, latency, and safety counters.",
}

UTILS_FILE_RESPONSIBILITIES = {
    "utils/__init__.py": "Utilities package marker.",
    "utils/language_detector.py": "Counts Arabic and English characters and classifies text as ar, en, mixed, or unknown.",
}


def simple_inventory_rows(folder: str, mapping: dict[str, str]) -> list[tuple[str, str]]:
    rows = []
    for path in sorted((ROOT / folder).glob("*.py"), key=lambda p: p.name.lower()):
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, mapping.get(rel, f"{folder}/ Python module present in this workspace.")))
    return rows


def tools_config_rows() -> list[tuple[str, str]]:
    return [
        ("WEATHER_DEFAULT_CITY / WEATHER_DEFAULT_LATITUDE / WEATHER_DEFAULT_LONGITUDE", "weather.py and live_data.py use these when no explicit weather city/location is available."),
        ("WEB_SEARCH_ENABLED", "live_data.py only force-enables generic search when this flag is true."),
        ("WEB_SEARCH_MAX_RESULTS", "Bounds search result count passed from live_data.py to web_search.py."),
        ("VOICE_NORMALIZER_MAX_SEARCH_RESULTS", "Bounds search rows retained when live_data.py formats voice-ready search blocks."),
        ("WEB_SEARCH_TRUSTED_DOMAINS / WEB_SEARCH_BLOCKED_DOMAINS", "web_search.py boosts trusted hosts and removes blocked hosts."),
        ("WEB_SEARCH_TRUSTED_DOMAIN_BOOST / WEB_SEARCH_RECENCY_BOOST", "web_search.py scoring weights for reranking search results."),
        ("_TIMEOUT_SECONDS", "weather.py local Open-Meteo timeout: 5.0 seconds."),
        ("_SEARCH_TIMEOUT_SECONDS", "web_search.py local DuckDuckGo timeout: 6.0 seconds."),
        ("_LIVE_DATA_TIMEOUT", "live_data.py outer future timeout: 7.0 seconds."),
    ]


def add_phase9_content(document: Document) -> None:
    document.add_heading("4.9 tools/", level=2)
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "tools/ contains small side-effect-light helper tools used to enrich answers or shortcut deterministic calculations. "
        "The runtime live-data path detects weather or web-search needs, fetches only the relevant tool output, formats nonempty "
        "results into labeled voice-ready prompt blocks, and returns an empty string when no live data is needed or a fetch fails."
    )

    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], simple_inventory_rows("tools", TOOLS_FILE_RESPONSIBILITIES))
    add_caption(document, "Table 4.9.1: Inventory of Python files directly under tools/.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_tools_flow_diagram(), "Figure 4.9.1: Live data path from query detection through weather/search tools and voice-ready result blocks.")

    document.add_heading("Key Modules In Depth", level=3)
    document.add_paragraph(
        "weather.py calls the Open-Meteo forecast endpoint for current temperature, humidity, weather code, and wind speed. "
        "It maps WMO weather codes to short English condition text and returns an empty string on non-200 status or exceptions."
    )
    add_code_excerpt(document, code_excerpt("tools/weather.py", 12, 89))
    add_caption(document, "Listing 4.9.1: Open-Meteo current weather request and empty-string failure behavior.")

    document.add_paragraph(
        "web_search.py wraps DuckDuckGo text search. It runs the provider call in a worker with a hard timeout, drops blocked domains, "
        "scores remaining results with base rank, trusted-domain boost, and recency boost, then formats title/body/domain/date snippets."
    )
    add_code_excerpt(document, code_excerpt("tools/web_search.py", 149, 217))
    add_caption(document, "Listing 4.9.2: Web result scoring, filtering, ranking, and formatting.")
    add_code_excerpt(document, code_excerpt("tools/web_search.py", 220, 255))
    add_caption(document, "Listing 4.9.3: DuckDuckGo search timeout and graceful failure boundary.")

    document.add_paragraph(
        "live_data.py is the aggregator. It detects weather intent with English and Arabic terms, extracts a city when possible, detects web-search intent "
        "while excluding file-search phrases, optionally honors a force_search flag, and runs weather/search futures in parallel."
    )
    add_code_excerpt(document, code_excerpt("tools/live_data.py", 39, 128))
    add_caption(document, "Listing 4.9.4: Weather and web-search intent detection.")
    add_code_excerpt(document, code_excerpt("tools/live_data.py", 157, 242))
    add_caption(document, "Listing 4.9.5: Voice-ready result block formatting and parallel fetch orchestration.")

    document.add_paragraph(
        "calculator.py provides a fast deterministic math path. It only evaluates text that looks mathematical, normalizes Arabic-Indic digits and phrases, "
        "strips unsafe characters, rejects dangerous names, evaluates with only the math module and empty builtins, and returns None when it should fall through."
    )
    add_code_excerpt(document, code_excerpt("tools/calculator.py", 114, 164))
    add_caption(document, "Listing 4.9.6: Safe quick calculation gate and minimal eval namespace.")

    document.add_paragraph(
        "evaluate_nlu.py is a developer harness rather than a live-data tool. It loads routing cases, calls a decision-only path, and reports intent accuracy, "
        "slot accuracy, latency percentiles, unsafe executions, question false-fires, and margin violations without executing OS actions."
    )
    add_code_excerpt(document, code_excerpt("tools/evaluate_nlu.py", 49, 174))
    add_caption(document, "Listing 4.9.7: Decision-only NLU evaluation counters.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(document, "Live data selection:", [
        "query = user_query.strip()",
        "language = ar if query contains Arabic script else en",
        "weather_intent = detect weather keywords and optional city",
        "search_intent = None when weather already answers the query",
        "else search_intent = detect search/news/current keywords, excluding file-search phrases",
        "if no intent and force_search and WEB_SEARCH_ENABLED: search full query",
        "fetch weather/search in parallel when requested",
        "for each nonempty result: normalize for voice and wrap in [WEATHER] or [WEB_SEARCH]",
        "return joined blocks or empty string",
    ])
    add_pseudocode(document, "Search ranking:", [
        "collect more DuckDuckGo rows than requested",
        "drop blocklisted domains",
        "base score preserves provider order",
        "add trusted-domain boost when domain matches configured trusted list",
        "parse available publication date fields",
        "add recency boost with time decay",
        "sort by score descending and format top N rows",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Setting", "Role visible from tools/"], tools_config_rows())
    add_caption(document, "Table 4.9.2: Configuration symbols and local timeouts visible from tools/.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "tools/ behavior", "Evidence"], [
        ("Normal", "Weather and search requests return formatted blocks; calculator returns a formatted number for math-looking input; NLU eval reports decision-only statistics.", "weather.py, web_search.py, live_data.py, calculator.py, evaluate_nlu.py."),
        ("Degraded", "Open-Meteo, DuckDuckGo, missing provider, timeout, or exception paths return empty strings so callers can proceed without live context.", "weather.py get_weather(); web_search.py search_web(); live_data.py gather_live_data()."),
        ("Adversarial/Edge", "Broad Arabic question words are intentionally not generic search triggers; file-search phrases are excluded; calculator rejects dangerous names and non-finite/non-numeric results.", "live_data.py _detect_web_search_intent(); calculator.py quick_calc()."),
        ("Ambiguous Location", "Weather city extraction falls back to WEATHER_DEFAULT_CITY when no city is found or when the captured Arabic token is a time word such as today/now.", "live_data.py _detect_weather_intent()."),
    ])
    add_caption(document, "Table 4.9.3: tools/ behavior under normal, degraded, and edge conditions.")

    document.add_heading("4.10 utils/", level=2)
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "utils/ is currently a small shared-helper package. Its implemented helper is a language detector that classifies text by Arabic and English character counts. "
        "This is deliberately simpler than the heavier NLP language gates documented elsewhere: it exposes a reusable ar/en/mixed/unknown utility."
    )

    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], simple_inventory_rows("utils", UTILS_FILE_RESPONSIBILITIES))
    add_caption(document, "Table 4.10.1: Inventory of Python files directly under utils/.")

    document.add_heading("Key Helpers", level=3)
    add_code_excerpt(document, code_excerpt("utils/language_detector.py", 1, 31))
    add_caption(document, "Listing 4.10.1: Arabic/English character counting and language classification.")
    add_styled_table(document, ["Return value", "Condition"], [
        ("unknown", "No Arabic or English letters found."),
        ("ar", "Arabic letters present and English letters absent."),
        ("en", "English letters present and Arabic letters absent."),
        ("mixed", "Both Arabic and English letters present."),
    ])
    add_caption(document, "Table 4.10.2: utils.language_detector.detect_language() return contract.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "Together, tools/ and utils/ keep live facts and lightweight deterministic helpers out of the LLM core. Weather and search can augment prompts when live context is needed, "
        "calculator avoids unnecessary model calls for arithmetic, evaluate_nlu.py provides regression evidence, and utils/language_detector.py offers a reusable script-count primitive."
    )


def fill_phase9(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-tools", add_phase9_content)


UI_FILE_RESPONSIBILITIES = {
    "ui/bridge.py": "Optional FastAPI/uvicorn WebSocket bridge from engine events to desktop clients and UI commands back to route_command.",
    "ui/events.py": "Shared engine-to-UI event names, UI-to-engine command names, and JSON serialization helper.",
    "ui/tray.py": "Optional pystray system tray with state-colored icon, settings/log shortcuts, restart, demo toggle, and quit cleanup.",
}

DESKTOP_KEY_FILE_RESPONSIBILITIES = {
    "desktop/package.json": "Tauri/Vite/React package manifest, scripts, runtime dependencies, and test tooling.",
    "desktop/vite.config.ts": "Vite React/Tailwind config plus opt-in mock bridge server for dev:mock.",
    "desktop/src-tauri/tauri.conf.json": "Tauri app/window configuration: overlay and dashboard windows, transparency, tray-visible hidden startup.",
    "desktop/src-tauri/src/lib.rs": "Tauri commands, tray menu, overlay positioning, and hide-on-close window behavior.",
    "desktop/src/App.tsx": "Chooses overlay or dashboard view based on Tauri window label or browser fallback store state.",
    "desktop/src/protocol.ts": "Typed bridge protocol: engine events, UI commands, feature flags, config values, state colors, and WebSocket URL.",
    "desktop/src/hooks/useJarvisSocket.ts": "WebSocket connection hook with reconnect/backoff and JSON event dispatch into the store.",
    "desktop/src/stores/jarvisStore.ts": "Zustand store for connection state, dialogue state, config, transcripts, response, avatar settings, and persisted UI preferences.",
    "desktop/src/components/overlay/Overlay.tsx": "Frameless overlay view with wake-driven show/hide, avatar, transcript, prompt input, mute, and dashboard button.",
    "desktop/src/components/dashboard/Dashboard.tsx": "Control-center view for prompt, avatar, persona/model/language selections, feature flags, audio mute, and status.",
    "desktop/src/components/avatar/Avatar.tsx": "Avatar selector over aurora, glyph, glassai, and companion variants.",
    "desktop/src/components/overlay/PromptInput.tsx": "Text prompt form that sends text_command with simple Arabic/English detection.",
}


def ui_inventory_rows() -> list[tuple[str, str]]:
    rows = []
    for path in sorted((ROOT / "ui").glob("*.py"), key=lambda p: p.name.lower()):
        rel = path.relative_to(ROOT).as_posix()
        rows.append((rel, UI_FILE_RESPONSIBILITIES.get(rel, "UI Python module present in this workspace.")))
    return rows


def desktop_key_rows() -> list[tuple[str, str]]:
    rows = []
    for rel, desc in DESKTOP_KEY_FILE_RESPONSIBILITIES.items():
        path = ROOT / rel
        status = desc if path.exists() else "Expected key file was not found in this workspace."
        rows.append((rel, status))
    return rows


def desktop_toolchain_rows() -> list[tuple[str, str]]:
    return [
        ("Runtime shell", "Tauri 2 application with Rust commands and tray menu in desktop/src-tauri/src/lib.rs."),
        ("Frontend", "React 18, TypeScript, Vite 6, Tailwind/Vite plugin, Zustand store, and motion/Three/OGL-related visual dependencies from package.json."),
        ("Development", "npm run dev for Vite, npm run dev:mock for the opt-in mock bridge, npm run build for tsc plus Vite build, npm run test for Vitest."),
        ("Window model", "tauri.conf.json declares an overlay window and a dashboard window, both initially hidden; overlay is transparent, frameless, always-on-top, and skipped from the taskbar."),
        ("Protocol", "desktop/src/protocol.ts mirrors ui/events.py names and defines typed events, commands, config values, feature flags, state colors, and WebSocket URL."),
    ]


def ui_config_rows() -> list[tuple[str, str]]:
    return [
        ("UI_BRIDGE_ENABLED", "ui/bridge.py start() exits early when disabled."),
        ("UI_BRIDGE_HOST / UI_BRIDGE_PORT", "ui/bridge.py binds the FastAPI WebSocket endpoint at ws://host:port/ws."),
        ("LLM_MODEL / LLM_AUTO_SELECT_MODEL", "Included in bridge config events so the dashboard can display model state."),
        ("WAKE_WORD_MODE", "Included in bridge config events."),
        ("FEATURE_FLAGS", "Included in bridge config events and edited locally by dashboard toggles before sending feature_flag commands."),
        ("STT_BACKEND / TTS_DEFAULT_BACKEND / PERSONA_DEFAULT", "Included in bridge config events for dashboard display/control context."),
        ("PROJECT_ROOT / LOG_FILE", "ui/tray.py uses these for .env, project folder, and log shortcuts."),
        ("VITE_JARVIS_WS_URL", "desktop/src/protocol.ts uses this environment variable to target the real Python bridge; otherwise it defaults to mock ws://localhost:8765."),
    ]


def add_phase10_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "ui/ and desktop/ provide an optional visual layer over the Jarvis engine. The Python side can expose a WebSocket bridge "
        "and a small system tray, while the desktop side is a Tauri and React application with overlay and dashboard windows. "
        "The engine remains usable without either UI surface: bridge startup exits cleanly when disabled or when optional dependencies are missing, "
        "and the desktop application connects to the bridge rather than hosting the engine."
    )

    document.add_heading("4.11.1 ui/", level=2)
    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], ui_inventory_rows())
    add_caption(document, "Table 4.11.1: Inventory of Python files directly under ui/.")

    document.add_heading("Bridge Protocol", level=3)
    add_diagram_image(document, add_ui_bridge_diagram(), "Figure 4.11.1: Optional WebSocket protocol between the Jarvis engine and Tauri/React desktop clients.")
    add_code_excerpt(document, code_excerpt("ui/events.py", 1, 35))
    add_caption(document, "Listing 4.11.1: Event and command names shared by the Python bridge.")

    document.add_paragraph(
        "bridge.py creates an optional FastAPI WebSocket server at /ws. It records whether the bridge is enabled, host/port, connected clients, and muted state. "
        "If FastAPI or uvicorn are unavailable, start() logs that optional dependencies are unavailable and leaves the engine running without the bridge."
    )
    add_code_excerpt(document, code_excerpt("ui/bridge.py", 46, 96))
    add_caption(document, "Listing 4.11.2: Optional bridge startup, dependency guard, host/port, and state listener registration.")
    add_code_excerpt(document, code_excerpt("ui/bridge.py", 146, 207))
    add_caption(document, "Listing 4.11.3: WebSocket client loop and handled UI commands.")
    add_code_excerpt(document, code_excerpt("ui/bridge.py", 226, 306))
    add_caption(document, "Listing 4.11.4: route_command handoff, config/health events, and broadcast fan-out.")

    document.add_paragraph(
        "tray.py is another optional UI surface. If pystray is missing, start() logs that the tray icon is disabled. When available, it creates a state-colored icon, "
        "registers a dialogue-state listener, exposes menu shortcuts for .env/logs/project folder, toggles demo mode, restarts Jarvis, and performs shutdown cleanup on quit."
    )
    add_code_excerpt(document, code_excerpt("ui/tray.py", 43, 145))
    add_caption(document, "Listing 4.11.5: Tray menu actions, restart, .env demo-mode update, and shortcuts.")
    add_code_excerpt(document, code_excerpt("ui/tray.py", 176, 230))
    add_caption(document, "Listing 4.11.6: Optional tray startup, state listener, icon thread, and shutdown.")

    document.add_heading("4.11.2 desktop/", level=2)
    document.add_heading("Key File Inventory", level=3)
    add_styled_table(document, ["Key file", "Responsibility"], desktop_key_rows())
    add_caption(document, "Table 4.11.2: Key desktop/ files read for the Tauri+React overview.")

    document.add_heading("Toolchain And Window Model", level=3)
    add_styled_table(document, ["Area", "Observed implementation"], desktop_toolchain_rows())
    add_caption(document, "Table 4.11.3: Tauri+React toolchain and window model from desktop key files.")
    add_code_excerpt(document, code_excerpt("desktop/src-tauri/tauri.conf.json", 1, 53))
    add_caption(document, "Listing 4.11.7: Tauri overlay and dashboard window definitions.")
    add_code_excerpt(document, code_excerpt("desktop/src-tauri/src/lib.rs", 1, 59))
    add_caption(document, "Listing 4.11.8: Tauri commands for overlay/dashboard visibility and quit.")
    add_code_excerpt(document, code_excerpt("desktop/src-tauri/src/lib.rs", 60, 136))
    add_caption(document, "Listing 4.11.9: Tauri tray menu and hide-on-close behavior.")
    add_code_excerpt(document, code_excerpt("desktop/vite.config.ts", 1, 19))
    add_caption(document, "Listing 4.11.10: Vite config and opt-in mock bridge.")

    document.add_paragraph(
        "The desktop protocol mirrors the Python event names with TypeScript discriminated unions. The WebSocket URL defaults to a mock development server and can be pointed at the real Python bridge with VITE_JARVIS_WS_URL."
    )
    add_code_excerpt(document, code_excerpt("desktop/src/protocol.ts", 1, 128))
    add_caption(document, "Listing 4.11.11: Typed desktop bridge protocol and WebSocket URL selection.")
    add_code_excerpt(document, code_excerpt("desktop/src/hooks/useJarvisSocket.ts", 1, 72))
    add_caption(document, "Listing 4.11.12: Desktop WebSocket connection, event dispatch, and reconnect backoff.")

    document.add_paragraph(
        "Application state is stored in Zustand. Runtime events update dialogue state, transcripts, response, metrics, errors, and config. A subset of UI preferences is persisted to localStorage and rehydrated across separate Tauri windows."
    )
    add_code_excerpt(document, code_excerpt("desktop/src/stores/jarvisStore.ts", 1, 66))
    add_caption(document, "Listing 4.11.13: Desktop store state shape.")
    add_code_excerpt(document, code_excerpt("desktop/src/stores/jarvisStore.ts", 67, 163))
    add_caption(document, "Listing 4.11.14: Engine event reducer, persisted preferences, and cross-window sync.")

    document.add_paragraph(
        "App.tsx chooses overlay or dashboard according to the current Tauri window label; in a browser it falls back to the store's appView. The overlay shows when dialogue state is active and hides after idle; the dashboard exposes local controls and sends bridge commands."
    )
    add_code_excerpt(document, code_excerpt("desktop/src/App.tsx", 1, 24))
    add_caption(document, "Listing 4.11.15: Tauri window label selects overlay or dashboard.")
    add_code_excerpt(document, code_excerpt("desktop/src/components/overlay/Overlay.tsx", 17, 59))
    add_caption(document, "Listing 4.11.16: Wake-driven overlay show/hide lifecycle.")
    add_code_excerpt(document, code_excerpt("desktop/src/components/overlay/PromptInput.tsx", 1, 31))
    add_caption(document, "Listing 4.11.17: Prompt input sends text_command with simple language detection.")
    add_code_excerpt(document, code_excerpt("desktop/src/components/dashboard/Dashboard.tsx", 120, 232))
    add_caption(document, "Listing 4.11.18: Dashboard controls for config request, view hiding, language, model, feature flags, and mute.")
    add_code_excerpt(document, code_excerpt("desktop/src/components/avatar/Avatar.tsx", 1, 21))
    add_caption(document, "Listing 4.11.19: Avatar variant selection from store state.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(document, "Bridge event loop:", [
        "if UI_BRIDGE_ENABLED is false: return without starting",
        "if FastAPI or uvicorn is missing: log optional dependency state and return",
        "start uvicorn server on ws://host:port/ws in a daemon thread",
        "on client connect: accept socket, register client, send config event",
        "on text_command: route text on a worker thread, then broadcast response event",
        "on config_request or health_request: send/broadcast current bridge-local status",
        "on engine state change: broadcast state_changed event",
        "on broadcast failure: remove stale websocket client",
    ])
    add_pseudocode(document, "Desktop reconnect and dispatch:", [
        "set connectionStatus to connecting",
        "open WebSocket at JARVIS_WS_URL",
        "on open: reset reconnect delay and set connected",
        "on message: parse JSON as EngineEvent and dispatch into Zustand store",
        "on invalid JSON: dispatch error event",
        "on close: set disconnected and schedule reconnect with delay capped at 8000 ms",
        "send UICommand only when socket readyState is OPEN",
    ])

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Setting", "Role visible from ui/ or desktop/"], ui_config_rows())
    add_caption(document, "Table 4.11.4: UI and desktop configuration surface observed in this phase.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(document, ["Situation", "UI/desktop behavior", "Evidence"], [
        ("Normal", "Bridge starts on /ws, sends config on connection, broadcasts state and response events, desktop hook dispatches events into Zustand, and overlay/dashboard render from store state.", "ui/bridge.py, ui/events.py, desktop/src/protocol.ts, useJarvisSocket.ts, jarvisStore.ts, App.tsx."),
        ("Headless", "If the bridge is disabled or optional dependencies are missing, start() returns without raising; tray also disables itself when pystray is unavailable.", "ui/bridge.py start(); ui/tray.py start()."),
        ("Disconnected", "Desktop connection status becomes disconnected and the hook retries with exponential backoff; dashboard can show the disconnected status from the store.", "useJarvisSocket.ts; Dashboard.tsx."),
        ("UI Absent", "Tauri window commands affect only desktop windows; closeApp notes that the engine is separate, and Python bridge/tray are optional surfaces over the engine.", "desktop/src/lib/app.ts; ui/bridge.py; ui/tray.py."),
        ("Partial Commands", "setting_update and feature_flag commands are currently logged by the bridge rather than applied; the dashboard updates local UI state before sending them.", "ui/bridge.py _handle_message(); Dashboard.tsx."),
    ])
    add_caption(document, "Table 4.11.5: UI behavior under normal, headless, disconnected, absent, and partial-command conditions.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "ui/ and desktop/ add observability and hands-on control without making the assistant dependent on a graphical surface. "
        "The bridge exposes a narrow JSON protocol, the tray gives a lightweight operator menu, and the Tauri/React desktop turns engine state into an overlay and dashboard while preserving a headless runtime path."
    )


def fill_phase10(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-ui", add_phase10_content)


TEST_FILE_RESPONSIBILITIES = {
    "tests/__init__.py": "Marks tests/ as an importable package.",
    "tests/test_codeswitch_routing.py": "Exercises mixed Arabic/English command parsing for app, volume, folder, delete, rename, and move cases, including fall-through for ambiguous or nonexistent targets.",
    "tests/test_llm_routing_guard.py": "Checks that advice questions and corrupted career questions remain LLM queries instead of command chains, timers, or live-search commands; also checks repeated-question cleanup and low-value answer repair.",
    "tests/test_memory_layer.py": "Covers fast context, recent-turn context, semantic-recall bypass for short queries, and the latency ceiling around RAM-only memory access.",
    "tests/test_memory_reference_resolver.py": "Covers bilingual freshness-gated reference rewriting for last app/file and blocks vague destructive pronoun deletes.",
    "tests/test_pending_task_memory.py": "Covers RAM-only pending task slot filling, dispatch after all required slots are present, TTL expiry, and no-task behavior.",
    "tests/test_routing_safety.py": "Runs the labeled NLU eval set and asserts zero unsafe auto-executions and zero question-to-command false fires.",
    "tests/test_semantic_margin.py": "Asserts near-tie semantic routing defers instead of guessing while a clean app-open command still executes.",
    "tests/test_sentence_buffer.py": "Covers English and Arabic sentence buffering, punctuation boundaries, hard flush thresholds, and Arabic connector holding.",
    "tests/test_tts_prosody.py": "Covers voice polish normalization for repeated punctuation, markdown markers, ellipses, hyphen handling, and Arabic discourse/formal connectors.",
    "tests/test_voice_normalizer.py": "Covers spoken-form normalization for weather, units, URLs, search blocks, Arabic text, times, dates, speed, and percentages.",
    "tests/fixtures/nlu_eval_cases.jsonl": "Stores labeled routing-evaluation cases used by the safety tests and NLU evaluation harness.",
}


SCRIPT_RESPONSIBILITIES = {
    "scripts/generate_arabic_wake_data.py": "Generates unified English/Arabic wake-word WAVs for openWakeWord training; its internals are cross-referenced to Section 4.5.",
    "scripts/rebuild_graduation_doc.py": "Builds and incrementally fills docs/jarvis_documentation_book.docx, including figures, tables, code excerpts, placeholders, and per-phase dispatch.",
    "scripts/setup_windows.ps1": "Windows bootstrap helper that checks Python, upgrades pip, installs requirements.txt, creates .env from .env.example when absent, and runs core\\doctor.py.",
    "scripts/train_arabic_wake_model.py": "Trains and exports the unified English/Arabic wake-word ONNX model; its training flow is cross-referenced to Section 4.5.",
}


def test_inventory_rows() -> list[tuple[str, str, str]]:
    rows = []
    for relative, role in TEST_FILE_RESPONSIBILITIES.items():
        path = ROOT / relative
        if path.exists():
            detail = f"{path.stat().st_size} bytes"
        else:
            detail = "not present"
        rows.append((relative, detail, role))
    return rows


def fixture_stats_rows() -> list[tuple[str, str, str]]:
    path = ROOT / "tests" / "fixtures" / "nlu_eval_cases.jsonl"
    if not path.exists():
        return [("nlu_eval_cases.jsonl", "missing", "Fixture file is not present.")]
    counts: dict[str, int] = {}
    total = 0
    should_execute_false = 0
    should_clarify = 0
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        item = json.loads(line)
        total += 1
        category = str(item.get("category", "uncategorized"))
        counts[category] = counts.get(category, 0) + 1
        if item.get("should_execute") is False:
            should_execute_false += 1
        if item.get("should_clarify"):
            should_clarify += 1
    rows = [
        ("Total cases", str(total), "All non-empty JSONL rows in tests/fixtures/nlu_eval_cases.jsonl."),
        ("Cases where should_execute is false", str(should_execute_false), "Negative or non-executing cases in the fixture."),
        ("Cases where should_clarify is true", str(should_clarify), "Cases explicitly requiring clarification."),
    ]
    for category, count in sorted(counts.items()):
        rows.append((category, str(count), "Fixture category count."))
    return rows


def script_inventory_rows() -> list[tuple[str, str, str]]:
    rows = []
    for relative, role in SCRIPT_RESPONSIBILITIES.items():
        path = ROOT / relative
        if path.exists():
            detail = f"{path.stat().st_size} bytes"
        else:
            detail = "not present"
        rows.append((relative, detail, role))
    return rows


def model_asset_rows() -> list[tuple[str, str, str, str]]:
    model_root = ROOT / "models"
    if not model_root.exists():
        return [("models/", "absent", "n/a", "No model directory is present.")]
    rows = []
    for path in sorted(model_root.rglob("*")):
        if not path.is_file():
            continue
        relative = path.relative_to(ROOT).as_posix()
        if path.suffix.lower() == ".onnx" and "jarvis_unified" in relative:
            subsystem = "Wake-word model asset loaded by the audio wake-word runtime; see Sections 4.4 and 4.5."
        else:
            subsystem = "Model asset present under models/."
        rows.append((relative, path.suffix.lstrip(".") or "file", f"{path.stat().st_size} bytes", subsystem))
    if not rows:
        rows.append(("models/", "empty", "0 files", "No model files were listed."))
    return rows


def add_phase11_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "tests/, scripts/, and models/ are the repository's verification and support-asset layer. "
        "The tests assert routing safety, memory behavior, sentence buffering, and spoken text normalization. "
        "The scripts provide documentation generation, Windows setup, and wake-word data/model tooling. "
        "The models directory contains the packaged ONNX wake-word assets consumed by the audio runtime described in Sections 4.4 and 4.5."
    )

    document.add_heading("4.12 tests/", level=2)
    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Observed size", "What it verifies or stores"], test_inventory_rows())
    add_caption(document, "Table 4.12.1: Inventory of tests/ files and fixture coverage.")

    add_diagram_image(document, add_qa_assets_diagram(), "Figure 4.12.1: QA and support-asset flow across tests, scripts, models, and the runtime documentation.")

    document.add_heading("NLU Fixture Coverage", level=3)
    add_styled_table(document, ["Fixture category or metric", "Count", "Meaning"], fixture_stats_rows())
    add_caption(document, "Table 4.12.2: Counts derived from tests/fixtures/nlu_eval_cases.jsonl.")

    document.add_heading("Safety And Routing Tests", level=3)
    document.add_paragraph(
        "The safety suite runs labeled NLU cases through a decision-only route. It does not dispatch OS actions, mutate session memory, or call a network LLM. "
        "Its hard assertions are zero unsafe executions and zero question-to-command false fires."
    )
    add_code_excerpt(document, code_excerpt("tests/test_routing_safety.py", 1, 30))
    add_caption(document, "Listing 4.12.1: Safety invariants over the labeled NLU eval set.")
    add_code_excerpt(document, code_excerpt("tests/test_semantic_margin.py", 1, 26))
    add_caption(document, "Listing 4.12.2: Semantic near-tie guard and clean-match execution check.")
    add_code_excerpt(document, code_excerpt("tests/test_codeswitch_routing.py", 8, 43))
    add_caption(document, "Listing 4.12.3: Code-switch routing tests for apps, volume, folders, and ambiguous text.")
    add_code_excerpt(document, code_excerpt("tests/test_llm_routing_guard.py", 13, 69))
    add_caption(document, "Listing 4.12.4: LLM routing guard tests for advice questions and response cleanup.")

    document.add_heading("Memory And Multi-Turn Tests", level=3)
    add_code_excerpt(document, code_excerpt("tests/test_memory_layer.py", 8, 60))
    add_caption(document, "Listing 4.12.5: Fast memory context and LLM-context behavior tests.")
    add_code_excerpt(document, code_excerpt("tests/test_memory_reference_resolver.py", 9, 75))
    add_caption(document, "Listing 4.12.6: Freshness-gated app/file reference resolver tests.")
    add_code_excerpt(document, code_excerpt("tests/test_pending_task_memory.py", 9, 52))
    add_caption(document, "Listing 4.12.7: Pending task memory tests for slot filling and TTL expiry.")

    document.add_heading("Voice Text Utility Tests", level=3)
    add_code_excerpt(document, code_excerpt("tests/test_sentence_buffer.py", 6, 27))
    add_caption(document, "Listing 4.12.8: Sentence buffer boundary and connector-hold tests.")
    add_code_excerpt(document, code_excerpt("tests/test_tts_prosody.py", 1, 46))
    add_caption(document, "Listing 4.12.9: TTS prosody tests for repeated punctuation, markdown, dashes, and ellipses.")
    add_code_excerpt(document, code_excerpt("tests/test_voice_normalizer.py", 6, 35))
    add_caption(document, "Listing 4.12.10: Voice normalizer tests for weather, units, and URL cleanup.")

    document.add_heading("Test Philosophy", level=3)
    add_pseudocode(document, "Safety-oriented routing checks:", [
        "load labeled JSONL cases from tests/fixtures/nlu_eval_cases.jsonl",
        "for each utterance, call the decision-only router instead of the dispatcher",
        "record unsafe_execution when a case that must not execute would execute",
        "record question_false_fire when a question is routed as an executable command",
        "fail the suite unless unsafe_execution_count == 0",
        "fail the suite unless question_to_command_false_fires == 0",
        "for semantic near-ties, require clarification/defer behavior rather than silent execution",
    ])

    document.add_heading("4.13 scripts/", level=2)
    document.add_heading("Developer Tool Inventory", level=3)
    add_styled_table(document, ["Script", "Observed size", "Role"], script_inventory_rows())
    add_caption(document, "Table 4.13.1: scripts/ inventory and developer support roles.")
    document.add_paragraph(
        "rebuild_graduation_doc.py is the active generator for this Word book: it uses python-docx and Pillow, stores the output at docs/jarvis_documentation_book.docx, "
        "creates generated figures under docs/generated_figures, and fills one placeholder per requested phase. setup_windows.ps1 is a local Windows bootstrap helper. "
        "The wake-word data and training scripts are present here, but their pipeline is documented in Section 4.5 to avoid duplicating training internals."
    )
    add_code_excerpt(document, code_excerpt("scripts/rebuild_graduation_doc.py", 1, 31))
    add_caption(document, "Listing 4.13.1: Documentation generator imports, output paths, and purpose.")
    add_code_excerpt(document, function_excerpt("scripts/rebuild_graduation_doc.py", "insert_generated_content"))
    add_caption(document, "Listing 4.13.2: Placeholder-preserving insertion helper used by phase fills.")
    add_code_excerpt(document, code_excerpt("scripts/setup_windows.ps1", 1, 22))
    add_caption(document, "Listing 4.13.3: Windows setup script steps.")
    add_code_excerpt(document, code_excerpt("scripts/generate_arabic_wake_data.py", 1, 14))
    add_caption(document, "Listing 4.13.4: Wake data-generation script purpose and local-data warning.")
    add_code_excerpt(document, code_excerpt("scripts/train_arabic_wake_model.py", 1, 10))
    add_caption(document, "Listing 4.13.5: Wake model training script purpose; implementation details are covered in Section 4.5.")

    document.add_heading("4.14 models/", level=2)
    document.add_heading("Model Asset Inventory", level=3)
    add_styled_table(document, ["Model asset", "Format", "Observed size", "Subsystem"], model_asset_rows())
    add_caption(document, "Table 4.14.1: Model files listed under models/.")
    document.add_paragraph(
        "The listed ONNX files live under models/jarvis_unified/. The primary artifact name matches the default unified wake-word model path documented earlier, "
        "while the backup ONNX file preserves another local model artifact in the same subsystem. Section 4.4 documents runtime wake-model loading and Section 4.5 documents the dataset/training artifact boundary."
    )

    document.add_heading("Behavior And QA Coverage", level=3)
    add_styled_table(document, ["Area", "What a passing suite or verified asset proves", "Evidence in this phase"], [
        ("Routing safety", "Labeled NLU cases do not produce unsafe auto-executions or question-to-command false fires.", "test_routing_safety.py; nlu_eval_cases.jsonl."),
        ("Ambiguity handling", "Bare near-tie commands defer instead of guessing, while clean commands still execute.", "test_semantic_margin.py."),
        ("Code-switch parsing", "Mixed Arabic/English commands map to app, volume, folder, and file-operation intents when resolvable.", "test_codeswitch_routing.py."),
        ("Memory behavior", "Fast context stays RAM-oriented, follow-up references resolve only with fresh context, and pending tasks expire.", "test_memory_layer.py; test_memory_reference_resolver.py; test_pending_task_memory.py."),
        ("Voice text shaping", "Sentence buffering, prosody cleanup, and voice normalization keep spoken output cleaner before TTS.", "test_sentence_buffer.py; test_tts_prosody.py; test_voice_normalizer.py."),
        ("Developer repeatability", "The book build, Windows setup, wake-data generation, wake training, and packaged ONNX assets are discoverable from repository files.", "scripts/ inventory; models/ inventory."),
    ])
    add_caption(document, "Table 4.14.2: Behavior and QA coverage supplied by tests, scripts, and model assets.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "Phase 11 makes the support layer explicit: tests encode the project's safety and regression contracts, scripts encode repeatable developer workflows, "
        "and models/ stores the wake-word artifacts that connect the offline pipeline to the runtime audio subsystem."
    )


def fill_phase11(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-tests", add_phase11_content)


def add_main_excerpt(document: Document) -> None:
    source = (ROOT / "main.py").read_text(encoding="utf-8")
    excerpt = "\n".join(source.rstrip().splitlines()[:36])
    add_code_excerpt(document, excerpt)


def get_config_defaults() -> dict[str, str]:
    text = (ROOT / "core" / "config.py").read_text(encoding="utf-8")
    tree = ast.parse(text)
    defaults: dict[str, str] = {}
    for node in ast.walk(tree):
        if not isinstance(node, ast.Call):
            continue
        if not isinstance(node.func, ast.Name):
            continue
        if node.func.id not in {"_env", "_env_int", "_env_float", "_env_bool", "_env_list"}:
            continue
        if not node.args or not isinstance(node.args[0], ast.Constant):
            continue
        key = str(node.args[0].value)
        if len(node.args) < 2:
            defaults[key] = ""
            continue
        try:
            defaults[key] = repr(ast.literal_eval(node.args[1]))
        except Exception:
            defaults[key] = ast.unparse(node.args[1])
    return defaults


def add_config_excerpt(document: Document) -> None:
    lines = (ROOT / "core" / "config.py").read_text(encoding="utf-8").splitlines()
    excerpt = "\n".join(lines[20:44] + ["", lines[66]])
    add_code_excerpt(document, excerpt)


def dependency_rows() -> list[tuple[str, str, str]]:
    req = (ROOT / "requirements.txt").read_text(encoding="utf-8")
    rows = []
    current = "General"
    packages: list[str] = []
    for line in req.splitlines():
        stripped = line.strip()
        if stripped.startswith("# ==="):
            if packages:
                rows.append((current, ", ".join(packages[:7]), "runtime"))
            current = stripped.strip("#= ").replace("TIER ", "Tier ")
            packages = []
        elif stripped and not stripped.startswith("#"):
            packages.append(stripped.split("#", 1)[0].strip())
    if packages:
        rows.append((current, ", ".join(packages[:7]), "runtime"))
    training = (ROOT / "requirements-training.txt").read_text(encoding="utf-8")
    train_packages = [line.strip() for line in training.splitlines() if line.strip() and not line.startswith("#")]
    rows.append(("Training only", ", ".join(train_packages), "wake-word development"))
    return rows


def root_inventory_rows() -> list[tuple[str, str]]:
    return [
        ("main.py", "Thin launcher: migrates legacy paths, parses --demo-mode, starts optional tray/bridge, then calls core.orchestrator.run()."),
        ("core/", "Orchestration, routing, configuration, memory logic, persona/response shaping, diagnostics, and handlers."),
        ("audio/", "Wake-word detection, microphone capture, VAD, STT, and TTS runtime components."),
        ("nlp/", "Text normalization, code-switch routing, semantic/fuzzy/keyword routing, and slot/entity extraction."),
        ("llm/", "Ollama/Claude clients, prompt construction, sentence buffering, and structured NLU fallback."),
        ("os_control/", "Windows side-effect adapters, policy/risk gates, confirmation, persistence, and verification."),
        ("tools/", "Live-data and utility tools such as weather, web search, and calculator."),
        ("utils/", "Shared helpers consumed by other layers."),
        ("ui/ and desktop/", "Optional tray, WebSocket bridge, and Tauri desktop UI; the engine can run without them."),
        ("data/", "Runtime artifact home for logs, memory, indices, state, KB, vectors, VAD data, and wake samples."),
        ("models/", "Tracked model assets, including the deployed wake-word ONNX model."),
        ("scripts/", "Developer and wake-word training/maintenance scripts."),
        ("tests/", "Regression and safety checks for routing, language, memory, TTS, and other behavior."),
        ("requirements.txt", "Production runtime dependency set, grouped by feature tier."),
        ("requirements-training.txt", "Development-only dependencies for wake-word training."),
        (".env.example", "Configuration template; every documented JARVIS_* key carries an explicit default or blank override slot."),
        ("README.md", "Top-level operator guide and architecture summary."),
    ]


def data_layout_rows() -> list[tuple[str, str]]:
    rows = []
    descriptions = {
        "chroma_memory": "ChromaDB/vector-memory artifacts from earlier or alternate semantic-memory storage.",
        "index": "Search-index cache files, including the configured jarvis_index.db path.",
        "kb": "Offline knowledge-base index and metadata files.",
        "logs": "Structured runtime logs and action audit logs.",
        "memory": "SQLite/JSON memory store; documented in depth in Section 4.3.",
        "state": "Runtime state database and small persisted state artifacts.",
        "vad": "VAD-related runtime artifacts or cached data.",
        "vectors": "Vector semantic-memory store configured by JARVIS_VECTOR_MEMORY_DIR.",
        "wake_samples": "Wake-word enrollment and adaptive-retraining samples; training pipeline is documented in Section 4.5.",
    }
    for child in sorted((ROOT / "data").iterdir(), key=lambda p: p.name.lower()):
        if child.is_dir():
            rows.append((child.name + "/", descriptions.get(child.name, "Runtime data directory present in this workspace.")))
        else:
            rows.append((child.name, "Top-level runtime data file."))
    return rows


def add_phase1_content(document: Document) -> None:
    defaults = get_config_defaults()
    figure = add_root_layout_diagram()

    document.add_heading("4.1 Repository Layout", level=2)
    document.add_paragraph(
        "Chapter 4 documents the implementation from the repository root outward. "
        "This root section fixes the boundary conditions: what starts the assistant, "
        "where configuration enters, which dependency tiers are installed, and where "
        "runtime artifacts are placed. Subsystem mechanics are intentionally deferred "
        "to their own folder chapters."
    )
    add_diagram_image(document, figure, "Figure 4.1.1: Repository layout and root-level responsibility boundaries.")
    add_styled_table(document, ["Root item", "Responsibility"], root_inventory_rows())
    add_caption(document, "Table 4.1.1: Root-level files and folders with their implementation responsibilities.")

    document.add_heading("4.2 Entry Point main.py", level=2)
    document.add_paragraph(
        "The launcher keeps policy and assistant behavior outside the root file. It performs only process setup: "
        "legacy data-path migration is imported and executed, --demo-mode sets JARVIS_DEMO_MODE, optional tray and "
        "WebSocket bridge startup are attempted defensively, and the assistant engine is handed to core.orchestrator.run()."
    )
    add_main_excerpt(document)
    add_caption(document, "Listing 4.1.1: The root entry point delegates runtime behavior to the core orchestrator.")

    document.add_heading("4.3 Configuration core/config.py", level=2)
    document.add_paragraph(
        "Configuration is centralized in core/config.py. The file loads .env from the project root with override=True, "
        "then normalizes environment variables through typed helpers for strings, integers, floats, booleans, and lists. "
        "DATA_DIR defaults to the project data/ directory, and config.py ensures logs, memory, index, state, kb, and "
        "vectors exist on import."
    )
    add_config_excerpt(document)
    add_caption(document, "Listing 4.1.2: The configuration helper pattern and DATA_DIR root from core/config.py.")
    representative_keys = [
        ("JARVIS_DATA_DIR", defaults.get("JARVIS_DATA_DIR", "data"), "Root for runtime artifacts."),
        ("JARVIS_UI_BRIDGE_ENABLED", defaults.get("JARVIS_UI_BRIDGE_ENABLED", "True"), "Enables the optional desktop UI bridge."),
        ("JARVIS_UI_BRIDGE_PORT", defaults.get("JARVIS_UI_BRIDGE_PORT", "9720"), "Local WebSocket bridge port."),
        ("JARVIS_MAX_RECORD_DURATION", defaults.get("JARVIS_MAX_RECORD_DURATION", "8.0"), "Maximum command recording window."),
        ("JARVIS_WAKE_WORD_UNIFIED_ONNX_PATH", defaults.get("JARVIS_WAKE_WORD_UNIFIED_ONNX_PATH", "'models/jarvis_unified/jarvis_unified.onnx'"), "Bilingual wake-word model path."),
        ("JARVIS_STT_BACKEND", defaults.get("JARVIS_STT_BACKEND", "'hybrid_elevenlabs'"), "Primary STT strategy."),
        ("JARVIS_LLM_BACKEND", defaults.get("JARVIS_LLM_BACKEND", "'ollama'"), "LLM backend selection."),
        ("JARVIS_LLM_MODEL", defaults.get("JARVIS_LLM_MODEL", "'qwen3:4b'"), "Pinned/default local model before hardware auto-selection."),
        ("JARVIS_SEMANTIC_MIN_MARGIN", defaults.get("JARVIS_SEMANTIC_MIN_MARGIN", "0.08"), "Semantic router near-tie guard."),
        ("JARVIS_MEMORY_BACKEND", defaults.get("JARVIS_MEMORY_BACKEND", "'sqlite'"), "Primary memory persistence backend."),
        ("JARVIS_SECOND_FACTOR_PIN", defaults.get("JARVIS_SECOND_FACTOR_PIN", "'1234'"), "Spoken PIN for sensitive actions."),
        ("JARVIS_CONTROLS_VERIFY_STATE", defaults.get("JARVIS_CONTROLS_VERIFY_STATE", "True"), "Require read-back verification for controls."),
        ("JARVIS_WEATHER_CITY", defaults.get("JARVIS_WEATHER_CITY", "'Cairo'"), "Default weather location."),
        ("JARVIS_LOG_FILE", "data/logs/jarvis.log unless overridden", "Main runtime log path."),
    ]
    add_styled_table(document, ["Key", "Default from config.py", "Purpose"], representative_keys)
    add_caption(document, "Table 4.1.2: Representative JARVIS_* configuration keys at the root boundary.")

    document.add_heading("4.4 Environment and Dependency Files", level=2)
    document.add_paragraph(
        ".env.example is the operator-facing configuration template. It documents root paths, logs, security, audio, "
        "NLP, LLM, memory, controls, wake-word, and live-data settings with defaults or intentionally blank override "
        "slots. Secrets such as API keys are configuration inputs, not code constants."
    )
    add_styled_table(document, ["Tier", "Representative packages", "Role"], dependency_rows())
    add_caption(document, "Table 4.1.3: Dependency tiers declared by requirements.txt and requirements-training.txt.")

    document.add_heading("4.5 The data/ Runtime Tree", level=2)
    document.add_paragraph(
        "The root perspective treats data/ as the runtime artifact tree. The memory store and wake-word samples are "
        "listed here only at a glance because Section 4.3 documents persisted memory in depth, and Section 4.5 documents "
        "wake-word data/training provenance."
    )
    add_styled_table(document, ["data/ item", "Root-level meaning"], data_layout_rows())
    add_caption(document, "Table 4.1.4: Top-level data/ runtime directories observed in this workspace.")

    document.add_heading("4.6 Root-Perspective Startup Sequence", level=2)
    add_styled_table(
        document,
        ["Step", "Root-visible behavior", "Deferred detail"],
        [
            ("1", "Import core.data_migration.migrate_legacy_paths() and execute it before runtime starts.", "Migration mechanics are covered in core/."),
            ("2", "Parse --demo-mode; if present, set JARVIS_DEMO_MODE=1 for downstream modules.", "Demo overlay behavior is covered with routing/core behavior."),
            ("3", "Attempt to start the optional system tray without making it a hard dependency.", "UI tray behavior is covered in Section 4.11."),
            ("4", "Attempt to start the optional WebSocket bridge for the desktop UI.", "Bridge protocol is covered in Section 4.11."),
            ("5", "Call core.orchestrator.run() to begin the assistant loop.", "Wake/listen/process mechanics are covered in core/ and audio/."),
        ],
    )
    add_caption(document, "Table 4.1.5: Startup behavior visible from the repository root.")
    add_callout(
        document,
        "Phase 1 deliberately stops at the root boundary. It does not describe subsystem internals; later phases fill those bookmarked sections with folder-local evidence.",
        "Callout OK",
    )


def build_phase1_fragment() -> Document:
    document = Document()
    configure_styles(document)
    add_phase1_content(document)
    return document


def fill_phase1(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        build_skeleton(doc_path)
    document = Document(str(doc_path))
    try:
        placeholder = find_placeholder_paragraph(document, "c4")
    except ValueError:
        build_skeleton(doc_path)
        document = Document(str(doc_path))
        placeholder = find_placeholder_paragraph(document, "c4")

    body = document._body._element
    insertion_index = body.index(placeholder._element)
    bookmark_id = "999"
    for child in placeholder._element.iter():
        if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == "c4":
            bookmark_id = child.get(qn("w:id")) or bookmark_id
            break
    original_len = len(body)
    add_phase1_content(document)

    # python-docx inserts new body blocks immediately before the final sectPr.
    new_elements = list(body)[original_len - 1 : -1]
    for element in new_elements:
        body.remove(element)
    body.remove(placeholder._element)
    for offset, element in enumerate(new_elements):
        if offset == 0 and element.tag == qn("w:p"):
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), bookmark_id)
            start.set(qn("w:name"), "c4")
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), bookmark_id)
            element.insert(0, start)
            element.append(end)
        body.insert(insertion_index + offset, element)

    document.save(str(doc_path))
    return doc_path


def insert_generated_content(doc_path: Path, placeholder_id: str, add_content) -> Path:
    document = Document(str(doc_path))
    placeholder = find_placeholder_paragraph(document, placeholder_id)
    body = document._body._element
    insertion_index = body.index(placeholder._element)
    bookmark_id = "999"
    for child in placeholder._element.iter():
        if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == placeholder_id:
            bookmark_id = child.get(qn("w:id")) or bookmark_id
            break

    original_len = len(body)
    add_content(document)
    new_elements = list(body)[original_len - 1 : -1]
    for element in new_elements:
        body.remove(element)
    body.remove(placeholder._element)
    for offset, element in enumerate(new_elements):
        if offset == 0 and element.tag == qn("w:p"):
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), bookmark_id)
            start.set(qn("w:name"), placeholder_id)
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), bookmark_id)
            element.insert(0, start)
            element.append(end)
        body.insert(insertion_index + offset, element)
    document.save(str(doc_path))
    return doc_path


def add_phase2_content(document: Document) -> None:
    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "The core/ package is the runtime control plane for Jarvis. It receives recognized utterances from the audio "
        "layer, normalizes and routes them, verifies intent safety, dispatches to handlers or LLM response paths, and "
        "records memory, metrics, and diagnostics. In the finished Methodology chapter this section should cross-reference "
        "Figure 3.1; for now it documents the folder-local implementation evidence."
    )

    document.add_heading("File Inventory", level=3)
    add_styled_table(document, ["File", "Responsibility"], core_inventory_rows())
    add_caption(document, "Table 4.2.1: Inventory of every Python file in core/ and core/handlers/.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_core_flow_diagram(), "Figure 4.2.1: Core runtime flow from orchestrator loop through routing, verification, dispatch, memory, and metrics.")

    document.add_heading("Key Modules In Depth", level=3)
    document.add_paragraph(
        "orchestrator.py owns startup ordering and the long-lived wake/listen/process loop. It logs startup state, "
        "checks elevation, installs shutdown handling, initializes command services, warms latency-critical components, "
        "attaches interrupt targets, plays the blocking greeting, starts cache/adaptive wake background work, and then "
        "records utterances into the processing pipeline."
    )
    add_code_excerpt(document, code_excerpt("core/orchestrator.py", 1693, 1717))
    add_caption(document, "Listing 4.2.1: orchestrator.run() startup boundary.")
    add_code_excerpt(document, code_excerpt("core/orchestrator.py", 1860, 1879))
    add_caption(document, "Listing 4.2.2: wake-source delay, recording phase, and partial transcript callback.")

    document.add_paragraph(
        "command_router.py implements the cascade. It starts with parser fast-path confidence, then code-switch routing, "
        "semantic top-k/margin routing, keyword NLP, gated tool/structured-LLM tiers, and finally LLM_QUERY fallback. "
        "The route verifier is called before final dispatch metadata is logged."
    )
    add_code_excerpt(document, code_excerpt("core/command_router.py", 5102, 5160))
    add_caption(document, "Listing 4.2.3: routing cascade tiers before LLM fallback.")
    add_code_excerpt(document, code_excerpt("core/command_router.py", 5399, 5420))
    add_caption(document, "Listing 4.2.4: route_verifier hand-off from command_router.py.")

    document.add_paragraph(
        "route_verifier.py consolidates candidate checks into a RouteDecision with action values execute, clarify, "
        "confirm, or llm. The order is schema existence, required slots, question-opener demotion, entity confidence, "
        "risk gate, confidence/clarification verdict, and policy permission."
    )
    add_code_excerpt(document, code_excerpt("core/route_verifier.py", 91, 110))
    add_caption(document, "Listing 4.2.5: route_verifier.verify() input normalization and schema gate.")
    add_code_excerpt(document, code_excerpt("core/route_verifier.py", 116, 153))
    add_caption(document, "Listing 4.2.6: verifier missing-slot, question, entity, and risk decisions.")

    document.add_paragraph(
        "intent_schema.py is the folder's intent contract. Each IntentSpec declares the domain, required and optional slots, "
        "risk tier, fast-execute eligibility, and bilingual examples. The schema contains 38 registered intents in this "
        "workspace, including OS, file, safety, meta, metrics, and LLM_QUERY domains."
    )
    add_code_excerpt(document, code_excerpt("core/intent_schema.py", 40, 75))
    add_caption(document, "Listing 4.2.7: IntentSpec and a bilingual OS_APP_OPEN schema entry.")

    document.add_paragraph(
        "Memory is split by latency. memory_manager.py exposes get_fast_context() for RAM-only working slots and "
        "get_llm_context() for recent turns plus bounded semantic recall. session_memory.py owns the in-process API; "
        "memory_store.py provides SQLite and vector persistence adapters. The physical files these APIs read/write are "
        "deferred to Section 4.3."
    )
    add_code_excerpt(document, code_excerpt("core/memory_manager.py", 31, 52))
    add_caption(document, "Listing 4.2.8: RAM-only fast memory context.")
    add_code_excerpt(document, code_excerpt("core/memory_manager.py", 54, 85))
    add_caption(document, "Listing 4.2.9: LLM memory context with bounded semantic recall.")
    add_code_excerpt(document, code_excerpt("core/memory_store.py", 40, 67))
    add_caption(document, "Listing 4.2.10: SQLite memory store connection and WAL configuration.")

    document.add_paragraph(
        "metrics.py supplies shared stage timers and latency buckets used throughout routing, recording, LLM, TTS, memory, "
        "and diagnostics. doctor.py reports dependency and feature availability, while hardware_detect.py selects model "
        "tiers from RAM/GPU conditions."
    )
    add_code_excerpt(document, code_excerpt("core/metrics.py", 952, 976))
    add_caption(document, "Listing 4.2.11: stage_timer and record_stage_timing.")
    add_code_excerpt(document, code_excerpt("core/hardware_detect.py", 13, 25))
    add_caption(document, "Listing 4.2.12: Qwen3 model-tier table.")

    document.add_heading("Algorithms", level=3)
    add_pseudocode(
        document,
        "Verifier decision order:",
        [
            "candidate = routed ParsedCommand-like object",
            "spec = intent_schema.get_spec(candidate.intent)",
            "if spec is missing: return llm(no_schema_entry)",
            "if required slot is blank: return clarify(missing_slot)",
            "if question opener targets command intent: return llm(question_opener_detected)",
            "if any entity score < 0.45: return clarify(low_entity_confidence)",
            "risk = per-action override or schema risk",
            "if risk is high or medium: return confirm(risk)",
            "if assess_intent_confidence asked to clarify: return clarify(low_confidence)",
            "if policy blocks permission: return llm(policy_blocked)",
            "return execute(ok)",
        ],
    )
    add_pseudocode(
        document,
        "Hardware tier selection:",
        [
            "ram_gb = detect_total_ram_gb()",
            "gpu = detect_gpu_available(ollama_base_url)",
            "for each tier ordered high -> minimal:",
            "    if ram_gb >= min_ram and (not gpu_required or gpu):",
            "        return model, num_ctx, lightweight_num_ctx, ram_gb, gpu",
            "fallback = qwen3:0.6b with 1024 context",
        ],
    )

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(document, ["Key", "Default", "Core-owned purpose"], core_config_rows())
    add_caption(document, "Table 4.2.2: Core-owned JARVIS_* configuration surface relevant to routing, memory, startup, and diagnostics.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(
        document,
        ["Situation", "Core behavior", "Evidence"],
        [
            ("Normal", "orchestrator initializes services, records speech, route_command resolves the first confident tier, verifier logs action/reason, dispatch returns a shaped response.", "orchestrator.py run loop; command_router.py route_command; response_shaper.py."),
            ("Degraded", "If semantic routing is unavailable or not ready, command_router marks semantic_pending/unavailable and continues to keyword/tool/LLM paths; if SQLite memory fails, SessionMemory falls back to JSON.", "command_router.py _try_semantic_routing; session_memory.py initialization."),
            ("Adversarial/Edge", "Question-shaped commands are demoted to LLM, missing slots ask clarification, medium/high risk actions require confirmation, and policy blocks return an LLM-style non-execution decision.", "route_verifier.py verify(); intent_schema.py risk fields."),
            ("Operational", "RuntimeCoordinator accepts wake interrupts only in interruptible phases and cancels attached TTS/LLM work; blocked phases play/log the blocked path.", "runtime_coordinator.py RuntimeCoordinator.request_interrupt()."),
        ],
    )
    add_caption(document, "Table 4.2.3: Core behavior under normal, degraded, and edge conditions.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "core/ advances the project objectives by making the assistant deterministic at the decision boundary: fast routes "
        "stay cheap, uncertain routes ask or defer, risky routes confirm, and every turn leaves timing/memory evidence. "
        "It also keeps bilingual English/Egyptian-Arabic handling visible in schemas, templates, language gates, and persona/voice shaping while leaving physical data artifacts to their dedicated chapters."
    )


def fill_phase2(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-core", add_phase2_content)


def add_phase3_content(document: Document) -> None:
    sqlite_summary = sqlite_memory_summary()
    json_summary = json_memory_summary()

    document.add_heading("Purpose", level=3)
    document.add_paragraph(
        "data/memory/ is the durable memory artifact layer behind the in-process memory APIs documented in Section 4.2. "
        "This phase reads only the memory files on disk. The observed store preserves recent turns, working slots, app/file "
        "references, language history, clarification preferences, and preference-like state used for continuity and "
        "personalization across sessions."
    )

    document.add_heading("File and Subfolder Inventory", level=3)
    add_styled_table(document, ["Artifact", "Format", "What it stores", "Reader/writer relationship"], memory_file_inventory_rows())
    add_caption(document, "Table 4.3.1: Files observed directly under data/memory/.")

    document.add_heading("Internal Data Flow", level=3)
    add_diagram_image(document, add_memory_data_flow_diagram(), "Figure 4.3.1: Persisted memory data flow from turn/slot writes to later context reconstruction.")

    document.add_heading("Key Artifacts In Depth", level=3)
    document.add_paragraph(
        "jarvis_memory.db is the primary observed SQLite artifact. Its PRAGMA values report journal_mode="
        f"{sqlite_summary.get('journal_mode')}, page_count={sqlite_summary.get('page_count')}, and "
        f"page_size={sqlite_summary.get('page_size')}. The database currently contains "
        f"{sqlite_summary['counts'].get('turns')} rows in turns and {sqlite_summary['counts'].get('slots')} rows in slots."
    )
    table_rows = []
    for item in sqlite_summary["tables"]:
        sql = str(item.get("sql") or "").replace("\n", " ")
        sql = re.sub(r"\s+", " ", sql).strip()
        table_rows.append((item["name"], sql[:420]))
    add_styled_table(document, ["SQLite table", "Observed schema"], table_rows)
    add_caption(document, "Table 4.3.2: SQLite table schema observed in jarvis_memory.db.")

    index_rows = []
    for item in sqlite_summary["indexes"]:
        sql = str(item.get("sql") or "implicit SQLite index").replace("\n", " ")
        index_rows.append((item["name"], sql[:260]))
    add_styled_table(document, ["SQLite index", "Definition"], index_rows)
    add_caption(document, "Table 4.3.3: SQLite indexes observed in jarvis_memory.db.")

    document.add_paragraph(
        "The turns table stores bounded recent conversation rows with timestamp, user text, assistant text, language, and "
        "intent. To avoid leaking transcript content in the thesis artifact, the table below records only ids, language, "
        "intent, and text lengths."
    )
    add_styled_table(
        document,
        ["Recent id", "Timestamp", "Language", "Intent", "User chars", "Assistant chars"],
        [
            (
                str(row.get("id", "")),
                str(row.get("timestamp", "")),
                str(row.get("language", "")),
                str(row.get("intent", "")),
                str(row.get("user_len", "")),
                str(row.get("assistant_len", "")),
            )
            for row in sqlite_summary["turn_sample"]
        ],
    )
    add_caption(document, "Table 4.3.4: Sanitized recent-turn sample from the turns table.")

    document.add_paragraph(
        "The slots table stores JSON-encoded working state. Observed slot names include app/file references, pending "
        "confirmation state, clarification preferences, app and command usage, response mode, audio/STT profiles, "
        "preferred language, and bilingual language history."
    )
    add_styled_table(document, ["Slot name", "JSON category", "updated_at"], sqlite_summary["slots"])
    add_caption(document, "Table 4.3.5: Slot names and value categories observed in the slots table.")

    document.add_paragraph(
        "jarvis_memory.json is present as a JSON artifact with top-level keys "
        f"{', '.join(json_summary.get('keys') or [])}. It contains {json_summary.get('turn_count')} turn objects, "
        f"preferred_language={json_summary.get('preferred_language') or 'blank'}, and "
        f"{len(json_summary.get('context_slots') or [])} context slot keys. The SQLite slots table also contains "
        "__legacy_json_imported__, which records that the JSON import path has already been consumed."
    )
    add_styled_table(document, ["JSON context slot", "Observed type"], json_summary.get("context_slots") or [])
    add_caption(document, "Table 4.3.6: Context slot keys observed in jarvis_memory.json.")

    document.add_paragraph(
        "jarvis_memory.db-shm and jarvis_memory.db-wal are SQLite WAL-mode companions. The observed WAL file is zero bytes "
        "at inspection time, which is consistent with no pending uncheckpointed pages. The SHM file is the shared-memory "
        "sidecar used by SQLite while WAL mode is active."
    )

    document.add_heading("Algorithms Reflected By The Data", level=3)
    add_pseudocode(
        document,
        "Observed turn/slot persistence pattern:",
        [
            "on each remembered turn:",
            "    append or retain a row in turns(timestamp, user, assistant, language, intent)",
            "    update slots for working references such as last_app, last_file, language_history",
            "    encode complex values as JSON strings in slots.value",
            "    keep SQLite in WAL mode for safer concurrent reads/writes",
            "later context reconstruction:",
            "    read recent rows from turns",
            "    read named slots and decode JSON",
            "    combine language history, references, preferences, and recent turns",
        ],
    )
    add_pseudocode(
        document,
        "Data-only retention observation:",
        [
            f"observed turns row count = {sqlite_summary['counts'].get('turns')}",
            "observed turn ids are high while retained row count is small",
            "therefore the data suggests a bounded recent-turn store",
            "the exact pruning trigger is not claimed here because this phase did not read source code",
        ],
    )

    document.add_heading("Configuration Surface", level=3)
    add_styled_table(
        document,
        ["Storage setting", "Observed path/value", "Meaning from data perspective"],
        [
            ("Memory directory", "data/memory/", "Folder containing the persisted memory artifacts inspected in this phase."),
            ("Primary database", "data/memory/jarvis_memory.db", "SQLite store with turns and slots tables."),
            ("WAL sidecars", "jarvis_memory.db-wal and jarvis_memory.db-shm", "SQLite journal/shared-memory companions."),
            ("Legacy/debug JSON", "data/memory/jarvis_memory.json", "JSON state file with turns and context_slots."),
            ("Language state", "preferred_language slot plus language_history arrays", "Observed bilingual continuity state."),
        ],
    )
    add_caption(document, "Table 4.3.7: Memory storage surface visible from persisted artifacts.")

    document.add_heading("Behavior In Different Situations", level=3)
    add_styled_table(
        document,
        ["Situation", "Observed artifact behavior", "Defense note"],
        [
            ("Normal", "SQLite database contains turns and slots; WAL mode is active; language_history includes English and Arabic entries.", "Supports continuity and bilingual context retention."),
            ("Degraded", "JSON file remains available as a legacy/debug state artifact alongside SQLite.", "The data itself shows a migration/import marker but this chapter does not claim source-level fallback mechanics."),
            ("Adversarial/Edge", "Transcript text is stored in turns, while thesis excerpts sanitize content to lengths only; slots use JSON values that can represent empty pending tasks and cleared confirmation tokens.", "Sensitive or noisy state can be audited without exposing full utterance text in documentation."),
            ("Corruption/Concurrency", "WAL and SHM companions are present; WAL length can be zero after checkpointing.", "SQLite journaling artifacts are self-contained in data/memory/."),
        ],
    )
    add_caption(document, "Table 4.3.8: Memory data behavior under normal, degraded, and edge situations.")

    document.add_heading("Contribution Summary", level=3)
    document.add_paragraph(
        "The persisted memory store advances continuity by keeping recent turns, working references, preferences, and "
        "language history across launches. Its observed English/Arabic language-history slots and preferred-language state "
        "support bilingual interaction without requiring every turn to rediscover context from scratch."
    )


def fill_phase3(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    return insert_generated_content(doc_path, "c4-memdata", add_phase3_content)


def add_conceptual_overview_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_1_1_1_conceptual_overview.png"
    width, height = 1900, 620
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 28)
        small_font = ImageFont.truetype("arial.ttf", 21)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()
    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((70, 40), "Figure 1.1 - Conceptual turn cycle", fill=colors["ink"], font=title_font)
    draw.text((70, 100), "The seven-stage loop that every spoken turn passes through, bilingually.", fill=colors["muted"], font=small_font)
    stages = [
        ("Wake", "unified EN/EGY ONNX", colors["blue"]),
        ("Record", "VAD-gated capture", colors["blue"]),
        ("STT", "language-locked transcript", colors["purple"]),
        ("Understand", "routing cascade", colors["purple"]),
        ("Act / Answer", "OS control or LLM", colors["amber"]),
        ("Speak", "streamed TTS", colors["green"]),
    ]
    x = 60
    y = 260
    box_w, box_h, gap = 260, 150, 55
    for i, (title, subtitle, color) in enumerate(stages):
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + box_h), fill=color)
        draw.text((x + 26, y + 34), title, fill=colors["ink"], font=font)
        draw.text((x + 26, y + 82), subtitle, fill=colors["muted"], font=small_font)
        if i < len(stages) - 1:
            ax = x + box_w
            draw.line((ax, y + box_h // 2, ax + gap, y + box_h // 2), fill=colors["line"], width=5)
            draw.polygon(
                [(ax + gap, y + box_h // 2), (ax + gap - 14, y + box_h // 2 - 8), (ax + gap - 14, y + box_h // 2 + 8)],
                fill=colors["line"],
            )
        x += box_w + gap
    draw.text((70, 460), "Loop closes back to Wake after Speak; a wake-word interrupt during Act/Answer/Speak cancels the in-flight stage.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_layered_architecture_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_3_1_1_layered_architecture.png"
    width, height = 1900, 1160
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 50)
        font = ImageFont.truetype("arial.ttf", 27)
        small_font = ImageFont.truetype("arial.ttf", 21)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()
    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((70, 35), "Figure 3.1 - Layered architecture", fill=colors["ink"], font=title_font)
    draw.text((70, 95), "Five horizontal layers; core/ is the control plane that every other layer reports to.", fill=colors["muted"], font=small_font)

    layers = [
        ("I/O Layer", "audio/ (wake, VAD, STT, TTS)", colors["blue"], 150),
        ("Understanding Layer", "nlp/ (code-switch, semantic, fuzzy, keyword, slots)", colors["purple"], 320),
        ("Control Plane", "core/ (orchestrator, router, verifier, memory, persona, metrics)", colors["amber"], 490),
        ("Action / Answer Layer", "os_control/, tools/, llm/ (verified side effects and generation)", colors["green"], 660),
        ("Presentation Layer", "ui/ + desktop/ (optional tray + WebSocket bridge + Tauri UI)", colors["grey"], 830),
    ]
    for title, subtitle, color, y in layers:
        draw.rounded_rectangle((90, y, 1810, y + 130), radius=18, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((90, y, 102, y + 130), fill=color)
        draw.text((130, y + 24), title, fill=colors["ink"], font=font)
        draw.text((130, y + 72), subtitle, fill=colors["muted"], font=small_font)
        if y != 830:
            draw.line((950, y + 130, 950, y + 190), fill=colors["line"], width=5)
            draw.polygon([(950, y + 190), (942, y + 176), (958, y + 176)], fill=colors["line"])
    draw.text((70, 990), "data/ (persisted memory, vectors, logs, wake-word artifacts) and models/ back every layer above without owning behavior.", fill=colors["ink"], font=font)
    draw.text((70, 1035), "Cross-reference: audio/ = Sec 4.4, nlp/ = Sec 4.6, core/ = Sec 4.2, os_control/tools/llm = Sec 4.7-4.9, ui/desktop = Sec 4.11.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_runtime_pipeline_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_3_2_1_runtime_pipeline.png"
    width, height = 1900, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 46)
        font = ImageFont.truetype("arial.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()
    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((60, 30), "Figure 3.2 - Runtime pipeline", fill=colors["ink"], font=title_font)
    draw.text((60, 82), "Single turn, per orchestrator.run() - top row then bottom row, left to right.", fill=colors["muted"], font=small_font)
    stages = [
        ("wake_word.py", "EMA+peak gate", colors["blue"]),
        ("mic.py", "VAD capture", colors["blue"]),
        ("stt.py", "lang-locked text", colors["purple"]),
        ("command_router.py", "routing cascade", colors["purple"]),
        ("route_verifier.py", "execute/clarify/confirm/llm", colors["amber"]),
        ("dispatch", "handlers / os_control / llm", colors["amber"]),
        ("response_shaper.py", "voice-safe text", colors["green"]),
        ("tts.py", "sentence-streamed audio", colors["green"]),
    ]
    box_w, box_h, gap_x, gap_y = 440, 150, 30, 40
    cols = 4
    start_x, start_y = 55, 170
    for i, (title, subtitle, color) in enumerate(stages):
        row, col = divmod(i, cols)
        x = start_x + col * (box_w + gap_x)
        y = start_y + row * (box_h + gap_y)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=14, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 10, y + box_h), fill=color)
        draw.text((x + 24, y + 30), title, fill=colors["ink"], font=font)
        draw.text((x + 24, y + 84), subtitle, fill=colors["muted"], font=small_font)
        if col < cols - 1:
            ax = x + box_w
            draw.line((ax, y + box_h // 2, ax + gap_x, y + box_h // 2), fill=colors["line"], width=4)
            draw.polygon([(ax + gap_x, y + box_h // 2), (ax + gap_x - 10, y + box_h // 2 - 6), (ax + gap_x - 10, y + box_h // 2 + 6)], fill=colors["line"])
        elif row == 0:
            wrap_y = y + box_h
            draw.line((x + box_w, y + box_h // 2, x + box_w + gap_x // 2, y + box_h // 2), fill=colors["line"], width=4)
            draw.line((x + box_w + gap_x // 2, y + box_h // 2, x + box_w + gap_x // 2, wrap_y + gap_y // 2), fill=colors["line"], width=4)
            draw.line((x + box_w + gap_x // 2, wrap_y + gap_y // 2, start_x, wrap_y + gap_y // 2), fill=colors["line"], width=4)
            draw.line((start_x, wrap_y + gap_y // 2, start_x, wrap_y + gap_y), fill=colors["line"], width=4)
            draw.polygon(
                [(start_x, wrap_y + gap_y), (start_x - 8, wrap_y + gap_y - 10), (start_x + 8, wrap_y + gap_y - 10)],
                fill=colors["line"],
            )
    draw.text((60, 600), "Memory and metrics read/write around every stage (Sec 4.2/4.3); a wake-word interrupt during dispatch/response/TTS cancels that stage (runtime_coordinator.py).", fill=colors["muted"], font=small_font)
    draw.text((60, 640), "Latency tiering: parser/code-switch/semantic tiers resolve in low single-digit milliseconds; only unresolved turns reach the LLM tier.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_routing_cascade_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_3_3_1_routing_cascade.png"
    width, height = 1200, 1360
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 46)
        font = ImageFont.truetype("arial.ttf", 25)
        small_font = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()
    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((50, 35), "Figure 3.3 - Routing cascade", fill=colors["ink"], font=title_font)
    draw.text((50, 92), "Earliest tier that resolves the command wins; each tier hands off on a miss.", fill=colors["muted"], font=small_font)
    tiers = [
        ("1. Parser fast-path", "regex/keyword, ~0 ms", colors["blue"]),
        ("2. Code-switch router", "verb+entity map, ~2 ms", colors["blue"]),
        ("3. Semantic router", "MiniLM top-k + margin, ~10 ms", colors["purple"]),
        ("4. Keyword/fuzzy NLP", "rapidfuzz over noisy STT", colors["purple"]),
        ("5. Tool-calling LLM", "Ollama/Claude structured tools", colors["amber"]),
        ("6. Structured LLM NLU", "schema JSON, verifier-gated, opt-in", colors["amber"]),
        ("7. General LLM chat", "no command resolved", colors["grey"]),
    ]
    x = 90
    y = 160
    box_w, box_h, gap = 1020, 130, 26
    for i, (title, subtitle, color) in enumerate(tiers):
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + box_h), fill=color)
        draw.text((x + 28, y + 24), title, fill=colors["ink"], font=font)
        draw.text((x + 28, y + 70), subtitle, fill=colors["muted"], font=small_font)
        if i < len(tiers) - 1:
            draw.line((x + box_w // 2, y + box_h, x + box_w // 2, y + box_h + gap), fill=colors["line"], width=4)
            draw.polygon(
                [(x + box_w // 2, y + box_h + gap), (x + box_w // 2 - 8, y + box_h + gap - 10), (x + box_w // 2 + 8, y + box_h + gap - 10)],
                fill=colors["line"],
            )
        y += box_h + gap
    draw.text((90, y + 30), "route_verifier.py gates every tier's candidate before dispatch: schema/slots, question guard, risk tier, and policy.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_memory_scheme_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_3_4_1_memory_scheme.png"
    width, height = 1000, 1000
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 44)
        font = ImageFont.truetype("arial.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 19)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()
    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((45, 30), "Figure 3.4 - Memory scheme", fill=colors["ink"], font=title_font)
    draw.text((45, 82), "Fast RAM context vs richer LLM context, both backed by data/memory/.", fill=colors["muted"], font=small_font)
    nodes = [
        ("session_memory.py", "turns, slots, preferences API", colors["blue"], 60, 190),
        ("memory_manager.py", "fast vs LLM context split", colors["purple"], 60, 380),
        ("Fast context", "RAM slots, no I/O on hot path", colors["green"], 520, 300),
        ("LLM context", "recent turns + bounded vector recall", colors["amber"], 520, 460),
        ("memory_store.py", "SQLite + vector adapters", colors["grey"], 60, 570),
        ("data/memory/*.db", "persisted turns/slots (Sec 4.3)", colors["green"], 520, 620),
        ("data/vectors/ (ChromaDB)", "async semantic recall", colors["amber"], 520, 760),
    ]
    for title, subtitle, color, x, y in nodes:
        draw.rounded_rectangle((x, y, x + 420, y + 130), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((x, y, x + 12, y + 130), fill=color)
        draw.text((x + 26, y + 26), title, fill=colors["ink"], font=font)
        draw.text((x + 26, y + 72), subtitle, fill=colors["muted"], font=small_font)
    for start, end in [((480, 250), (520, 340)), ((480, 445), (520, 500)), ((480, 635), (520, 660)), ((480, 640), (520, 800)), ((270, 320), (270, 380))]:
        draw.line((*start, *end), fill=colors["line"], width=4)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 10, ey - 6), (ex - 10, ey + 6)], fill=colors["line"])
    draw.text((45, 920), "LLM context write path is asynchronous so vector recall never blocks the fast conversational path.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_hardware_tier_diagram() -> Path:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "figure_3_5_1_hardware_tier_selection.png"
    width, height = 1150, 1030
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 42)
        font = ImageFont.truetype("arial.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 19)
    except OSError:
        title_font = font = small_font = ImageFont.load_default()
    colors = {
        "ink": "#1a2230", "muted": "#5b6675", "line": "#dde3ec",
        "soft": "#f6f8fb", "blue": "#1f6feb", "purple": "#7c3aed",
        "amber": "#b26a00", "green": "#188a42", "grey": "#68707c",
    }
    draw.text((40, 30), "Figure 3.5 - Hardware-tier selection", fill=colors["ink"], font=title_font)
    draw.text((40, 82), "core/hardware_detect.py picks the largest Qwen3 model that fits detected RAM/GPU.", fill=colors["muted"], font=small_font)
    rows = [
        ("16 GB+, GPU", "high -> qwen3:8b", "num_ctx 8192", colors["green"]),
        ("12 GB+, any", "medium -> qwen3:4b", "num_ctx 4096", colors["blue"]),
        ("8 GB, GPU", "medium -> qwen3:4b", "num_ctx 4096", colors["blue"]),
        ("8 GB, no GPU", "low -> qwen3:1.7b", "num_ctx 2048", colors["amber"]),
        ("< 8 GB, any", "minimal -> qwen3:0.6b", "num_ctx 1024", colors["grey"]),
    ]
    y = 170
    for label, tier, ctx, color in rows:
        draw.rounded_rectangle((40, y, 1110, y + 130), radius=16, fill=colors["soft"], outline=colors["line"], width=3)
        draw.rectangle((40, y, 52, y + 130), fill=color)
        draw.text((75, y + 20), label, fill=colors["ink"], font=font)
        draw.text((75, y + 58), tier, fill=colors["ink"], font=font)
        draw.text((75, y + 96), ctx, fill=colors["muted"], font=small_font)
        y += 150
    draw.text((40, 930), "Missing models are auto-pulled via Ollama on first run.", fill=colors["muted"], font=small_font)
    draw.text((40, 960), "Override with JARVIS_LLM_MODEL or disable via JARVIS_LLM_AUTO_SELECT.", fill=colors["muted"], font=small_font)
    image.save(path, quality=95)
    return path


def add_phase12_c1_content(document: Document) -> None:
    document.add_heading("1.1 Background", level=2)
    document.add_paragraph(
        "Jarvis is a local-first Windows voice assistant that pairs a wake-word-triggered voice pipeline with an "
        "intent-routing cascade and a locally hosted LLM (Ollama, Qwen3 family), falling back to general conversation "
        "only when no faster tier resolves the command. It targets any Windows 10/11 PC with 8GB+ RAM and no GPU "
        "requirement, and treats Egyptian Arabic and English as equally first-class from wake word through STT, "
        "routing, and TTS, including code-switched utterances within a single sentence."
    )
    diagram = add_conceptual_overview_diagram()
    add_diagram_image(document, diagram, "Figure 1.1: Conceptual turn cycle - wake, record, STT, understand, act/answer, speak.")

    document.add_heading("1.2 Problem Statement", level=2)
    document.add_paragraph(
        "Commercial voice assistants are cloud-dependent, English-first, and treat code-switched Arabic/English speech "
        "as an edge case rather than the normal case for a bilingual Egyptian user. They also generally cannot verify "
        "that a spoken command actually changed the state of the operating system before reporting success, and they "
        "offer no local-only operating mode that keeps conversation and OS-control data on the user's own machine. "
        "Jarvis addresses this by running the full pipeline locally (STT, LLM, TTS all have local fallbacks), verifying "
        "OS side effects after execution rather than assuming API calls succeeded, and routing mixed-language input "
        "through a cascade tuned for Egyptian Arabic/English code-switching rather than a single-language grammar."
    )

    document.add_heading("1.3 Objectives", level=2)
    add_styled_table(
        document,
        ["ID", "Objective", "How it is measured in this book"],
        [
            ("O1", "Accurate bilingual understanding of Egyptian Arabic and English, including code-switching.", "Sec 4.6 nlp/ margin/fuzzy behavior tables; Sec 4.12 code-switch and semantic-margin test suites."),
            ("O2", "Low-latency response for the common case, reserving the LLM for what earlier tiers cannot resolve.", "Sec 3.4 routing cascade tiers; Sec 4.2 command_router cascade; Sec 4.12 latency-oriented tests."),
            ("O3", "Reliable hands-free activation without false wakes or missed wakes.", "Sec 4.4 wake decision algorithm; Sec 4.5 wake-word dataset class balance and feature evidence."),
            ("O4", "Verified control of the Windows OS rather than fire-and-forget API calls.", "Sec 4.8 os_control verify-after-execute contract and behavior table."),
            ("O5", "Safety-by-risk-tier for destructive or sensitive actions.", "Sec 4.8 policy/risk_policy, confirmation/second_factor PIN flow; Sec 4.12 routing-safety tests."),
            ("O6", "Continuity and personalization across sessions via durable memory.", "Sec 4.2 session/memory manager; Sec 4.3 persisted data/memory/ artifacts and schema."),
        ],
    )
    add_caption(document, "Table 1.1: Project objectives O1-O6 and where this book substantiates each one.")

    document.add_heading("1.4 Scope", level=2)
    document.add_paragraph(
        "In scope: wake-word detection and adaptive retraining, hybrid cloud/local STT and TTS, the multi-tier intent "
        "routing cascade, verified Windows OS control, local LLM chat/tool-calling, durable SQLite/vector memory, an "
        "optional desktop UI bridge, and the automated test/eval suite documented in Chapter 4. Out of scope: a mobile "
        "or cross-platform client, cloud-hosted multi-user deployment, and any UI surface required for the engine to "
        "function, since the engine is designed to run headless."
    )

    document.add_heading("1.5 Motivation", level=2)
    document.add_paragraph(
        "The motivating use case is a single Windows user who speaks a natural mix of Egyptian Arabic and English in "
        "the same sentence - e.g. asking to \"open Chrome\" or \"شغّل Spotify\" - and expects the assistant to "
        "understand both without switching modes. Sections 4.6 and 4.4 show this is treated as the default input shape "
        "throughout the pipeline (STT language lock plus code-switch acceptance, a dedicated code-switch router tier, "
        "and bilingual TTS voice profiles) rather than bolted on as a translation layer."
    )

    document.add_heading("1.6 Report Organization", level=2)
    add_styled_table(
        document,
        ["Chapter", "Content"],
        [
            ("1", "Introduction: background, problem statement, objectives, scope, motivation."),
            ("2", "Literature Review: prior art in commercial assistants, local LLM runtimes, STT, wake word, VAD, TTS, semantic routing and memory."),
            ("3", "Methodology: design principles and the five whole-system diagrams (architecture, pipeline, routing, memory, hardware tiers)."),
            ("4", "Implementation: one section per top-level folder (4.1 root through 4.14 models/), built folder-by-folder from the real source and data."),
            ("5", "Conclusion & Future Work: objectives achieved, limitations, and next steps."),
        ],
    )
    add_caption(document, "Table 1.2: Reading guide for this document.")


def add_phase12_c2_content(document: Document) -> None:
    document.add_heading("2.1 Commercial Voice Assistants", level=2)
    document.add_paragraph(
        "Cloud assistants such as Siri, Google Assistant, and Alexa are optimized for broad single-language coverage "
        "and cloud-side processing; mixed-language utterances are typically handled by locale switching rather than "
        "per-utterance code-switch tolerance. Jarvis instead runs primarily on-device and treats Egyptian Arabic/English "
        "code-switching as a first-class input shape (Sec 4.6), and does not require a persistent cloud connection for "
        "its core loop (Sec 4.4, 4.7)."
    )

    document.add_heading("2.2 Local LLM Runtimes", level=2)
    document.add_paragraph(
        "Ollama serves as the local model runtime, hosting Qwen3-family models chosen automatically by "
        "core/hardware_detect.py (Sec 3.6, Sec 4.7). Running the assistant tier locally avoids per-request cloud "
        "latency and keeps conversation content on the user's machine; the optional Claude API path in "
        "llm/tool_caller.py is documented as an alternate backend rather than the primary one (Sec 4.7)."
    )

    document.add_heading("2.3 Speech-to-Text: Faster-Whisper and ElevenLabs", level=2)
    document.add_paragraph(
        "Faster-Whisper provides a local, hardware-sized STT fallback (model tier chosen by "
        "recommend_whisper_runtime(), Sec 3.6), while ElevenLabs provides higher cloud-quality transcription when an "
        "API key is configured. The key finding carried through Sec 4.4 is the per-utterance language lock: rather than "
        "decoding with language=None and letting the model guess, Jarvis fixes the primary language per utterance while "
        "still accepting code-switched words in the other script, and applies a confidence floor plus a hallucination "
        "guard (log-probability and compression ratio) to reject noise instead of passing along fabricated text."
    )

    document.add_heading("2.4 Wake-Word Detection and Dataset Practice", level=2)
    document.add_paragraph(
        "openWakeWord provides the ONNX runtime convention that audio/wake_word.py builds on: a single unified model "
        "detecting both \"Jarvis\" and \"جارفيس\" rather than per-language variants, decided via an "
        "EMA-smoothed score plus a raw peak threshold and confirm-frame debounce (Sec 4.4). Sec 4.5 documents the "
        "supporting dataset practice: class-balanced positive/negative WAV splits and precomputed fixed-shape feature "
        "tensors (41x96 float32), with negative examples deliberately outnumbering positive ones to bias toward "
        "false-wake resistance."
    )

    document.add_heading("2.5 Voice Activity Detection: Silero", level=2)
    document.add_paragraph(
        "Silero VAD (ONNX) provides the primary speech/non-speech gate for recording, with an energy-based fallback "
        "when the ONNX model is unavailable (Sec 4.4). This two-tier VAD design mirrors the wake-word module's own "
        "primary/fallback pattern, keeping the pipeline usable in degraded environments."
    )

    document.add_heading("2.6 Text-to-Speech: edge-tts and ElevenLabs", level=2)
    document.add_paragraph(
        "edge-tts supplies a free, no-API-key baseline voice for both languages (ar-EG-SalmaNeural, en-US-AriaNeural); "
        "ElevenLabs is used when configured for higher-quality synthesis. audio/tts.py streams sentence-by-sentence so "
        "playback starts before the full response is generated, and a deterministic (non-LLM) voice normalizer "
        "converts numbers, dates, and place names to spoken form before either backend receives the text (Sec 4.4)."
    )

    document.add_heading("2.7 Semantic Routing and Memory", level=2)
    document.add_paragraph(
        "sentence-transformers (multilingual MiniLM) backs the semantic router's top-k plus margin acceptance rule "
        "(Sec 4.6, Sec 3.4), letting paraphrased commands resolve without an LLM call. ChromaDB backs the optional "
        "vector-memory recall path documented in Sec 4.3 and Sec 3.5, providing bounded semantic recall for the LLM "
        "context without blocking the fast conversational path."
    )

    document.add_heading("2.8 Positioning", level=2)
    document.add_paragraph(
        "Jarvis's contribution is not any single component in isolation - openWakeWord, Faster-Whisper, Ollama, and "
        "sentence-transformers are all existing open-source building blocks - but their integration into a cascade "
        "that is bilingual by default, verifies OS side effects rather than assuming them, and degrades gracefully "
        "component-by-component (Sec 4.4, 4.6, 4.8) instead of failing the whole turn when one optional dependency is "
        "missing."
    )


def add_phase12_c3_content(document: Document) -> None:
    document.add_heading("3.1 Design Principles", level=2)
    add_styled_table(
        document,
        ["Principle", "What it means in this codebase"],
        [
            ("Latency tiering", "Cheaper, earlier routing tiers (parser, code-switch, semantic) are tried before the LLM tier, so most turns never pay LLM latency (Sec 3.4, Sec 4.2)."),
            ("One LLM call per turn", "llm/prompt_builder.py assembles persona, language, memory, and KB context into a single prompt package rather than chaining multiple model calls (Sec 4.7)."),
            ("Verify before claiming", "os_control/ adapters read back real OS state after acting and only report success if it verifiably changed (Sec 4.8)."),
            ("Graceful degradation", "Optional dependencies (Silero, ElevenLabs, sentence-transformers, WinRT, pywin32) each have a documented fallback path rather than a hard failure (Sec 4.4, 4.6, 4.8)."),
            ("Safety by risk tier", "Destructive or sensitive actions route through policy/risk_policy and a spoken-PIN second factor before execution (Sec 4.8)."),
        ],
    )
    add_caption(document, "Table 3.1: Design principles evidenced across the implementation chapters.")

    document.add_heading("3.2 Architecture", level=2)
    document.add_paragraph(
        "Jarvis is organized into five horizontal layers plus a persisted data layer that backs all of them: an I/O "
        "layer (audio/) at the boundary with the user, an understanding layer (nlp/) that turns transcripts into "
        "candidate intents, a control plane (core/) that owns routing, verification, memory, and persona, an "
        "action/answer layer (os_control/, tools/, llm/) that performs verified side effects or generates language, and "
        "an optional presentation layer (ui/, desktop/) that never gates core functionality."
    )
    diagram = add_layered_architecture_diagram()
    add_diagram_image(document, diagram, "Figure 3.1: Layered architecture from I/O through the optional presentation layer.")

    document.add_heading("3.3 Runtime Pipeline", level=2)
    document.add_paragraph(
        "Within orchestrator.run(), a single turn flows wake -> record -> STT -> route -> verify -> dispatch -> shape "
        "-> speak. Each stage is documented in its owning folder chapter; this figure shows how those chapters compose "
        "into one pass, with the wake-word interrupt gate (runtime_coordinator.py) able to cancel dispatch, response "
        "shaping, or TTS mid-stream (Sec 4.2, 4.4, 4.7)."
    )
    diagram = add_runtime_pipeline_diagram()
    add_diagram_image(document, diagram, "Figure 3.2: Runtime pipeline for a single turn, stage by stage.")

    document.add_heading("3.4 Routing Cascade", level=2)
    document.add_paragraph(
        "command_router.py tries seven tiers in order and dispatches on the first one that resolves the command: a "
        "regex/keyword parser fast-path, a code-switch shortcut router, a semantic router, keyword/fuzzy matching, a "
        "tool-calling LLM tier, an opt-in structured-LLM NLU fallback gated behind route_verifier.py, and finally "
        "general LLM chat for anything that is not a command. The semantic tier's margin-acceptance rule requires the "
        "best candidate to clear an absolute confidence floor (best >= sigma) and to lead the runner-up by a minimum "
        "gap (best - second >= delta); failing either check defers to a later tier instead of guessing on a near-tie "
        "(Sec 4.6)."
    )
    diagram = add_routing_cascade_diagram()
    add_diagram_image(document, diagram, "Figure 3.3: Routing cascade tiers, earliest resolving tier wins.")

    document.add_heading("3.5 Memory Scheme", level=2)
    document.add_paragraph(
        "session_memory.py exposes two context shapes to the rest of the system: a fast, RAM-only context (recent "
        "slots and references, no I/O on the hot path) and a richer LLM context (recent turns plus bounded vector "
        "recall) assembled by memory_manager.py. Both are backed by the persisted store documented in Sec 4.3: a "
        "SQLite database (turns and slots tables, WAL journaling) is the primary store, while data/vectors/ (ChromaDB "
        "with all-MiniLM-L6-v2 embeddings) supplies semantic recall asynchronously so it never blocks the fast path."
    )
    diagram = add_memory_scheme_diagram()
    add_diagram_image(document, diagram, "Figure 3.4: Memory scheme - fast RAM context and LLM context over the same persisted store.")

    document.add_heading("3.6 Models", level=2)
    add_styled_table(
        document,
        ["Model", "Role", "Source"],
        [
            ("Qwen3 family (0.6b/1.7b/4b/8b)", "Local chat/tool-calling LLM tier, selected by hardware tier.", "Ollama, auto-pulled on first run."),
            ("Unified wake-word ONNX (jarvis_unified.onnx)", "Bilingual EN/EGY wake detection, one model for both trigger phrases.", "Trained via scripts/train_arabic_wake_model.py over data in Sec 4.5; loaded in Sec 4.4."),
            ("Faster-Whisper (base/small/medium)", "Local STT fallback, sized by VRAM/RAM via recommend_whisper_runtime().", "faster-whisper package."),
            ("Silero VAD (ONNX)", "Primary speech/non-speech gate, with energy-based fallback.", "onnxruntime model asset."),
            ("all-MiniLM-L6-v2 / multilingual MiniLM", "Sentence embeddings for semantic routing and vector memory recall.", "sentence-transformers."),
            ("edge-tts / ElevenLabs voices", "Bilingual TTS voice profiles (ar-EG-SalmaNeural, en-US-AriaNeural, or cloud voices).", "edge-tts package / ElevenLabs API."),
        ],
    )
    add_caption(document, "Table 3.2: Models used across the pipeline and where each is loaded.")
    diagram = add_hardware_tier_diagram()
    add_diagram_image(document, diagram, "Figure 3.5: Hardware-tier selection for the local LLM model.")
    add_styled_table(
        document,
        ["RAM", "GPU", "Tier", "Model", "num_ctx", "lightweight_ctx"],
        [
            ("16 GB+", "Yes", "high", "qwen3:8b", "8192", "4096"),
            ("12 GB+", "Any", "medium", "qwen3:4b", "4096", "2048"),
            ("8 GB", "Yes", "medium", "qwen3:4b", "4096", "2048"),
            ("8 GB", "No", "low", "qwen3:1.7b", "2048", "1024"),
            ("< 8 GB", "Any", "minimal", "qwen3:0.6b", "1024", "512"),
        ],
    )
    add_caption(document, "Table 3.3: core/hardware_detect.py model tiers, copied from _MODEL_TIERS.")

    document.add_heading("3.7 Dependencies", level=2)
    add_styled_table(
        document,
        ["Tier", "Representative packages", "Role"],
        [
            ("Tier 1: Voice pipeline (core)", "python-dotenv, numpy, sounddevice, onnxruntime, faster-whisper, openwakeword, edge-tts, soundfile", "Wake word, capture, STT, TTS - required to run at all."),
            ("Tier 2: Orchestration (core)", "httpx, psutil, rapidfuzz", "Ollama/HTTP client, hardware detection, fuzzy intent matching."),
            ("Tier 3: Optional LLM backends", "anthropic", "Optional Claude API backend; Ollama remains primary."),
            ("Tier 4: Intelligent features", "chromadb, sentence-transformers, duckduckgo-search, ddgs", "Semantic routing, vector memory, live web search; degrade gracefully if absent."),
            ("Tier 5: OS integration (Windows)", "comtypes, pycaw, wmi, winsdk, screen-brightness-control, pyperclip, watchdog, pywin32, elevenlabs, pystray, Pillow", "Volume/brightness/radio control, clipboard, Outlook, tray icon."),
            ("Tier 6: Desktop UI bridge (optional)", "fastapi, uvicorn[standard]", "WebSocket bridge to the desktop app; engine runs without it."),
            ("Training only (requirements-training.txt)", "scipy, torch", "Wake-word model retraining (scripts/train_arabic_wake_model.py), not needed to run Jarvis."),
        ],
    )
    add_caption(document, "Table 3.4: Dependency tiers, copied from requirements.txt/requirements-training.txt.")

    document.add_heading("3.8 Key Algorithms Overview", level=2)
    add_styled_table(
        document,
        ["Algorithm", "Where documented", "One-line summary"],
        [
            ("Wake decision", "Sec 4.4", "EMA-smoothed score plus raw peak threshold, confirm-frame debounce, and cooldown before triggering."),
            ("Per-utterance language lock", "Sec 4.4", "STT fixes the primary language per utterance instead of decoding with language=None, while still accepting code-switched words."),
            ("Semantic margin routing", "Sec 4.6, Sec 3.4", "Accept only if best score clears an absolute floor and leads the runner-up by a minimum margin; otherwise defer to the next tier."),
            ("Route verifier decision", "Sec 4.2", "Converts a routed candidate into execute / clarify / confirm / llm based on schema completeness, question-shape guard, and risk tier."),
            ("Verify-after-execute", "Sec 4.8", "Every OS-control adapter re-reads real state after acting and only reports success if the state verifiably changed."),
            ("Wake-word feature/training/eval pipeline", "Sec 4.5", "Fixed-shape (41x96) feature tensors from class-balanced WAV splits feed model training; evaluation happens offline before a model ships."),
        ],
    )
    add_caption(document, "Table 3.5: Key algorithms and their detailed treatment elsewhere in this book.")


def fill_phase12(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    insert_generated_content(doc_path, "c1", add_phase12_c1_content)
    insert_generated_content(doc_path, "c2", add_phase12_c2_content)
    insert_generated_content(doc_path, "c3", add_phase12_c3_content)
    return doc_path


def remove_orphaned_c4_utils_section(doc_path: Path = DOC_PATH) -> None:
    """Delete the leftover Phase-9-scaffold heading/placeholder.

    utils/ content was written into the real "4.10 utils/" section under the
    c4-tools placeholder; the separate c4-utils placeholder from the Phase 0
    skeleton was never filled and its heading is orphaned. Remove both the
    stray heading and the placeholder paragraph so no empty section remains.
    """
    document = Document(str(doc_path))
    paras = document.paragraphs
    target_index = None
    for i, p in enumerate(paras):
        if p.text.strip() == "4.10 utils/ cross-reference placeholder":
            target_index = i
            break
    if target_index is None:
        return
    to_remove = [paras[target_index]]
    if target_index + 1 < len(paras) and "[[FILL:c4-utils]]" in paras[target_index + 1].text:
        to_remove.append(paras[target_index + 1])
    for paragraph in to_remove:
        delete_paragraph(paragraph)
    document.save(str(doc_path))


def add_phase13_title_content(document: Document) -> None:
    document.add_paragraph("Jarvis", style="Title")
    document.add_paragraph("A Local-First Bilingual (English / Egyptian Arabic) Windows Voice Assistant")
    document.add_paragraph("Project Documentation Book - prepared for the Final Discussion")
    document.add_paragraph(
        "Wake word -> streaming capture + STT -> intent routing cascade -> verified action / local LLM -> streamed TTS"
    )
    document.add_paragraph("Generated from the Jarvis repository source and runtime data, folder-by-folder.")
    document.add_paragraph("")
    document.add_paragraph("Prepared by: Abdelrhman Yousef Mahrous")
    document.add_paragraph("Supervised by: Dr. Ibrahim Mubarak")
    document.add_paragraph("2026")


def add_phase13_ack_content(document: Document) -> None:
    document.add_paragraph(
        "This project builds on a stack of open-source and hosted components documented throughout Chapter 4: "
        "openWakeWord and Silero for on-device audio detection, Faster-Whisper and ElevenLabs for bilingual speech "
        "recognition, Ollama and the Qwen3 model family for local language understanding and generation, edge-tts "
        "for free-tier speech synthesis, sentence-transformers and ChromaDB for semantic routing and memory recall, "
        "and Open-Meteo/DuckDuckGo for live data. Their availability as free or self-hostable tools is what makes a "
        "fully local-first bilingual assistant practical on ordinary consumer hardware."
    )
    document.add_paragraph(
        "Equal acknowledgment goes to the Egyptian Arabic / English bilingual user experience that motivated this "
        "project: treating code-switched speech as the normal case, not an edge case, shaped nearly every design "
        "decision documented in this book, from the unified wake-word model in Sec 4.4 to the code-switch router in "
        "Sec 4.6."
    )


def add_phase13_abstract_content(document: Document) -> None:
    document.add_paragraph(
        "Jarvis is a local-first Windows voice assistant built for a bilingual Egyptian Arabic / English user. A "
        "single unified wake-word model detects both \"Jarvis\" and \"جارفيس\"; a hybrid cloud/local speech-to-text "
        "path locks the primary language per utterance while still accepting code-switched words; a seven-tier "
        "routing cascade resolves most commands in low single-digit milliseconds and reserves a locally hosted Qwen3 "
        "language model (selected automatically for the detected hardware) for anything the faster tiers cannot "
        "resolve; every operating-system side effect is verified by re-reading real OS state rather than assumed from "
        "an API call's return value; and durable SQLite plus vector memory give the assistant continuity across "
        "sessions. This book documents that system for a final defense: Chapter 1 states the problem and objectives, "
        "Chapter 2 reviews the prior art each component builds on, Chapter 3 lays out the whole-system methodology and "
        "diagrams, Chapter 4 documents every top-level folder of the implementation against the real source and data "
        "on disk, and Chapter 5 closes with an honest accounting of what is achieved and what remains as future work."
    )
    document.add_heading("Reading Guide", level=3)
    document.add_paragraph(
        "Readers preparing for the defense should read Chapters 1 and 3 first for the whole-system picture, then use "
        "Chapter 4 as a reference: each subsection (4.1 through 4.14) is self-contained, maps to exactly one "
        "repository folder, and can be read in isolation without the others."
    )


def add_phase13_abbr_content(document: Document) -> None:
    add_styled_table(
        document,
        ["Abbreviation", "Meaning"],
        [
            ("ASR / STT", "Automatic Speech Recognition / Speech-to-Text."),
            ("TTS", "Text-to-Speech."),
            ("VAD", "Voice Activity Detection."),
            ("NLU", "Natural Language Understanding."),
            ("LLM", "Large Language Model."),
            ("ONNX", "Open Neural Network Exchange, the runtime model format used for wake-word and VAD models."),
            ("WAL", "Write-Ahead Logging, the SQLite journal mode used by the memory store."),
            ("KB", "Knowledge Base (offline retrieval index)."),
            ("PIN", "Personal Identification Number, used as the spoken second factor for sensitive actions."),
            ("EGY", "Egyptian Arabic."),
            ("WinRT", "Windows Runtime API, used for no-admin Wi-Fi/Bluetooth/Airplane-mode control."),
            ("FAR / FRR", "False-Accept Rate / False-Reject Rate, wake-word evaluation metrics."),
            ("p50 / p95", "50th / 95th percentile latency."),
            ("EMA", "Exponential Moving Average, used to smooth the wake-word detection score."),
            ("RAM", "Random-Access Memory, used for hardware-tier model selection."),
            ("GPU", "Graphics Processing Unit."),
        ],
    )


def _collect_headings_for_contents(document: Document) -> list[tuple[int, str]]:
    """Collect Heading 1/2 text for the static contents list.

    Heading 3 is intentionally excluded: it is the per-folder chapter
    template's subsection name (Purpose, File Inventory, Algorithms, ...)
    repeated across all eleven Chapter 4 subsections, so listing it would
    add ~90 near-duplicate, non-navigational entries instead of a usable TOC.
    """
    entries = []
    for paragraph in document.paragraphs:
        style = paragraph.style.name
        text = paragraph.text.strip()
        if not text:
            continue
        if style == "Heading 2" and text == "Document Build Note":
            continue
        if style == "Heading 1":
            entries.append((0, text))
        elif style == "Heading 2" and text not in FRONT_MATTER_TITLES:
            entries.append((1, text))
    return entries


def add_contents_static_list(document: Document) -> None:
    document.add_paragraph(
        "Static contents list (regenerate the live Table of Contents field in Word/LibreOffice with References -> "
        "Update Table for page numbers):",
        style="Caption",
    )
    source = Document(str(DOC_PATH))
    for depth, text in _collect_headings_for_contents(source):
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
        if depth == 0:
            paragraph.runs[0].bold = True


def add_figures_static_list(document: Document) -> None:
    document.add_paragraph(
        "Static List of Figures (regenerate the live field in Word/LibreOffice for page numbers):",
        style="Caption",
    )
    source = Document(str(DOC_PATH))
    count = 0
    for paragraph in source.paragraphs:
        if paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Figure"):
            document.add_paragraph(paragraph.text.strip())
            count += 1
    if count == 0:
        document.add_paragraph("No figures were found in the document at generation time.")


def add_tables_static_list(document: Document) -> None:
    document.add_paragraph(
        "Static List of Tables (regenerate the live field in Word/LibreOffice for page numbers):",
        style="Caption",
    )
    source = Document(str(DOC_PATH))
    count = 0
    for paragraph in source.paragraphs:
        if paragraph.style.name == "Caption" and paragraph.text.strip().startswith("Table"):
            document.add_paragraph(paragraph.text.strip())
            count += 1
    if count == 0:
        document.add_paragraph("No tables were found in the document at generation time.")


def add_phase13_c5_content(document: Document) -> None:
    document.add_heading("5.1 Objectives Achieved", level=2)
    add_styled_table(
        document,
        ["ID", "Objective", "Evidence in this book"],
        [
            ("O1", "Bilingual EGY/English understanding, including code-switching.", "Sec 4.6 code-switch router, semantic margin rule, and fuzzy/keyword fallback; Sec 4.4 STT language lock with code-switch acceptance; Sec 4.12 code-switch and semantic-margin test suites."),
            ("O2", "Low-latency response, reserving the LLM tier for what earlier tiers cannot resolve.", "Sec 3.4 seven-tier cascade; Sec 4.2 command_router implementation; parser/code-switch/semantic tiers resolve without any LLM call."),
            ("O3", "Reliable hands-free wake activation.", "Sec 4.4 EMA/peak/confirm-frame wake decision; Sec 4.5 class-balanced dataset (negatives outnumber positives) biasing toward false-wake resistance."),
            ("O4", "Verified control of the Windows OS.", "Sec 4.8 verify-after-execute contract implemented across native_ops, radio_ops, and windows_toggles, each re-reading real state after acting."),
            ("O5", "Safety-by-risk-tier for destructive/sensitive actions.", "Sec 4.8 policy/risk_policy tiers and the spoken-PIN second-factor confirmation flow; Sec 4.12 routing-safety tests asserting zero unsafe auto-executions."),
            ("O6", "Continuity/personalization via durable memory.", "Sec 4.2 session/memory manager fast-vs-LLM context split; Sec 4.3 persisted SQLite turns/slots schema and vector-memory recall path."),
        ],
    )
    add_caption(document, "Table 5.1: Objectives O1-O6 against the evidence documented in Chapter 4.")
    document.add_paragraph(
        "Across the eleven implementation chapters (Sec 4.1-4.14), every objective stated in Chapter 1 is backed by "
        "real, inspected source code or on-disk data rather than by design intent alone: file inventories were "
        "cross-checked against the actual folders, and data-chapter claims (Sec 4.3, 4.5, 4.12, 4.14) were verified "
        "against live SQLite queries, NumPy tensor shapes, and JSONL fixture counts rather than estimated."
    )

    document.add_heading("5.2 Limitations", level=2)
    document.add_paragraph(
        "Several gaps are visible directly from the implementation chapters rather than inferred. The wake-word "
        "chapter (Sec 4.5) found no training-loop source code inside the deployed-relevant dataset folder itself, so "
        "the exact loss function, epoch count, and threshold-selection procedure are documented only as the training "
        "script's stated purpose and cross-referenced to scripts/train_arabic_wake_model.py rather than reproduced "
        "step-by-step; a smaller sibling Arabic-only wake dataset was also observed but is out of scope for the "
        "chapter, which explicitly targets the unified deployed model instead. The memory chapter (Sec 4.3) noted "
        "that some persisted-state claims (e.g. SQLite journal_mode) reflect a point-in-time snapshot and can drift as "
        "the store continues to be written during normal use. The desktop UI chapter (Sec 4.11) summarizes the "
        "Tauri+React toolchain rather than exhaustively documenting every component, per the folder-by-folder scope "
        "rule for that phase. Several optional dependencies (ElevenLabs, sentence-transformers, WinRT, pywin32) are "
        "documented with their fallback behavior but were not exercised in a live degraded-hardware test as part of "
        "this book."
    )

    document.add_heading("5.3 Future Work", level=2)
    add_styled_table(
        document,
        ["Area", "Possible next step"],
        [
            ("Wake-word evaluation", "Add a dedicated held-out eval harness with reported FAR/FRR inside the training folder itself, rather than relying on the deployed model's field behavior alone (extends Sec 4.5)."),
            ("Structured LLM NLU", "Promote the opt-in schema-constrained NLU fallback (Sec 4.7) from off-by-default to a measured, gradually-enabled path once more labeled cases are collected (Sec 4.12 fixtures)."),
            ("Memory consolidation", "Add an explicit long-term consolidation/summarization policy for the turns table beyond the currently-observed bounded recent-turn pattern (Sec 4.3)."),
            ("Desktop UI parity", "Extend the desktop bridge/tray coverage (Sec 4.11) so more engine event types have a first-class dashboard view."),
            ("Cross-platform scope", "Evaluate which os_control/ adapters would need Linux/macOS equivalents if the assistant were ever ported beyond Windows (Sec 4.8)."),
        ],
    )
    add_caption(document, "Table 5.2: Candidate future-work items grounded in gaps observed while documenting Chapter 4.")


def add_phase13_refs_content(document: Document) -> None:
    add_styled_table(
        document,
        ["Reference", "Role in Jarvis"],
        [
            ("openWakeWord", "ONNX wake-word detection convention consumed by audio/wake_word.py (Sec 4.4) and produced by the training pipeline (Sec 4.5)."),
            ("Silero VAD", "Primary voice-activity-detection backend in audio/vad.py (Sec 4.4)."),
            ("Faster-Whisper", "Local speech-to-text fallback in audio/stt.py, sized by core/hardware_detect.py (Sec 4.4, 3.6)."),
            ("ElevenLabs", "Optional cloud STT/TTS backend for higher-quality bilingual audio (Sec 4.4)."),
            ("Ollama", "Local LLM runtime serving the Qwen3 model family (Sec 4.7, 3.6)."),
            ("Qwen3 model family", "Local chat/tool-calling models selected by hardware tier (Sec 3.6, core/hardware_detect.py)."),
            ("edge-tts", "Free-tier text-to-speech backend, default voices ar-EG-SalmaNeural / en-US-AriaNeural (Sec 4.4)."),
            ("sentence-transformers", "Multilingual embeddings backing the semantic router and vector memory recall (Sec 4.6, 4.3, 2.7)."),
            ("ChromaDB", "Vector store for semantic memory recall under data/vectors/ (Sec 4.3, 3.5)."),
            ("Open-Meteo", "No-API-key weather data source used by tools/weather.py (Sec 4.9)."),
            ("DuckDuckGo (duckduckgo-search / ddgs)", "No-API-key web search source used by tools/web_search.py (Sec 4.9)."),
            ("Repository source and data", "All file inventories, code excerpts, schema samples, and behavior tables in Chapter 4 are drawn directly from the Jarvis repository source and data/ contents as read at documentation time."),
        ],
    )
    add_caption(document, "Table R.1: Consolidated references and their role in the system.")


def add_phase13_appendix_content(document: Document) -> None:
    document.add_heading("Appendix A: Source-Code Listing Reference", level=2)
    add_styled_table(
        document,
        ["Area", "Folder(s)", "Documented in"],
        [
            ("Root package", "repository root", "Sec 4.1"),
            ("Runtime control plane", "core/, core/handlers/", "Sec 4.2"),
            ("Persisted memory", "data/memory/", "Sec 4.3"),
            ("Audio I/O", "audio/", "Sec 4.4"),
            ("Wake-word data/training", "wake word data/jarvis_unified_training/", "Sec 4.5"),
            ("Understanding layer", "nlp/", "Sec 4.6"),
            ("Local LLM", "llm/, llm/prompts/", "Sec 4.7"),
            ("Windows OS control", "os_control/", "Sec 4.8"),
            ("Live-data tools", "tools/", "Sec 4.9"),
            ("Shared utilities", "utils/", "Sec 4.10"),
            ("UI bridge and desktop app", "ui/, desktop/", "Sec 4.11"),
            ("Tests, scripts, models", "tests/, scripts/, models/", "Sec 4.12-4.14"),
        ],
    )
    add_caption(document, "Table A.1: Where each area of the source tree is documented in this book.")

    document.add_heading("Appendix B: Configuration Reference", level=2)
    document.add_paragraph(
        "Every environment key referenced throughout Chapter 4 has a working default in .env.example at the "
        "repository root (Sec 4.1). Copy it to .env and edit only the keys relevant to your setup; representative "
        "keys are tabulated per-folder in each Configuration Surface subsection (e.g. Table 4.1.2, 4.4.2, 4.8.3)."
    )

    document.add_heading("Appendix C: Installation Guide", level=2)
    add_styled_table(
        document,
        ["Step", "Command / action"],
        [
            ("1. Install runtime dependencies", "python -m pip install -r requirements.txt"),
            ("2. (Optional) install training-only dependencies", "python -m pip install -r requirements-training.txt"),
            ("3. Copy the environment template", "copy .env.example .env"),
            ("4. Start the local LLM runtime", "ollama serve"),
            ("5. Run the assistant", "python main.py"),
            ("6. Verify the install", "python core/doctor.py"),
        ],
    )
    add_caption(document, "Table C.1: Installation steps, copied from README.md.")

    document.add_heading("Appendix D: User Manual", level=2)
    document.add_paragraph(
        "Say \"Jarvis\" or \"جارفيس\" to wake the assistant - both are detected by the same unified model (Sec 4.4). "
        "Speak your command or question in English, Egyptian Arabic, or a natural mix of both (Sec 4.6). Destructive "
        "or sensitive actions (permanent delete, rename/move, shutdown) will ask for a spoken PIN before executing "
        "(default 1234, configurable via JARVIS_SECOND_FACTOR_PIN, Sec 4.8) - never read anything back that could be "
        "overheard and replayed. If a control genuinely cannot be verified as changed, Jarvis reports that honestly "
        "instead of claiming success."
    )

    document.add_heading("Appendix E: Testing Scenarios", level=2)
    document.add_paragraph(
        "The automated suite documented in Sec 4.12 exercises: code-switch routing across apps/volume/folders/delete "
        "intents; LLM routing guards so advice questions do not misfire as commands; semantic near-tie deferral; "
        "fast-context and LLM-context memory behavior; freshness-gated app/file reference resolution; pending-task "
        "slot filling and TTL expiry; sentence-buffer boundaries for English and Arabic; TTS prosody normalization; "
        "and voice-normalizer output for weather, units, and URLs. The labeled NLU evaluation fixture (154 cases) "
        "backs the safety invariant that zero unsafe actions auto-execute and no question is misrouted as a command."
    )

    document.add_heading("Appendix F: Glossary", level=2)
    document.add_paragraph(
        "See the Abbreviations section in the front matter for acronym expansions. Domain-specific terms used "
        "throughout Chapter 4 (margin acceptance, verify-after-execute, risk tier, fast context vs. LLM context) are "
        "each defined at first use in their owning subsection and cross-referenced from Chapter 3."
    )


def fill_phase13(doc_path: Path = DOC_PATH) -> Path:
    if not doc_path.exists():
        raise FileNotFoundError(f"Missing {doc_path}; run --phase 0 first.")
    remove_orphaned_c4_utils_section(doc_path)
    insert_generated_content(doc_path, "title", add_phase13_title_content)
    insert_generated_content(doc_path, "ack", add_phase13_ack_content)
    insert_generated_content(doc_path, "abstract", add_phase13_abstract_content)
    insert_generated_content(doc_path, "abbr", add_phase13_abbr_content)
    insert_generated_content(doc_path, "c5", add_phase13_c5_content)
    insert_generated_content(doc_path, "refs", add_phase13_refs_content)
    insert_generated_content(doc_path, "appendix", add_phase13_appendix_content)
    # Contents/Figures/Tables are generated last so they reflect the fully assembled document.
    insert_generated_content(doc_path, "contents", add_contents_static_list)
    insert_generated_content(doc_path, "figures", add_figures_static_list)
    insert_generated_content(doc_path, "tables", add_tables_static_list)
    return doc_path


def add_build_note(document: Document) -> None:
    document.add_heading("Document Build Note", level=2)
    add_callout(
        document,
        "This book is built folder-by-folder. Each implementation phase reads one target area, "
        "fills one placeholder, verifies the result, and leaves all other sections untouched.",
    )
    add_styled_table(
        document,
        ["Convention", "Phase 0 rule"],
        [
            ("Figure numbering", "Use Figure 4.<folder>.<n> inside Chapter 4 sections."),
            ("Table numbering", "Use Table 4.<folder>.<n> inside Chapter 4 sections."),
            ("Folder template", "Purpose; file inventory; data-flow diagram; key modules; algorithms; configuration; behavior table; contribution."),
            ("Accuracy", "Claims must map to files read in that phase; uncertain details are marked as unverified rather than asserted."),
            ("Language scope", "Egyptian-Arabic and English behavior are both first-class where the code supports them."),
        ],
    )


def add_front_matter(document: Document, bookmark_start: int) -> int:
    document.add_paragraph("Jarvis Documentation Book", style="Title")
    document.add_paragraph("Phase 0 scaffold for a production-grade thesis/defense document.")
    add_build_note(document)
    document.add_page_break()

    bookmark_id = bookmark_start
    for placeholder_id, title in FRONT_MATTER:
        document.add_heading(title, level=1)
        if placeholder_id == "contents":
            paragraph = document.add_paragraph()
            add_field(paragraph, r'TOC \o "1-3" \h \z \u', "Table of Contents will be updated in Phase 13.")
        elif placeholder_id == "figures":
            paragraph = document.add_paragraph()
            add_field(paragraph, r'TOC \h \z \c "Figure"', "List of Figures will be updated in Phase 13.")
        elif placeholder_id == "tables":
            paragraph = document.add_paragraph()
            add_field(paragraph, r'TOC \h \z \c "Table"', "List of Tables will be updated in Phase 13.")
        add_placeholder(document, placeholder_id, bookmark_id)
        bookmark_id += 1
        document.add_page_break()
    return bookmark_id


def build_skeleton(doc_path: Path = DOC_PATH) -> Path:
    document = Document()
    configure_document(document)
    configure_styles(document)

    bookmark_id = add_front_matter(document, 1)

    for placeholder_id, heading in CHAPTERS:
        document.add_heading(heading, level=1)
        add_placeholder(document, placeholder_id, bookmark_id)
        bookmark_id += 1
        if placeholder_id != "c4":
            document.add_page_break()

    for placeholder_id, heading in IMPLEMENTATION_SECTIONS:
        document.add_heading(heading, level=2)
        add_placeholder(document, placeholder_id, bookmark_id)
        bookmark_id += 1

    document.add_page_break()

    for placeholder_id, heading in BACK_MATTER:
        document.add_heading(heading, level=1)
        add_placeholder(document, placeholder_id, bookmark_id)
        bookmark_id += 1
        document.add_page_break()

    doc_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(doc_path))
    return doc_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build or update the Jarvis documentation book.")
    parser.add_argument(
        "--phase",
        choices=["0", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12", "13"],
        default="0",
        help="Phase to run. Phase 0 rebuilds the skeleton; Phase 1 fills the root package section.",
    )
    args = parser.parse_args()
    if args.phase == "13":
        path = fill_phase13()
        print(f"Filled Phase 13 in {path}")
    elif args.phase == "12":
        path = fill_phase12()
        print(f"Filled Phase 12 in {path}")
    elif args.phase == "11":
        path = fill_phase11()
        print(f"Filled Phase 11 in {path}")
    elif args.phase == "10":
        path = fill_phase10()
        print(f"Filled Phase 10 in {path}")
    elif args.phase == "9":
        path = fill_phase9()
        print(f"Filled Phase 9 in {path}")
    elif args.phase == "8":
        path = fill_phase8()
        print(f"Filled Phase 8 in {path}")
    elif args.phase == "7":
        path = fill_phase7()
        print(f"Filled Phase 7 in {path}")
    elif args.phase == "6":
        path = fill_phase6()
        print(f"Filled Phase 6 in {path}")
    elif args.phase == "5":
        path = fill_phase5()
        print(f"Filled Phase 5 in {path}")
    elif args.phase == "4":
        path = fill_phase4()
        print(f"Filled Phase 4 in {path}")
    elif args.phase == "3":
        path = fill_phase3()
        print(f"Filled Phase 3 in {path}")
    elif args.phase == "2":
        path = fill_phase2()
        print(f"Filled Phase 2 in {path}")
    elif args.phase == "1":
        path = fill_phase1()
        print(f"Filled Phase 1 in {path}")
    else:
        path = build_skeleton()
        print(f"Created {path}")


if __name__ == "__main__":
    main()
